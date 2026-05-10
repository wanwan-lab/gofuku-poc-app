"""
商品在庫・販売管理 (Streamlit)

st.secrets に以下を設定してください（例は .streamlit/secrets.toml）。

必須キー:
  GEMINI_API_KEY
  GAS_UPLOAD_URL           … 画像を Google ドライブに保存する Web アプリ（GAS）の URL
  GAS_API_KEY                … GAS Web アプリ呼び出し用の共有キー（payload の apiKey に付与）
  GOOGLE_DRIVE_FOLDER_ID     … 保存先フォルダID（GAS に渡す）
  google_service_account     … サービスアカウントJSONの各フィールド（[google_service_account] セクション）
    または GOOGLE_SERVICE_ACCOUNT_JSON … JSON文字列1本

スプレッドシート運用時に追加で必須:
  GOOGLE_SPREADSHEET_ID      … 記録用スプレッドシートID（``INVENTORY_SOURCE=csv`` のときは不要で、共有 ``inventory.csv`` を使用）

任意:
  GEMINI_MODEL_NAME          … Gemini モデル ID（未設定時は下記 DEFAULT_GEMINI_MODEL）
  GEMINI_VOUCHER_MODEL_NAME  … 証憑（納品書等）画像解析専用モデル（未設定時は GEMINI_MODEL_NAME と同じ既定）
  GOOGLE_WORKSHEET_NAME      … ワークシート名（未設定時は DEFAULT_WORKSHEET_NAME）
  GAS_UPLOAD_TIMEOUT_SECONDS … GAS への POST タイムアウト秒（既定 300、1〜3600 にクランプ）
  FALLBACK_IMAGE_URL_WHEN_GAS_UNCONFIGURED … GAS 未設定時に台帳の画像URL列へ入れるプレースホルダ URL
  APP_PASSWORD               … アプリ画面の簡易ログイン用（平文。GitHub には secrets.toml をコミットしないこと）
  INVENTORY_SOURCE           … ``csv`` | ``sheet`` 。未指定時は **GOOGLE_SPREADSHEET_ID があるとき sheet**、無いとき **csv**（リポジトリ直下 ``inventory.csv`` または環境変数 ``GOFUKU_INVENTORY_CSV``）。

※ 画像の Gemini 解析は **google-generativeai** を使用します。モデル名は ``GEMINI_MODEL_NAME``（既定は flash 系プレビュー）。
※ **登録（インプット）** ページの証憑取込で、納品書・請求書・領収書を画像・PDF・Excel・Word から解析し入庫（購入）として台帳に追記できます（確定前に表で編集可能）。
※ PDF/Excel/Word 取込には ``pypdf`` / ``pymupdf`` / ``openpyxl`` / ``python-docx`` を使用します（requirements.txt）。
※ アップロード画像は任意。商品写真は Pillow で長辺最大1280px・JPEG品質80に変換してから解析・ドライブ保存します。
※ 証憑ファイルを台帳に確定反映するとき、Drive 保存用には画像のみ長辺最大2000px・JPEG品質75に再圧縮します（PDF/xlsx/docx は原本のまま）。
※ 台帳日時・撮影日時未取得時の現在時刻は **pytz** の ``Asia/Tokyo``（JST）です。

サイドバーで **登録（インプット）** / **ギャラリー（カタログ）** / **在庫一覧** / **集計・分析（ダッシュボード）** を切り替えられます。
在庫データは **共有の inventory.csv**（ローカル）または **Google スプレッドシート**（``INVENTORY_SOURCE`` で選択）に読み書きします。
列定義・CSV 入出力は **app.py 内に内包**しています。

スプレッドシート1行目はヘッダーとして次の列順を想定:
  日時 | 商品名 | 仕入先・取引先 | 数量 | 仕入金額（税抜） | 仕入金額（税込）
  | 販売予定金額（税抜） | 販売予定金額（税込） | 実売金額（税抜） | 実売金額（税込） | 粗利 | ステータス（在庫中/販売済） | メモ（任意） | 在庫カテゴリー（任意・構成比用） | 画像URL | 管理ID | 最後に確認した日付（棚卸日） | 証憑記録日時 | 証憑URL
  | 仕入日時 | 入庫種別 | 浮貸日時 | 販売日時 | 出庫種別
  ※在庫は **1点につき1行** を基本とし、台帳の **数量** 列は主に入出庫集計用です（未入力・空は **1** として扱います）。
  ※写真は **1枚まで** アップロードできます。写真があるときは1回だけドライブに保存し、**複数行を同時に登録する** ときは **全行に同じ画像URL** を入れます。
  ※「在庫カテゴリー」は分析・構成比用の列です。**手動の仕入れ登録では必須**、空欄の行はキャッシュ JSON や和装向けキーワードで補完されます（在庫一覧の AI 一括で台帳列を埋められます）。
  ※「管理ID」列は自動採番（例: G00000001）のシリアルです。既存行の末尾に列を追加しても列位置はずれません。
  ※「**日時**」列（A列）は **その行が最後に台帳へ保存された時点の JST 時刻**（登録・販売反映・一覧からの保存など）です。**仕入日時** は仕入の暦（EXIF 等を ``record_datetime`` に渡した値）、**入庫種別** は登録画面の区分（入庫（購入）・入庫（浮貸）等）です。**販売日時**・**出庫種別** は販売確定時に記録します（顧客返品・戻入は **販売管理** の出庫区分で記録します）。
  ※旧シートの「入出庫種別」列は読み込み時に **入庫種別** へ移して無視します（ヘッダーは新列順に更新されます）。
  ※「仕入金額（税抜）」「仕入金額（税込）」は **1点あたりの行合計**（台帳の各行は1点）です。
  ※旧シートに「仕入単価（税抜）」列が残っている場合は、読み込み時にその列を除いて新しい列構成に揃えます。
  ※新規登録画面では仕入金額（税込）の計算に使う消費税を **10% / 8% / 非課税** から選べます（既定は10%）。
  ※「販売予定金額（税抜）」「実売金額（税抜）」には **1点あたりの税抜金額** を保存し、税込総額列はその税抜行合計に、仕入行と同じ税率で四捨五入します。
  ※金額列（仕入〜粗利まで）は書き込み時に表示形式 **#,##0** を適用します。
  ※粗利は税抜ベースで「販売済」なら実売金額（税抜）−原価、「在庫中」なら販売予定金額（税抜）−原価（いずれも1点あたり）。台帳保存時に再計算します。
  ※「最後に確認した日付（棚卸日）」は棚卸作業用の任意列です（YYYY-MM-DD 推奨）。1人棚卸しの進捗把握に使います。
  ※ **販売管理** タブで管理ID（G########）を指定して出庫を記録します。出庫（販売）・出庫（除外）・出庫（浮貸）で **販売済** にするときは実売が必須です（除外で **対象外** のときは実売不要）。出庫（浮貸）・出庫（戻入）で **在庫中** のままにするときは **浮貸日時** に確定時の JST（または手入力）を記録します。**出庫（返品）** は **販売済** 行を **在庫中** に戻し、販売日時・出庫種別を更新します。
  ※「証憑記録日時」は証憑取込の **確定ボタンを押した JST 時刻**（recorded_at に相当）。「証憑URL」はその証憑を GAS 経由で Drive に保存したときの表示 URL（evidence_url）です。
  ※台帳一覧から手動で在庫行を販売済に編集する場合は、**販売日時**・**出庫種別**・実売・ステータスを整合させてください（保存時に変更があった行の「日時」は自動で更新されます）。
"""

from __future__ import annotations

import contextlib
import base64
import difflib
import hashlib
import io
import json
import math
import os
import re
import uuid
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable, Iterator

import altair as alt
import numpy as np
import google.generativeai as genai
import gspread
import pandas as pd
import pytz
import requests
import streamlit as st
from google.oauth2 import service_account
from PIL import Image, ImageOps

# --- スプレッドシート列・税率（app 単体で完結：サブモジュール未コミットでも Cloud で動く） ---
COL_DATETIME = "日時"
# 旧シート互換（読み込み時のみ。列は台帳から廃止）
LEGACY_COL_MOVEMENT_TYPE = "入出庫種別"
COL_QTY = "数量"
LEGACY_COL_SALE_SOURCE_MGMT_ID = "販売元管理ID"
COL_NAME = "商品名"
COL_SUPPLIER = "仕入先・取引先"
COL_PRICE_EXCL = "仕入金額（税抜）"
COL_PRICE_INCL = "仕入金額（税込）"
LEGACY_COL_UNIT_PRICE = "仕入単価（税抜）"
COL_PLANNED_SALE = "販売予定金額（税抜）"
COL_PLANNED_SALE_INCL = "販売予定金額（税込）"
COL_ACTUAL_SALE = "実売金額（税抜）"
COL_ACTUAL_SALE_INCL = "実売金額（税込）"
COL_GROSS_PROFIT = "粗利"
COL_STOCK_STATUS = "ステータス（在庫中/販売済）"
COL_IMAGE_URL = "画像URL"
COL_SALE_IMAGE_URL = "販売画像URL"
COL_MEMO = "メモ"
COL_CATEGORY = "在庫カテゴリー"
COL_MANAGEMENT_ID = "管理ID"
COL_LAST_STOCKTAKE = "最後に確認した日付（棚卸日）"
COL_LOAN_DATETIME = "浮貸日時"
# 証憑取込（recorded_at / evidence_url に相当）
COL_VOUCHER_RECORDED_AT = "証憑記録日時"
COL_VOUCHER_EVIDENCE_URL = "証憑URL"
# 仕入と販売を同一行で分離（「日時」は行の最終更新時刻。仕入の暦は下記の仕入日時・入庫種別）
COL_PURCHASE_DATETIME = "仕入日時"
COL_PURCHASE_MOVEMENT = "入庫種別"
COL_SALE_DATETIME = "販売日時"
COL_SALE_OUTBOUND_TYPE = "出庫種別"

STATUS_IN_STOCK = "在庫中"
STATUS_SOLD = "販売済"
STATUS_EXCLUDED = "対象外"
STOCK_STATUS_OPTIONS: tuple[str, ...] = (
    STATUS_IN_STOCK,
    STATUS_SOLD,
    STATUS_EXCLUDED,
)

OUTBOUND_KIND_EXCLUDE = "出庫（除外）"
OUTBOUND_KIND_DISPOSAL_LEGACY = "出庫（処分）"


def _is_outbound_exclude_kind(kind: str) -> bool:
    """出庫（除外）および台帳に残る旧ラベル「出庫（処分）」を同一区分として扱う。"""
    k = (kind or "").strip()
    return k == OUTBOUND_KIND_EXCLUDE or k == OUTBOUND_KIND_DISPOSAL_LEGACY


def _movement_is_outbound(mv: str) -> bool:
    """入出庫種別が出庫（販売・浮貸など）かどうか。"""
    return (mv or "").strip().startswith("出庫")


CONSUMPTION_TAX_RATE = 0.10
CONSUMPTION_TAX_CHOICE_TO_RATE: dict[str, float] = {
    "10%": 0.10,
    "8%": 0.08,
    "非課税": 0.0,
}

EXPECTED_HEADERS: list[str] = [
    COL_DATETIME,
    COL_NAME,
    COL_SUPPLIER,
    COL_QTY,
    COL_PRICE_EXCL,
    COL_PRICE_INCL,
    COL_PLANNED_SALE,
    COL_PLANNED_SALE_INCL,
    COL_ACTUAL_SALE,
    COL_ACTUAL_SALE_INCL,
    COL_GROSS_PROFIT,
    COL_STOCK_STATUS,
    COL_MEMO,
    COL_CATEGORY,
    COL_IMAGE_URL,
    COL_SALE_IMAGE_URL,
    COL_MANAGEMENT_ID,
    COL_LAST_STOCKTAKE,
    COL_VOUCHER_RECORDED_AT,
    COL_VOUCHER_EVIDENCE_URL,
    COL_PURCHASE_DATETIME,
    COL_PURCHASE_MOVEMENT,
    COL_LOAN_DATETIME,
    COL_SALE_DATETIME,
    COL_SALE_OUTBOUND_TYPE,
]

_INVENTORY_CSV_DEFAULT_NAME = "inventory.csv"


def _inventory_csv_path() -> Path:
    raw = (os.environ.get("GOFUKU_INVENTORY_CSV") or "").strip()
    if raw:
        return Path(raw).expanduser().resolve()
    return Path(__file__).resolve().parent / _INVENTORY_CSV_DEFAULT_NAME


@st.cache_data(show_spinner=False)
def _inventory_csv_read_df_cached(mtime_ns: int, path_str: str) -> pd.DataFrame:
    """CSV の内容を mtime 付きでキャッシュ（同一内容の連続再描画を軽くする）。"""
    _ = mtime_ns
    path = Path(path_str)
    if not path.exists():
        return pd.DataFrame(columns=EXPECTED_HEADERS)
    df = pd.read_csv(path, encoding="utf-8-sig")
    if LEGACY_COL_MOVEMENT_TYPE in df.columns:
        if COL_PURCHASE_MOVEMENT not in df.columns:
            df[COL_PURCHASE_MOVEMENT] = ""
        m = df[COL_PURCHASE_MOVEMENT].fillna("").astype(str).str.strip() == ""
        df.loc[m, COL_PURCHASE_MOVEMENT] = (
            df.loc[m, LEGACY_COL_MOVEMENT_TYPE].fillna("").astype(str)
        )
        df = df.drop(columns=[LEGACY_COL_MOVEMENT_TYPE], errors="ignore")
    for c in EXPECTED_HEADERS:
        if c not in df.columns:
            df[c] = ""
    return df[EXPECTED_HEADERS].copy()


def _inventory_csv_read_df() -> pd.DataFrame:
    path = _inventory_csv_path()
    if not path.exists():
        return pd.DataFrame(columns=EXPECTED_HEADERS)
    try:
        mt = int(path.stat().st_mtime_ns)
    except OSError:
        mt = 0
    return _inventory_csv_read_df_cached(mt, str(path.resolve()))


def _inventory_csv_write_df(df: pd.DataFrame) -> None:
    path = _inventory_csv_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    out = df.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    out.to_csv(path, index=False, encoding="utf-8-sig")


# --- st.secrets のキー名（文字列リテラルの散在を避ける） ---
SECRET_GEMINI_API_KEY = "GEMINI_API_KEY"
SECRET_GEMINI_MODEL_NAME = "GEMINI_MODEL_NAME"
SECRET_GEMINI_VOUCHER_MODEL_NAME = "GEMINI_VOUCHER_MODEL_NAME"
SECRET_GAS_UPLOAD_URL = "GAS_UPLOAD_URL"
SECRET_GAS_API_KEY = "GAS_API_KEY"
SECRET_GAS_UPLOAD_TIMEOUT_SECONDS = "GAS_UPLOAD_TIMEOUT_SECONDS"
SECRET_GOOGLE_DRIVE_FOLDER_ID = "GOOGLE_DRIVE_FOLDER_ID"
SECRET_GOOGLE_SPREADSHEET_ID = "GOOGLE_SPREADSHEET_ID"
SECRET_GOOGLE_WORKSHEET_NAME = "GOOGLE_WORKSHEET_NAME"
SECRET_GOOGLE_SERVICE_ACCOUNT_JSON = "GOOGLE_SERVICE_ACCOUNT_JSON"
SECRET_GOOGLE_SERVICE_ACCOUNT_SECTION = "google_service_account"
SECRET_APP_PASSWORD = "APP_PASSWORD"
SECRET_INVENTORY_SOURCE = "INVENTORY_SOURCE"
SECRET_FALLBACK_IMAGE_URL_WHEN_GAS_UNCONFIGURED = "FALLBACK_IMAGE_URL_WHEN_GAS_UNCONFIGURED"

# --- secrets に無いときの既定（非機密のデフォルトのみ） ---
DEFAULT_GEMINI_MODEL = "gemini-3-flash-preview"
DEFAULT_WORKSHEET_NAME = "在庫履歴"
DEFAULT_GAS_FALLBACK_IMAGE_URL = "https://example.com/?gofuku-app=skipped-no-gas-secrets"
DEFAULT_GAS_UPLOAD_TIMEOUT_SECONDS = 300

# --- 画像アップロード前処理 ---
UPLOAD_JPEG_MAX_LONG_EDGE = 1280
# Gemini 画像解析・写真照合API向け（トークン削減・コスト最適化）
GEMINI_ANALYSIS_MAX_LONG_EDGE = 1024
GEMINI_ANALYSIS_JPEG_QUALITY = 78
GEMINI_IMAGE_ANALYSIS_CACHE_ENTRIES = 128
UPLOAD_JPEG_QUALITY = 80
# 仕入れ登録タブから Drive へ保存する商品写真（長辺・品質）
PURCHASE_DRIVE_JPEG_MAX_LONG_EDGE = 2000
PURCHASE_DRIVE_JPEG_QUALITY = 75

# --- 証憑ファイルを Google ドライブへ送る直前の軽量化（画像のみ） ---
VOUCHER_DRIVE_JPEG_MAX_LONG_EDGE = 2000
VOUCHER_DRIVE_JPEG_QUALITY = 75

SHEET_AMOUNT_NUMBER_PATTERN = "#,##0"
TZ_JP = pytz.timezone("Asia/Tokyo")
LEDGER_DATA_EDITOR_KEY = "inventory_ledger_data_editor"
INV_GALLERY_PAGE_SIZE = 30
STOCKTAKE_CAND_PAGE_SIZE = 5
# 棚卸し登録: AI が返す候補の最大件数（UI は STOCKTAKE_CAND_PAGE_SIZE 件ずつページング）
STOCKTAKE_CAND_AI_MAX = 40
# 棚卸し登録: 表記ゆれ・洋服・雑貨でも候補を拾うため既定をやや低め（無関係行はプロンプトで除外指示）
STOCKTAKE_CAND_MIN_CONFIDENCE = 0.14
# 棚卸しAI照合: 台帳コンテキストの最大行数（対象リスト内で上限）
STOCKTAKE_AI_CONTEXT_MAX_LINES = 2000
# 一括棚卸の台帳保存を分割（API タイムアウト・行数制限の回避）
STOCKTAKE_SHEET_SAVE_MAX_IDS = 40
# 販売管理AI照合: 台帳コンテキストの最大行数（在庫中を広く拾うため通常より大きめ）
SALES_AI_CONTEXT_MAX_LINES = 2000
# 棚卸 multiselect: ボタンで値を変えるときウィジェットキーと競合しないよう pending キーに取り直す
_PENDING_STOCKTAKE_ASSIST_BATCH_MIDS = "_pending_stocktake_assist_batch_mids"
_PENDING_STOCKTAKE_MULTI_DONE_MIDS = "_pending_stocktake_multi_done_mids"
SESSION_KEY_INV_SHEET_CACHE_BUST = "_inv_sheet_cache_bust"
# 在庫一覧: 棚卸し「今回の対象リスト」を台帳フォルダに JSON で永続化（アプリ終了後も維持）
STOCKTAKE_WORK_SESSION_FILENAME = "stocktake_work_session.json"
# 分析ダッシュボード: 商品名＋仕入先をキーにしたカテゴリー推定キャッシュ（.gitignore の *.json でコミットされない想定）
INVENTORY_CATEGORY_CACHE_FILENAME = "inventory_category_cache.json"
# 移行前の session_state キー（読み込み時にファイルへ移す）
_SESSION_KEY_STOCKTAKE_WORK_REMAINING_LEGACY = "_inv_stocktake_work_remaining_mids"
VOUCHER_DATA_EDITOR_KEY = "voucher_inventory_preview_editor"
LEDGER_PICK_PLACEHOLDER = "（選ばない）"


# 写真→台帳照合（仕入 AI・棚卸し・販売の写真紐付け）向け。和装専門店以外・雑貨・アパレルでも迷いにくくする。
_GEMINI_LEDGER_PHOTO_MATCH_GUIDANCE_JA = """
照合の考え方（業種は限定しない）:
- 在庫は **和服に限らず** 洋服・帽子・バッグ・アクセ・雑貨・アパレル全般・一点物があり得る。写真の品がリストのどれに相当するか、**色・シルエット・素材感・ブランド表記・タグ**と、API に渡る台帳抜粋の **管理ID・商品名・在庫カテゴリー・仕入先名（メーカー名）・在庫状態** で推測する（**メモ・金額・原価・税・連絡先は API には送らない**）。**どの行を選ぶかの判断に、台帳の金額や写真から読める価格が一致するかは使わない。**
- **飲料・カップ麺などパッケージにロゴや商品名が大きく写る品**は文字・形状で照合しやすい。**無包装の衣料・帽子・布小物**は似た見た目が多く判別が難しい。首元・ウエスト・内側の **ケアタグ・サイズ・品番・紙タグ・下札** が写っていれば必ず読み取り、リストの **商品名** と突き合わせる。
- タグが読めない衣料では、**同じ仕入先名で商品名の語が写真の品とかぶる行**を候補に含めてよい。**無関係な別仕入先**の行は入れない。
- 商品名が **略称・英字・カタカナ・型番のみ** で、写真の見え方と文字が違っても同一在庫と判断できるならその management_id を選ぶ。
- 柄は **和柄だけでなく** 無地・ストライプ・チェック・ロゴ・プリント等も手がかりにする。
- 衣料・帽子で確信が低いときは **候補数を増やし** confidence を **0.12〜0.35** 程度まで下げてよい（明らかに無関係な行は入れない）。
""".strip()

# 証憑取込（Gemini への共通指示・JSON 仕様）
VOUCHER_EXTRACTION_RULES = """注意点:
- 商品名が曖昧な場合（例：着物一式）は、入力内の他の情報から「振袖」や「訪問着」などのカテゴリーを推測し、items[].category と name の両方に反映できるよう整理してください。
- 手書きの修正や合計金額がある場合は、最新の数値を優先してください。
- 出力は、プログラムで解析可能な純粋な JSON オブジェクトのみとしてください（説明文・Markdown・コードフェンスは禁止）。"""

VOUCHER_JSON_SPEC = """JSON の構造（キーは次のみ。値の型を守ること）:
{
  "supplier_name": "仕入先名（文字列。読めなければ空文字）",
  "purchase_date": "YYYY-MM-DD（暦日のみ。読めなければ空文字）",
  "items": [
    {
      "name": "商品名",
      "quantity": 整数（1以上。読めなければ1）,
      "unit_price": 1点または当該明細行の税抜単価・金額（円の整数。読めなければ null）,
      "category": "推定区分（振袖・訪問着・帯など。不明なら空文字）"
    }
  ]
}
unit_price は原則として税抜の円整数。明細が税込のみのときは税抜に換算して整数で入れてください。"""


def check_password() -> bool:
    """認証済みになるまで在庫アプリ本体を起動しない。未認証時は認証UIのみ表示し st.stop() する。"""
    if st.session_state.get("authenticated"):
        return True

    st.set_page_config(page_title="認証", layout="centered")
    st.header("認証")

    raw_pw = st.secrets.get(SECRET_APP_PASSWORD)
    if raw_pw is None:
        st.error(
            f"`.streamlit/secrets.toml` に **{SECRET_APP_PASSWORD}** を設定してください。"
            "（ローカルで `secrets.toml` が無い場合は作成してください）"
        )
        st.stop()
    expected = str(raw_pw).strip()

    if not expected:
        st.error(f"{SECRET_APP_PASSWORD} が空です。secrets.toml を確認してください。")
        st.stop()

    with st.form("auth_screen_form"):
        password = st.text_input("パスワード", type="password")
        submitted = st.form_submit_button(
            "ログイン（パスワード入力後は Enter でも送信できます）"
        )

    if submitted:
        if (password or "").strip() == expected:
            st.session_state["authenticated"] = True
            st.rerun()
        st.error("パスワードが正しくありません。")

    st.stop()


def _apply_inventory_amount_number_formats(ws) -> None:
    """金額系の列（仕入金額（税抜）〜粗利まで）に、2行目以降で #,##0 を適用する。"""
    idx_start = EXPECTED_HEADERS.index(COL_PRICE_EXCL)
    idx_end = EXPECTED_HEADERS.index(COL_GROSS_PROFIT)
    end_row = max(int(ws.row_count), 2)
    ws.spreadsheet.batch_update(
        {
            "requests": [
                {
                    "repeatCell": {
                        "range": {
                            "sheetId": ws.id,
                            "startRowIndex": 1,
                            "endRowIndex": end_row,
                            "startColumnIndex": idx_start,
                            "endColumnIndex": idx_end + 1,
                        },
                        "cell": {
                            "userEnteredFormat": {
                                "numberFormat": {
                                    "type": "NUMBER",
                                    "pattern": SHEET_AMOUNT_NUMBER_PATTERN,
                                }
                            }
                        },
                        "fields": "userEnteredFormat.numberFormat",
                    }
                }
            ]
        }
    )


def jst_now() -> datetime:
    """現在の日本時間（JST・timezone-aware）。"""
    return datetime.now(TZ_JP)


def jst_now_str() -> str:
    """スプレッドシート用の日時文字列（JST・秒まで）。"""
    return jst_now().strftime("%Y-%m-%d %H:%M:%S")


def capture_datetime_jst_from_bytes(raw: bytes) -> str | None:
    """画像バイナリの EXIF から撮影日時を読み、JST の壁時計として解釈して文字列化する。

    リサイズ前の元データに対して呼ぶこと（EXIF 失効前に取得する）。
    EXIF にタイムゾーンが無いため、取得値は **日本のローカル時刻** として ``Asia/Tokyo`` に固定する。
    失敗時は ``None``（呼び出し側で ``jst_now_str()`` をデフォルトにする）。
    """
    try:
        with Image.open(io.BytesIO(raw)) as img:
            exif = img.getexif()
        if not exif:
            return None
        dt_s: str | None = None
        try:
            from PIL.ExifTags import IFD

            sub = exif.get_ifd(IFD.Exif)
            if sub:
                dt_s = sub.get(36867) or sub.get(36868)  # DateTimeOriginal, DateTimeDigitized
            if not dt_s:
                sub0 = exif.get_ifd(IFD.IFD0)
                if sub0:
                    dt_s = sub0.get(306)  # DateTime
        except Exception:
            pass
        if not dt_s:
            dt_s = exif.get(36867) or exif.get(306)
        if not dt_s:
            return None
        if isinstance(dt_s, bytes):
            dt_s = dt_s.decode("utf-8", errors="ignore")
        dt_s = str(dt_s).strip()
        naive: datetime | None = None
        for fmt in ("%Y:%m:%d %H:%M:%S", "%Y-%m-%d %H:%M:%S"):
            try:
                naive = datetime.strptime(dt_s, fmt)
                break
            except ValueError:
                continue
        if naive is None:
            return None
        aware = TZ_JP.localize(naive)
        return aware.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return None


def capture_datetime_jst_from_upload(uploaded) -> str | None:
    """UploadedFile から EXIF 日時を取得（内部は :func:`capture_datetime_jst_from_bytes`）。"""
    try:
        return capture_datetime_jst_from_bytes(uploaded.getvalue())
    except Exception:
        return None


def _resize_long_edge_max(img: Image.Image, max_edge: int) -> Image.Image:
    w, h = img.size
    long_edge = max(w, h)
    if long_edge <= max_edge:
        return img
    scale = max_edge / float(long_edge)
    nw = max(1, int(round(w * scale)))
    nh = max(1, int(round(h * scale)))
    return img.resize((nw, nh), Image.Resampling.LANCZOS)


def prepare_upload_image_jpeg(
    raw: bytes,
    *,
    max_long_edge: int | None = None,
    quality: int | None = None,
) -> tuple[bytes, str]:
    """Gemini 送信用・GAS 保存用の共通前処理。

    EXIF 向き補正のうえ、長辺リサイズと JPEG 再エンコードを行う。
    ``max_long_edge`` / ``quality`` を省略したときは :data:`UPLOAD_JPEG_MAX_LONG_EDGE` と
    :data:`UPLOAD_JPEG_QUALITY` を使う（Gemini 向けの既定）。

    Returns:
        (jpeg_bytes, mime_type)  mime_type は常に ``image/jpeg`` 。
    """
    mx = int(max_long_edge) if max_long_edge is not None else UPLOAD_JPEG_MAX_LONG_EDGE
    q = int(quality) if quality is not None else UPLOAD_JPEG_QUALITY
    mx = max(256, min(mx, 8192))
    q = max(40, min(q, 95))
    img = Image.open(io.BytesIO(raw))
    img = ImageOps.exif_transpose(img)
    rgba = img.convert("RGBA")
    bg = Image.new("RGB", rgba.size, (255, 255, 255))
    bg.paste(rgba, mask=rgba.getchannel("A"))
    img = bg
    img = _resize_long_edge_max(img, mx)
    buf = io.BytesIO()
    img.save(
        buf,
        format="JPEG",
        quality=q,
        optimize=True,
        progressive=True,
    )
    return buf.getvalue(), "image/jpeg"


def _secret_str(key: str, default: str = "") -> str:
    """設定文字列を ``st.secrets.get`` で取得。欠損・空は default。"""
    v = st.secrets.get(key, default)
    if v is None:
        return default
    s = str(v).strip()
    return s if s else default


def _uses_local_inventory_csv() -> bool:
    """``INVENTORY_SOURCE=csv`` なら True。``sheet`` 明示なら False。未指定時はスプレッドシートIDが無ければ CSV。"""
    v = _secret_str(SECRET_INVENTORY_SOURCE, "").lower()
    if v == "csv":
        return True
    if v == "sheet":
        return False
    return not _secret_str(SECRET_GOOGLE_SPREADSHEET_ID)


def _secret_int(
    key: str, default: int, *, min_value: int = 1, max_value: int = 10_000
) -> int:
    raw = st.secrets.get(key)
    if raw is None:
        return default
    s = str(raw).strip()
    if not s:
        return default
    try:
        return max(min_value, min(max_value, int(s)))
    except ValueError:
        return default


def _gas_upload_timeout_seconds() -> int:
    return _secret_int(
        SECRET_GAS_UPLOAD_TIMEOUT_SECONDS,
        DEFAULT_GAS_UPLOAD_TIMEOUT_SECONDS,
        min_value=30,
        max_value=3600,
    )


def _fallback_image_url_when_gas_unconfigured() -> str:
    return _secret_str(
        SECRET_FALLBACK_IMAGE_URL_WHEN_GAS_UNCONFIGURED,
        DEFAULT_GAS_FALLBACK_IMAGE_URL,
    )


def _gemini_model_name() -> str:
    return _secret_str(SECRET_GEMINI_MODEL_NAME, DEFAULT_GEMINI_MODEL)


def _gemini_voucher_model_name() -> str:
    """証憑画像解析用モデル。専用キーが空なら通常の Gemini モデル名にフォールバック。"""
    v = _secret_str(SECRET_GEMINI_VOUCHER_MODEL_NAME, "")
    return v if v else _gemini_model_name()


def _load_service_account_info() -> dict[str, Any]:
    raw_json = st.secrets.get(SECRET_GOOGLE_SERVICE_ACCOUNT_JSON)
    if raw_json is not None and str(raw_json).strip():
        if isinstance(raw_json, str):
            return json.loads(raw_json)
        raise ValueError(
            f"{SECRET_GOOGLE_SERVICE_ACCOUNT_JSON} は文字列である必要があります。"
        )
    ga = st.secrets.get(SECRET_GOOGLE_SERVICE_ACCOUNT_SECTION)
    if ga is not None:
        if hasattr(ga, "to_dict"):
            return dict(ga.to_dict())
        return dict(ga)
    raise ValueError(
        "サービスアカウントが見つかりません。"
        f" [{SECRET_GOOGLE_SERVICE_ACCOUNT_SECTION}] セクションか "
        f"{SECRET_GOOGLE_SERVICE_ACCOUNT_JSON} を設定してください。"
    )


def _credentials():
    info = _load_service_account_info()
    scopes = ("https://www.googleapis.com/auth/spreadsheets",)
    return service_account.Credentials.from_service_account_info(info, scopes=scopes)


@st.cache_resource
def _gspread_client():
    return gspread.authorize(_credentials())


def _parse_json_from_model(text: str) -> dict[str, Any]:
    """モデル出力から JSON オブジェクトを抽出する（コードフェンスや前後の説明文を許容）。"""
    t = text.strip()
    fence = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", t)
    if fence:
        t = fence.group(1).strip()
    try:
        obj = json.loads(t)
    except json.JSONDecodeError:
        i0 = t.find("{")
        i1 = t.rfind("}")
        if i0 == -1 or i1 <= i0:
            raise
        obj = json.loads(t[i0 : i1 + 1])
    if isinstance(obj, list) and obj and isinstance(obj[0], dict):
        obj = obj[0]
    if not isinstance(obj, dict):
        raise ValueError("JSON がオブジェクト形式ではありません。")
    return obj


def _coerce_positive_int(val: Any, default: int = 1) -> int:
    try:
        n = int(float(str(val).strip()))
        return max(1, n)
    except Exception:
        return default


def _coerce_unit_price_yen(val: Any) -> int | None:
    """税抜単価（円）を整数にする。不明・null・空なら None（既存入力を上書きしない）。"""
    if val is None:
        return None
    s = str(val).strip()
    if not s or s.lower() in ("null", "none", "-", "不明"):
        return None
    s = re.sub(r"[,\s円￥¥]", "", s, flags=re.UNICODE)
    s = re.sub(r"(?i)yen", "", s)
    if not s or not re.search(r"\d", s):
        return None
    try:
        n = int(round(float(s)))
        return max(1, n)
    except Exception:
        return None


def _apply_gemini_json_to_session(
    result: dict[str, Any],
    df_ledger: pd.DataFrame | None = None,
) -> None:
    """Gemini の JSON をフォーム用 session_state に反映する（英日キー両対応）。"""
    r = result
    st.session_state.pop("_gemini_match_management_id", None)
    _qty_g = _coerce_positive_int(
        r.get("quantity") or r.get("数量") or r.get("qty") or 1,
        default=1,
    )
    st.session_state.field_row_quantity = min(2000, _qty_g)

    m = r.get("match")
    row_hit: pd.Series | None = None
    if isinstance(m, dict) and df_ledger is not None and not df_ledger.empty:
        mid_hit = str(m.get("management_id") or m.get("管理ID") or "").strip()
        if mid_hit and COL_MANAGEMENT_ID in df_ledger.columns:
            mask_id = df_ledger[COL_MANAGEMENT_ID].astype(str).str.strip() == mid_hit
            sub_hit = df_ledger.loc[mask_id]
            if len(sub_hit) == 1:
                row_hit = sub_hit.iloc[0]
                st.session_state["_gemini_match_management_id"] = mid_hit

    pn = str(
        r.get("product_name")
        or r.get("商品名")
        or r.get("name")
        or ""
    ).strip()
    su = str(
        r.get("supplier")
        or r.get("仕入先・取引先")
        or r.get("仕入先")
        or r.get("取引先")
        or r.get("vendor")
        or ""
    ).strip()
    if row_hit is not None:
        rpn = str(row_hit.get(COL_NAME, "") or "").strip()
        rsu = str(row_hit.get(COL_SUPPLIER, "") or "").strip()
        if rpn:
            pn = rpn
        if rsu:
            su = rsu
    match_conf_ok = isinstance(m, dict) and float(m.get("confidence") or 0) >= 0.75
    if match_conf_ok:
        if not pn:
            pn = str(m.get("product_name") or "").strip()
        if not su:
            su = str(m.get("supplier") or "").strip()
    if pn:
        st.session_state.field_product_name = pn
    if su:
        st.session_state.field_supplier = su

    line_yen: int | None = None
    if row_hit is not None:
        ly0 = _finite_int(row_hit.get(COL_PRICE_EXCL), 0)
        if ly0 > 0:
            line_yen = ly0
    if line_yen is None:
        line_yen = _coerce_unit_price_yen(
            r.get("line_price_excl")
            or r.get("line_excl_yen")
            or r.get("仕入金額（税抜）")
            or r.get("unit_price_excl")
            or r.get("unit_price")
            or r.get("product_unit_price_excl")
            or r.get("単価")
            or r.get("単価（税抜）")
        )
    if line_yen is None and match_conf_ok and isinstance(m, dict):
        line_yen = _coerce_unit_price_yen(
            m.get("line_price_excl")
            or m.get("unit_price_excl")
            or m.get("unit_price")
        )
    if line_yen is not None:
        st.session_state.field_line_excl_yen = line_yen

    kind = str(
        r.get("product_kind")
        or r.get("種類")
        or r.get("type")
        or r.get("商品カテゴリ")
        or ""
    ).strip()
    if not kind and pn:
        kind = pn
    st.session_state.ai_kind = kind

    vf = r.get("visual_features")
    if not vf:
        parts = [
            r.get(k)
            for k in (
                "色",
                "柄",
                "素材",
                "状態",
                "色柄",
                "備考",
                "color",
                "pattern",
                "material",
                "condition",
            )
            if r.get(k)
        ]
        vf = " / ".join(str(p) for p in parts)
    st.session_state.ai_features = str(vf or "")
    st.session_state.ai_parse_ran = True

    ic = str(
        r.get("inventory_category")
        or r.get("在庫カテゴリー")
        or r.get("stock_category")
        or r.get("category_label")
        or ""
    ).strip()
    if row_hit is not None and COL_CATEGORY in row_hit.index:
        rc = str(row_hit.get(COL_CATEGORY, "") or "").strip()
        if rc:
            ic = rc
    if isinstance(m, dict) and match_conf_ok:
        mic = str(m.get("inventory_category") or m.get("在庫カテゴリー") or "").strip()
        if mic and not ic:
            ic = mic
    if ic:
        st.session_state.field_inventory_category = ic[:80]


def _apply_gemini_sale_link_to_session(
    result: dict[str, Any],
    df_ledger: pd.DataFrame | None,
    *,
    fill_product_preview_fields: bool = True,
    restrict_to_sold: bool = False,
    restrict_to_float_loan_outbound: bool = False,
) -> None:
    """販売管理の「販売する管理ID」欄へ、写真照合で得た管理IDを session_state に反映する。

    ``restrict_to_sold=True`` … **販売済** の行のみ有効（出庫（返品）の写真照合用）。
    ``restrict_to_float_loan_outbound=True`` … **在庫中かつ出庫種別が出庫（浮貸）** の行のみ（出庫（戻入）の写真照合用）。
    """
    st.session_state.pop("_sale_link_management_id", None)
    st.session_state.pop("_sale_link_warn", None)
    r = result
    m = r.get("match")
    if not isinstance(m, dict):
        m = {}
    mid = str(
        m.get("management_id")
        or m.get("管理ID")
        or r.get("management_id")
        or ""
    ).strip()
    if (
        not mid
        and df_ledger is not None
        and not df_ledger.empty
    ):
        pn0 = str(
            m.get("product_name")
            or r.get("product_name")
            or r.get("商品名")
            or ""
        ).strip()
        su0 = str(
            m.get("supplier")
            or r.get("supplier")
            or r.get("仕入先・取引先")
            or r.get("仕入先")
            or ""
        ).strip()
        fr = _single_row_fuzzy_ledger_match(
            df_ledger,
            pn0,
            su0,
            only_float_loan_outbound=restrict_to_float_loan_outbound,
            only_in_stock=not restrict_to_sold and not restrict_to_float_loan_outbound,
            only_sold=restrict_to_sold,
            limit=14,
        )
        if fr is not None:
            mid = str(fr.get(COL_MANAGEMENT_ID, "") or "").strip()
            if mid:
                m = {
                    **m,
                    "management_id": mid,
                    "confidence": max(
                        float(m.get("confidence") or 0), 0.78
                    ),
                }
    conf = float(m.get("confidence") or r.get("confidence") or 0)
    row_hit: pd.Series | None = None
    if (
        mid
        and df_ledger is not None
        and not df_ledger.empty
        and COL_MANAGEMENT_ID in df_ledger.columns
    ):
        mask_mid = df_ledger[COL_MANAGEMENT_ID].astype(str).str.strip() == mid
        if restrict_to_float_loan_outbound:
            hits_ok = df_ledger.loc[
                mask_mid & _mask_ledger_in_stock_outbound_float_loan(df_ledger)
            ]
        else:
            hits_ok = df_ledger.loc[
                mask_mid
                & (
                    _mask_ledger_sold(df_ledger)
                    if restrict_to_sold
                    else _mask_ledger_in_stock(df_ledger)
                )
            ]
        if len(hits_ok) == 1:
            row_hit = hits_ok.iloc[0]
    if not mid:
        return
    if row_hit is None:
        if (
            df_ledger is not None
            and not df_ledger.empty
            and COL_MANAGEMENT_ID in df_ledger.columns
        ):
            hits_any = df_ledger.loc[
                df_ledger[COL_MANAGEMENT_ID].astype(str).str.strip() == mid
            ]
            if len(hits_any) == 1:
                stt_bad = _normalize_stock_status(
                    str(hits_any.iloc[0].get(COL_STOCK_STATUS, ""))
                )
                need = (
                    "出庫種別が出庫（浮貸）の在庫中"
                    if restrict_to_float_loan_outbound
                    else ("販売済" if restrict_to_sold else "在庫中")
                )
                st.session_state["_sale_link_warn"] = (
                    f"管理ID {mid} は「{stt_bad}」のためこの出庫区分では使えません。"
                    f"照合対象は **{need}** の行のみです。"
                )
                return
        st.session_state.field_sale_source_mgmt_id = mid
        st.session_state["_sale_link_warn"] = (
            f"管理ID {mid} が台帳に見つかりません。表記を確認するか手入力してください。"
        )
        return
    st.session_state.field_sale_source_mgmt_id = mid
    st.session_state["_sale_link_management_id"] = mid
    if fill_product_preview_fields and conf >= 0.72:
        rpn = str(row_hit.get(COL_NAME, "") or "").strip()
        rsu = str(row_hit.get(COL_SUPPLIER, "") or "").strip()
        if rpn:
            st.session_state.field_product_name = rpn
        if rsu:
            st.session_state.field_supplier = rsu
        ly = _finite_int(row_hit.get(COL_PRICE_EXCL), 0)
        if ly > 0:
            st.session_state.field_line_excl_yen = ly


def _gemini_input_image_from_upload(uploaded) -> Image.Image:
    """Gemini 写真照合用に長辺を抑えた JPEG と同等の RGB PIL を返す（Drive 保存用とは解像度が異なる）。"""
    jpeg_bytes, _ = prepare_upload_image_jpeg(
        uploaded.getvalue(),
        max_long_edge=GEMINI_ANALYSIS_MAX_LONG_EDGE,
        quality=GEMINI_ANALYSIS_JPEG_QUALITY,
    )
    return Image.open(io.BytesIO(jpeg_bytes)).convert("RGB")


def _pil_image_for_gemini(image_data: Any) -> Image.Image:
    """Gemini へ渡す PIL（RGB・EXIF 補正済み）。"""
    if isinstance(image_data, Image.Image):
        im = ImageOps.exif_transpose(image_data)
        return im.convert("RGB")
    if isinstance(image_data, (bytes, bytearray, memoryview)):
        im = Image.open(io.BytesIO(bytes(image_data)))
        im = ImageOps.exif_transpose(im)
        return im.convert("RGB")
    gv = getattr(image_data, "getvalue", None)
    if callable(gv):
        return _pil_image_for_gemini(gv())
    raise TypeError(
        "image_data は PIL.Image / bytes / getvalue() を持つオブジェクトである必要があります。"
    )


def _gemini_analysis_jpeg_bytes(image_data: Any) -> bytes:
    """Gemini 画像解析API向けに長辺を抑えた JPEG bytes（キャッシュキー・入力兼用）。"""
    if isinstance(image_data, (bytes, bytearray, memoryview)):
        raw = bytes(image_data)
    elif hasattr(image_data, "getvalue") and callable(getattr(image_data, "getvalue")):
        raw = image_data.getvalue()
    else:
        pil = _pil_image_for_gemini(image_data)
        buf = io.BytesIO()
        pil.save(buf, format="PNG")
        raw = buf.getvalue()
    return prepare_upload_image_jpeg(
        raw,
        max_long_edge=GEMINI_ANALYSIS_MAX_LONG_EDGE,
        quality=GEMINI_ANALYSIS_JPEG_QUALITY,
    )[0]


def _gemini_image_analysis_session_key(
    jpeg_bytes: bytes,
    inventory_context: str,
    prompt_mode: str,
    model_name: str,
) -> str:
    jh = hashlib.sha256(jpeg_bytes).hexdigest()
    ih = hashlib.sha256(inventory_context.encode("utf-8")).hexdigest()[:24]
    safe_mode = re.sub(r"[^a-zA-Z0-9_.-]", "_", prompt_mode)
    safe_model = re.sub(r"[^a-zA-Z0-9_.-]", "_", model_name)
    return f"_gemini_imgtxt_{jh}_{ih}_{safe_mode}_{safe_model}"


@st.cache_data(
    max_entries=GEMINI_IMAGE_ANALYSIS_CACHE_ENTRIES,
    show_spinner=False,
    hash_funcs={bytes: lambda b: hashlib.sha256(b).hexdigest()},
)
def _cached_gemini_product_image_analysis(
    jpeg_bytes: bytes,
    inventory_context: str,
    prompt_mode: str,
    model_name: str,
) -> str:
    """商品写真の Gemini 解析（同一 JPEG・同一コンテキストではキャッシュヒットし API を呼ばない）。"""
    api_key = _secret_str(SECRET_GEMINI_API_KEY)
    if not api_key:
        raise RuntimeError(
            f"{SECRET_GEMINI_API_KEY} が設定されていません。`.streamlit/secrets.toml` を確認してください。"
        )
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    subject = Image.open(io.BytesIO(jpeg_bytes)).convert("RGB")

    inv_stripped = (inventory_context or "").strip()
    inv_tail_for_match = ""
    if inv_stripped:
        inv_tail_for_match = f"""
次のリストは、すでに台帳にある行の抜粋です（**在庫中・販売済などステータス付き**。最大約400件。写真と同一・類似の商品がありそうなら必ず照合してください）。
各行には **管理ID・商品名・仕入先名・在庫カテゴリー・在庫状態** のみが含まれます（**原価・税込・連絡先・メモなど機密・金額情報は API に送信していません**。台帳に数値列があっても、**行の選び方に金額の一致は使わないこと**）。
{_GEMINI_LEDGER_PHOTO_MATCH_GUIDANCE_JA}

{inv_stripped}

照合するときは必ず "match" オブジェクトを付け、少なくとも次を含めてください:
- "management_id" (string): 上のリストにある行の **管理ID** と完全一致する値。リストには **在庫中・販売済などすべてのステータス** が含まれます。**写真と同一・類似の台帳上の1行** に対応するときは、その行の管理IDを選ぶこと（仕入れの入力補助のため。ステータスで候補を除外しない）。該当がなければ ""
- "product_name" (string): 台帳の商品名に合わせた確定案（推測でも可）
- "supplier" (string): 台帳の仕入先に合わせた確定案（推測でも可）
- "line_price_excl" (null のみ): API に仕入金額・原価は渡していないため **必ず null**
- "inventory_category" (string): リストの照合先行に **{COL_CATEGORY}** が載っていればその行と同一の文字列。リストに無い・該当行が空なら ""
- "confidence" (number): 0.0〜1.0 で、写真と台帳行が同一在庫である確信度

同一行が見つからない場合は management_id を "" にし、confidence は 0.4 未満にしてください。
"""
    if prompt_mode == "stocktake_match":
        if not inv_stripped:
            raise ValueError(
                "棚卸しの照合には台帳に在庫中の行が必要です（スプレッドシートを確認してください）。"
            )
        prompt = f"""**直後の画像** が、棚卸しのために撮影した **現物1点** です（衣料・アパレル・帽子・雑貨・一点物など **業種を限定しない** ）。
{_GEMINI_LEDGER_PHOTO_MATCH_GUIDANCE_JA}

次の **続きのテキスト** に示すリストは、**今回の棚卸作業でまだ未確認として残っている在庫中の行** に限定します（販売済は含みません。リスト外の管理IDは返さない）。
同じ型・同シリーズ・同仕入れロットで **複数行あることが多い** ので、写真に合いそうな行を **複数** 挙げてください（最大{STOCKTAKE_CAND_AI_MAX}件。アプリでは画面を{STOCKTAKE_CAND_PAGE_SIZE}件ずつに分けて表示します）。
**どの行が写真の品かは、金額の一致では判断しない**（商品名・在庫カテゴリー・仕入先名・タグ・色形などで推測する）。
台帳の「商品名」「仕入先・取引先」と現在画像の見た目・文脈を基に、同一商品の可能性を評価する。
**リストに無い管理IDは絶対に返さない** でください。JSON だけを返す（説明文・Markdown のコードフェンス禁止）。
---
{inv_stripped}

返却形式（キーは次のみ）:
- "stocktake_candidates" (array): 必須。各要素は object で、次のフィールドを持つ:
  - "management_id" (string): リストにある在庫中行の管理ID（G########）
  - "confidence" (number): 0.0〜1.0（写真と同一在庫である確信度。表示順はアプリ側で管理IDの昇順に整列する）
  - "product_name" (string): その行の商品名（参考）
  - "supplier" (string): その行の仕入先（参考）
  - "feature_observation" (string): この画像で見えた特徴の補足（任意）。なければ ""
  該当が1件も無いときは空配列 []。
  迷う場合は複数入れてよい（confidence が低いものも列挙してよい。ただし無関係な行は入れない）。

任意で互換用に "match" (object) を1件だけ付けてもよい（先頭候補と同じ内容でよい）。"""
        response = model.generate_content([prompt, subject])
        return response.text or ""

    if prompt_mode == "sale_link":
        if not inv_stripped:
            raise ValueError(
                "販売元の写真照合には、台帳に **在庫中** の行が少なくとも1行必要です（管理ID・商品名などが入っている行）。"
                "在庫がすべて販売済のときや、台帳の読み込みに失敗しているときは使えません。"
            )
        prompt = f"""**直後の画像** は、**販売時にどの在庫行に対応するか** を特定するための商品写真です（衣料・アパレル・帽子・雑貨など **業種を限定しない** 在庫）。
{_GEMINI_LEDGER_PHOTO_MATCH_GUIDANCE_JA}

次の **続きのテキスト** のリストは台帳の **在庫中** の行だけです（**販売済の行は含めていません**）。**行の選定に金額の一致は使わない。** 必ずこのリストの中からだけ management_id を選べ。
各行の「商品名」「仕入先・取引先」の一致度を優先し、画像全体の見た目で同一商品かを判断する。
JSON だけを返してください（説明文・コードフェンス禁止）。
---
{inv_stripped}

返却形式（キーは次のみ）:
- "match" (object): 必須。フィールド:
  - "management_id" (string): 選んだ行の管理ID（G########）。該当なしなら ""
  - "confidence" (number): 0.0〜1.0
  - "product_name" (string): その行の商品名（参考）
  - "supplier" (string): その行の仕入先（参考）
  - "line_price_excl" (null のみ): API に仕入金額は渡していないため **必ず null**
  - "feature_observation" (string): 現在画像で確認できた特徴の補足（メモ追記用）。なければ ""

該当がなければ management_id を ""、confidence は 0.25 以下にする。"""
        response = model.generate_content([prompt, subject])
        return response.text or ""

    if prompt_mode == "sale_link_sold":
        if not inv_stripped:
            raise ValueError(
                "返品照合には、台帳に **販売済** の行が少なくとも1行必要です。"
            )
        prompt = f"""**直後の画像** は、**販売済のどの行が返品・再入庫の対象か** を特定するための商品写真です。
{_GEMINI_LEDGER_PHOTO_MATCH_GUIDANCE_JA}

次の **続きのテキスト** のリストは台帳の **販売済** の行だけです。**行の選定に金額の一致は使わない。** 必ずこのリストの中からだけ management_id を選べ。
各行の「商品名」「仕入先・取引先」の一致度を優先し、画像全体の見た目で同一商品かを判断する。
JSON だけを返してください（説明文・コードフェンス禁止）。
---
{inv_stripped}

返却形式（キーは次のみ）:
- "match" (object): 必須。フィールド:
  - "management_id" (string): 選んだ行の管理ID（G########）。該当なしなら ""
  - "confidence" (number): 0.0〜1.0
  - "product_name" (string): その行の商品名（参考）
  - "supplier" (string): その行の仕入先（参考）
  - "line_price_excl" (null のみ): API に仕入金額は渡していないため **必ず null**
  - "feature_observation" (string): 現在画像で確認できた特徴の補足。なければ ""

該当がなければ management_id を ""、confidence は 0.25 以下にする。"""
        response = model.generate_content([prompt, subject])
        return response.text or ""

    if prompt_mode == "sale_link_float_loan":
        if not inv_stripped:
            raise ValueError(
                "戻入の写真照合には、台帳に **在庫中かつ出庫種別が出庫（浮貸）** の行が少なくとも1行必要です。"
                "該当がない場合や **管理ID** が空の行しかない場合は使えません。"
            )
        prompt = f"""**直後の画像** は、**出庫（戻入）でどの在庫行に対応するか** を特定するための商品写真です（浮貸からの戻しを紐付けるため）。
{_GEMINI_LEDGER_PHOTO_MATCH_GUIDANCE_JA}

次の **続きのテキスト** のリストは台帳のうち **在庫中** であり、かつ **出庫種別** が **出庫（浮貸）** の行だけです（それ以外の在庫行・販売済は含みません）。**行の選定に金額の一致は使わない。** 必ずこのリストの中からだけ management_id を選べ。
各行の「商品名」「仕入先・取引先」の一致度を優先し、画像全体の見た目で同一商品かを判断する。
JSON だけを返してください（説明文・コードフェンス禁止）。
---
{inv_stripped}

返却形式（キーは次のみ）:
- "match" (object): 必須。フィールド:
  - "management_id" (string): 選んだ行の管理ID（G########）。該当なしなら ""
  - "confidence" (number): 0.0〜1.0
  - "product_name" (string): その行の商品名（参考）
  - "supplier" (string): その行の仕入先（参考）
  - "line_price_excl" (null のみ): API に仕入金額は渡していないため **必ず null**
  - "feature_observation" (string): 現在画像で確認できた特徴の補足（メモ追記用）。なければ ""

該当がなければ management_id を ""、confidence は 0.25 以下にする。"""
        response = model.generate_content([prompt, subject])
        return response.text or ""

    schema_intro = f"""**直後の画像** は **小売・卸の在庫・売買用** の商品写真です（呉服に限らず洋服・帽子・バッグ・雑貨など）。この画像について次のキーだけを持つ JSON オブジェクトを 1 つだけ返してください。
説明文や Markdown のコードフェンスは付けず、JSON のみを出力してください。

必須キー（値の型を守ること）:
- "product_name" (string): 商品名として適切な短い名称。不明なら ""
- "supplier" (string): 仕入先・取引先として推測できる名称。不明なら ""
- "quantity" (integer): 写っている点数・束の本数などの推定。最低 1
- "inventory_category" (string): **在庫カテゴリー**（分析・構成比用の短いラベル。例: 帯、ジャケット、帽子、雑貨、飲料。業種は問わない）。20文字以内。推測できる場合は必ず入れる。本当に不明なら ""
- "product_kind" (string): 種類の推定（例: 振袖、訪問着、帯、ニット、シャツ、キャップ、ワンピース）。不明なら ""
- "color" (string): 色の推定。不明なら ""
- "pattern" (string): 柄の推定。不明なら ""
- "material" (string): 素材の推定。不明なら ""
- "condition" (string): 状態の推定。不明なら ""
- "unit_price_excl" (integer or null): 1点あたりの税抜の仕入金額（円）の推定。相場・品質から読めない場合は null（勝手に 1 にしない）
"""
    schema_footer = f"""任意: 台帳照合結果を "match" にまとめる（上記リストがあるときはできる限り付与）
  例: {{"management_id": "G00000001", "product_name": "…", "supplier": "…", "line_price_excl": null, "inventory_category": "帯", "confidence": 0.85}}
  リストの行に **{COL_CATEGORY}** が載っているときは、照合して "match" に management_id を入れる場合 **必ず** 同じ行の値を "inventory_category" に含める。リストにカテゴリーが無いときのみ省略可。
  **API に台帳の金額・原価は含めていない**ため、"match"."line_price_excl" は **必ず null**。
  不要・該当なしのときは "match" キー自体を省略してもよい。"""
    prompt = schema_intro + inv_tail_for_match + schema_footer
    response = model.generate_content([prompt, subject])
    return response.text or ""


def _consumption_tax_rate_from_choice_label(label: str) -> float:
    return CONSUMPTION_TAX_CHOICE_TO_RATE.get(label, CONSUMPTION_TAX_RATE)


def _finite_int(val: Any, default: int = 0) -> int:
    """NaN / inf / 非数値を default に落とし、有限な int にする（金額・数量用）。"""
    try:
        if val is None:
            return default
        if isinstance(val, (float, np.floating)):
            if not math.isfinite(float(val)) or pd.isna(val):
                return default
        if isinstance(val, str):
            t = val.strip()
            if not t or t.lower() in ("nan", "none", "<na>", "nat"):
                return default
            t = (
                t.replace(",", "")
                .replace("，", "")
                .replace("¥", "")
                .replace("\u00a5", "")
                .strip()
            )
            if not t:
                return default
            val = t
        n = float(pd.to_numeric(val, errors="coerce"))
        if not math.isfinite(n) or pd.isna(n):
            return default
        return int(round(n))
    except (OverflowError, ValueError, TypeError):
        return default


def _series_to_numeric_loose(s: pd.Series) -> pd.Series:
    """カンマ区切り・円記号付きのセルでも数値化する（スプレッドシート取り込み用）。"""
    if not isinstance(s, pd.Series):
        s = pd.Series(s)
    if s.dtype == object or pd.api.types.is_string_dtype(s):
        st = s.where(pd.notna(s), "").astype(str)
        st = st.str.replace(",", "", regex=False)
        st = st.str.replace("，", "", regex=False)
        st = st.str.replace("¥", "", regex=False)
        st = st.str.replace("\u00a5", "", regex=False)
        st = st.str.strip()
        st = st.replace({"nan": "", "None": "", "<NA>": ""})
        return pd.to_numeric(st, errors="coerce")
    return pd.to_numeric(s, errors="coerce")


def price_incl_tax(price_excl_yen: int, tax_rate: float | None = None) -> int:
    """税抜き行金額から税込円金額（四捨五入）。

    Args:
        price_excl_yen: 税抜き金額（円）
        tax_rate: 消費税率（例: 0.1）。None のときは :data:`CONSUMPTION_TAX_RATE`（10%）
    """
    ex = _finite_int(price_excl_yen, 0)
    r_raw = CONSUMPTION_TAX_RATE if tax_rate is None else tax_rate
    r = float(pd.to_numeric(r_raw, errors="coerce"))
    if not math.isfinite(r) or pd.isna(r):
        r = float(CONSUMPTION_TAX_RATE)
    return int(round(ex * (1 + r)))


def _infer_tax_rate_from_main_line(line_excl_yen: int, line_incl_yen: int) -> float:
    """仕入金額（税抜）と仕入金額（税込）から、登録時と同じ消費税区分を推定する。"""
    excl = _finite_int(line_excl_yen, 0)
    incl = _finite_int(line_incl_yen, 0)
    if excl <= 0:
        return float(CONSUMPTION_TAX_RATE)
    if incl <= excl:
        return 0.0
    for _label, rate in CONSUMPTION_TAX_CHOICE_TO_RATE.items():
        if price_incl_tax(excl, float(rate)) == incl:
            return float(rate)
    return float(CONSUMPTION_TAX_RATE)


def _estimate_excl_yen_from_incl_yen(incl_yen: int) -> int:
    """税込行金額から税抜行金額を推定（登録時の税率候補に税込が一致する組を採用）。"""
    incl_i = _finite_int(incl_yen, 0)
    if incl_i <= 0:
        return 0
    seen: set[float] = set()
    for rate in CONSUMPTION_TAX_CHOICE_TO_RATE.values():
        rr = float(rate)
        if rr in seen:
            continue
        seen.add(rr)
        if rr == 0.0:
            return incl_i
        ex = int(round(incl_i / (1 + rr)))
        if price_incl_tax(ex, rr) == incl_i:
            return ex
    return int(round(incl_i / (1 + float(CONSUMPTION_TAX_RATE))))


def _planned_actual_line_amounts(
    qty: int,
    planned_unit_excl: int,
    actual_unit_excl: int,
    status: str,
    tax_rate: float,
) -> tuple[int, int, int, int]:
    """販売予定・実売の税抜行計と税込行計（税抜列×数量を行合計にしてから税込）。"""
    q = max(1, _finite_int(qty, 1))
    pu, au = _finite_int(planned_unit_excl, 0), _finite_int(actual_unit_excl, 0)
    st = _normalize_stock_status(status)
    tr = float(pd.to_numeric(tax_rate, errors="coerce"))
    if not math.isfinite(tr) or pd.isna(tr):
        tr = float(CONSUMPTION_TAX_RATE)
    plex = pu * q if pu > 0 else 0
    pincl = price_incl_tax(plex, tr) if plex > 0 else 0
    aex = (au * q) if (st == STATUS_SOLD and au > 0) else 0
    aincl = price_incl_tax(aex, tr) if aex > 0 else 0
    return plex, pincl, aex, aincl


def analyze_image_with_gemini(
    image_data,
    *,
    inventory_context: str | None = None,
    prompt_mode: str = "full",
) -> str:
    """商品写真の Gemini 解析。画像は長辺約1024pxのJPEGに正規化し、同一内容は ``@st.cache_data`` と session で再利用する。"""
    api_key = _secret_str(SECRET_GEMINI_API_KEY)
    if not api_key:
        raise RuntimeError(
            f"{SECRET_GEMINI_API_KEY} が設定されていません。`.streamlit/secrets.toml` を確認してください。"
        )
    jpeg_bytes = _gemini_analysis_jpeg_bytes(image_data)
    inv_s = (inventory_context or "").strip()
    model_name = _gemini_model_name()
    ss_key = _gemini_image_analysis_session_key(
        jpeg_bytes, inv_s, prompt_mode, model_name
    )
    if ss_key in st.session_state:
        return st.session_state[ss_key]
    out = _cached_gemini_product_image_analysis(
        jpeg_bytes, inv_s, prompt_mode, model_name
    )
    st.session_state[ss_key] = out
    return out

def _voucher_configure_model():
    api_key = _secret_str(SECRET_GEMINI_API_KEY)
    if not api_key:
        raise RuntimeError(
            f"{SECRET_GEMINI_API_KEY} が設定されていません。`.streamlit/secrets.toml` を確認してください。"
        )
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(_gemini_voucher_model_name())


def analyze_voucher_document_with_gemini_images(images: list[Any]) -> str:
    """証憑のページ画像（1枚以上）を Gemini に渡し、supplier_name / purchase_date / items の JSON 文字列を返す。"""
    if not images:
        raise ValueError("画像がありません。")
    model = _voucher_configure_model()
    n = len(images)
    multi = (
        "複数のページ画像が順に渡されています。全体を通して一つの証憑として読み取ってください。"
        if n > 1
        else ""
    )
    prompt = f"""内部システムプロンプト:
あなたはプロの会計士です。提供された呉服店向け証憑（納品書・請求書・領収書）のページ画像を解析し、在庫管理に必要なデータを抽出してください。
{multi}

{VOUCHER_EXTRACTION_RULES}

{VOUCHER_JSON_SPEC}"""
    response = model.generate_content([prompt, *images])
    return response.text or ""


def analyze_voucher_document_with_gemini(image_data) -> str:
    """単一の証憑画像を Gemini に渡す（後方互換・内部は複数画像 API と共通）。"""
    return analyze_voucher_document_with_gemini_images([image_data])


def analyze_voucher_document_with_gemini_text(
    *, source_instruction: str, document_body: str
) -> str:
    """テキスト（PDF抽出・Markdown表・Word全文など）から同一 JSON 形式で抽出する。"""
    model = _voucher_configure_model()
    body = (document_body or "").strip()
    max_chars = 600_000
    if len(body) > max_chars:
        body = body[:max_chars] + "\n\n（長文のため先頭のみ。末尾を省略しました）"
    prompt = f"""内部システムプロンプト:
あなたはプロの会計士です。呉服店向け証憑（納品書・請求書・領収書）に関する次の入力から、在庫管理に必要なデータを抽出してください。

{source_instruction}

{VOUCHER_EXTRACTION_RULES}

{VOUCHER_JSON_SPEC}

--- 入力ここから ---
{body}
--- 入力ここまで ---"""
    response = model.generate_content(prompt)
    return response.text or ""


def _voucher_upload_suffix(filename: str) -> str:
    parts = (filename or "").rsplit(".", 1)
    return parts[-1].lower() if len(parts) == 2 else ""


def prepare_voucher_file_for_drive_storage(
    raw: bytes, original_filename: str
) -> tuple[bytes, str, str]:
    """証憑を Drive 保存用に整形する。JPG/PNG 等は長辺最大2000px・JPEG quality=75。PDF/xlsx/docx は原本のまま。

    Returns:
        (保存バイナリ, mime_type, 拡張子に使う文字列 ``.jpg`` 等)
    """
    suf = _voucher_upload_suffix(original_filename)
    if suf in ("jpg", "jpeg", "png", "webp"):
        img = Image.open(io.BytesIO(raw))
        img = ImageOps.exif_transpose(img)
        rgba = img.convert("RGBA")
        bg = Image.new("RGB", rgba.size, (255, 255, 255))
        bg.paste(rgba, mask=rgba.getchannel("A"))
        img = bg
        img = _resize_long_edge_max(img, VOUCHER_DRIVE_JPEG_MAX_LONG_EDGE)
        buf = io.BytesIO()
        img.save(
            buf,
            format="JPEG",
            quality=VOUCHER_DRIVE_JPEG_QUALITY,
            optimize=True,
            progressive=True,
        )
        return buf.getvalue(), "image/jpeg", ".jpg"
    mime_map = {
        "pdf": "application/pdf",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    mime = mime_map.get(suf, "application/octet-stream")
    ext = f".{suf}" if suf else ""
    return raw, mime, ext


def _voucher_drive_safe_filename(
    purchase_date: str, supplier: str, original_filename: str, ext_with_dot: str
) -> str:
    """``[仕入日]_[仕入先]_[元ファイル名].拡子`` 形式（ファイル名向けに禁則文字を除去）。"""
    d = (purchase_date or "").strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", d):
        d = jst_now().strftime("%Y-%m-%d")
    d_part = d.replace("-", "")
    sup = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", (supplier or "").strip())[:50].strip(
        "_"
    ) or "unknown"
    base = Path(str(original_filename)).name
    stem = Path(base).stem
    safe_stem = re.sub(
        r'[^\w\-_.\u3040-\u30ff\u4e00-\u9fff]', "_", stem
    ).strip("._")[:100] or "voucher"
    ext = (
        ext_with_dot
        if ext_with_dot.startswith(".")
        else (f".{ext_with_dot}" if ext_with_dot else "")
    )
    return f"{d_part}_{sup}_{safe_stem}{ext}"


def _gas_evidence_upload_ready() -> bool:
    """証憑を Google ドライブ（GAS 経由）に送るのに必要な secrets が揃っているか。"""
    return bool(
        _secret_str(SECRET_GAS_UPLOAD_URL)
        and _secret_str(SECRET_GAS_API_KEY)
        and _secret_str(SECRET_GOOGLE_DRIVE_FOLDER_ID)
    )


def _voucher_extract_pdf_text(pdf_bytes: bytes) -> str:
    """pypdf で PDF からプレーンテキストを抽出する。"""
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
    except Exception as e:
        raise ValueError(f"PDF を開けませんでした: {e}") from e
    if getattr(reader, "is_encrypted", False):
        try:
            auth = reader.decrypt("")
        except Exception as e:
            raise ValueError("パスワード付きPDFには未対応です。") from e
        if auth == 0:
            raise ValueError("パスワード付きPDFには未対応です。")
    chunks: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        chunks.append(t)
    return "\n".join(chunks).strip()


def _voucher_pdf_pages_to_images(
    pdf_bytes: bytes, *, max_pages: int = 15, zoom: float = 2.0
) -> list[Image.Image]:
    """スキャン PDF 等をページ画像にレンダリングする（pymupdf）。"""
    import fitz

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        n = min(doc.page_count, max_pages)
        if n <= 0:
            raise ValueError("PDF にページがありません。")
        out: list[Image.Image] = []
        for i in range(n):
            page = doc.load_page(i)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            out.append(
                Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            )
        return out
    finally:
        doc.close()


def _dataframe_to_markdown_table(df: pd.DataFrame, *, max_rows: int = 500) -> str:
    """pandas DataFrame を簡易 Markdown 表にする（外部パッケージ不要）。"""
    df2 = df.fillna("").head(max_rows).copy()
    cols = [str(c).replace("|", "/") for c in df2.columns]
    header = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join("---" for _ in cols) + " |"
    lines = [header, sep]
    for _, row in df2.iterrows():
        cells = [
            str(v).replace("|", "\\|").replace("\n", " ").strip()
            for v in row.tolist()
        ]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _voucher_excel_bytes_to_markdown(xlsx_bytes: bytes) -> str:
    """Excel を全シート読み、Markdown 形式のテキストにまとめる。"""
    bio = io.BytesIO(xlsx_bytes)
    try:
        xl = pd.ExcelFile(bio, engine="openpyxl")
    except Exception as e:
        raise ValueError(
            f"Excel（.xlsx）の読み込みに失敗しました。openpyxl 対応形式か確認してください: {e}"
        ) from e
    parts: list[str] = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(xl, sheet_name=sheet, header=0, engine="openpyxl")
        if df.empty:
            parts.append(f"## シート: {sheet}\n\n（データ行なし）\n")
        else:
            parts.append(
                f"## シート: {sheet}\n\n{_dataframe_to_markdown_table(df)}\n"
            )
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Excel に読み取れるシートがありません。")
    return text


def _voucher_docx_bytes_to_text(docx_bytes: bytes) -> str:
    """Word（.docx）から段落および表セルのテキストを抽出する。"""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    table_lines: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [
                cell.text.strip().replace("\n", " ") for cell in row.cells
            ]
            table_lines.append("\t".join(cells))
    blocks = paras + table_lines
    return "\n".join(blocks).strip()


class _VoucherRawBytesUpload:
    """``_gemini_input_image_from_upload`` 用の最小インターフェース。"""

    __slots__ = ("_data",)

    def __init__(self, data: bytes) -> None:
        self._data = data

    def getvalue(self) -> bytes:
        return self._data


def analyze_voucher_upload_bytes(raw: bytes, filename: str) -> str:
    """アップロード種別に応じて分岐し、Gemini から JSON 文字列を返す。"""
    suf = _voucher_upload_suffix(filename)
    if suf in ("jpg", "jpeg", "png", "webp"):
        img = _gemini_input_image_from_upload(_VoucherRawBytesUpload(raw))
        return analyze_voucher_document_with_gemini(img)
    if suf == "pdf":
        text = _voucher_extract_pdf_text(raw)
        if text:
            return analyze_voucher_document_with_gemini_text(
                source_instruction=(
                    "以下は PDF からテキスト抽出した内容です。"
                    "この請求・納品・領収に相当する情報から在庫データを抽出してください。"
                ),
                document_body=text,
            )
        imgs = _voucher_pdf_pages_to_images(raw)
        return analyze_voucher_document_with_gemini_images(imgs)
    if suf == "xlsx":
        md = _voucher_excel_bytes_to_markdown(raw)
        return analyze_voucher_document_with_gemini_text(
            source_instruction=(
                "以下は Excel（.xlsx）を Markdown 表に変換したものです。"
                "この表から仕入・納品に相当する在庫情報を抽出してください。"
            ),
            document_body=md,
        )
    if suf == "docx":
        t = _voucher_docx_bytes_to_text(raw)
        if not t:
            raise ValueError("Word 文書からテキストを抽出できませんでした。")
        return analyze_voucher_document_with_gemini_text(
            source_instruction=(
                "以下は Word（.docx）から抽出した全文です。"
                "この文書内の請求・納品・領収に相当する情報から在庫データを抽出してください。"
            ),
            document_body=t,
        )
    raise ValueError(
        f"未対応の拡張子です（.jpg / .png / .pdf / .xlsx / .docx のみ）: .{suf}"
    )


def _merge_voucher_items_for_preview(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """同一商品名（大文字小文字・前後空白無視）の数量を合算し、後から現れる単価で上書き（手書き修正の優先に近づける）。"""
    order: dict[str, dict[str, Any]] = {}
    key_order: list[str] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        name = str(it.get("name") or "").strip()
        if not name:
            continue
        key = name.casefold()
        q = _coerce_positive_int(it.get("quantity"), 1)
        up = _coerce_unit_price_yen(it.get("unit_price"))
        if up is None:
            up = _coerce_unit_price_yen(it.get("unit_price_excl"))
        cat = str(it.get("category") or "").strip()
        if key not in order:
            order[key] = {"商品名": name, "数量": q, "単価（税抜）": up, "カテゴリ": cat}
            key_order.append(key)
        else:
            row = order[key]
            row["数量"] = int(row["数量"]) + q
            if up is not None:
                row["単価（税抜）"] = int(up)
            if cat:
                row["カテゴリ"] = cat
    out: list[dict[str, Any]] = []
    for k in key_order:
        r = order[k]
        up = r["単価（税抜）"]
        if up is None:
            up = 1
        out.append(
            {
                "商品名": r["商品名"],
                "数量": max(1, int(r["数量"])),
                "単価（税抜）": max(1, int(up)),
                "カテゴリ": str(r.get("カテゴリ") or ""),
            }
        )
    return out


def _voucher_record_datetime_jst(purchase_date_str: str) -> str:
    """証憑の仕入日を台帳の日時列用に整形。YYYY-MM-DD のときはその日の JST 正午、それ以外は現在 JST。"""
    s = (purchase_date_str or "").strip()
    if re.fullmatch(r"\d{4}-\d{2}-\d{2}", s):
        return f"{s} 12:00:00"
    return jst_now_str()


def _init_voucher_sidebar_state() -> None:
    if "voucher_supplier_edit" not in st.session_state:
        st.session_state.voucher_supplier_edit = ""
    if "voucher_date_edit" not in st.session_state:
        st.session_state.voucher_date_edit = ""
    if "voucher_consumption_tax_choice" not in st.session_state:
        st.session_state.voucher_consumption_tax_choice = "10%"


def _confirm_voucher_import(
    df: pd.DataFrame | None,
    supplier: str,
    purchase_date_str: str,
    tax_choice_label: str,
) -> None:
    """証憑プレビュー表の内容を 1点1行で台帳に追記する。"""
    if df is None or df.empty:
        st.error("反映する行がありません。")
        return
    rows_spec: list[tuple[str, int, int, str]] = []
    for _, row in df.iterrows():
        name = str(row.get("商品名", "") or "").strip()
        if not name:
            continue
        q = max(1, _finite_int(row.get("数量"), 1))
        up = max(1, _finite_int(row.get("単価（税抜）"), 1))
        cat = str(row.get("カテゴリ", "") or "").strip()
        rows_spec.append((name, q, up, cat))
    if not rows_spec:
        st.error("商品名が入っている行がありません。")
        return
    total_q = sum(q for _, q, _, _ in rows_spec)
    ws: Any = None
    if not _uses_local_inventory_csv():
        ws = ensure_worksheet_header()
        if ws is None:
            st.warning("スプレッドシート未設定のため保存できません。")
            return
    tax_r = _consumption_tax_rate_from_choice_label(str(tax_choice_label))
    rec_dt = _voucher_record_datetime_jst(purchase_date_str)
    sup = (supplier or "").strip()
    recorded_at = jst_now_str()
    evidence_url = ""
    stash_b = st.session_state.get("voucher_stash_bytes")
    stash_name = str(st.session_state.get("voucher_stash_name") or "voucher")
    if stash_b is not None:
        if _gas_evidence_upload_ready():
            try:
                blob, mime, ext_dot = prepare_voucher_file_for_drive_storage(
                    bytes(stash_b), stash_name
                )
                fname = _voucher_drive_safe_filename(
                    recorded_at[:10], sup, stash_name, ext_dot
                )
                evidence_url = upload_image_to_drive(fname, mime, blob)
                fb = _fallback_image_url_when_gas_unconfigured()
                if not (evidence_url or "").strip() or (evidence_url or "").strip() == (
                    fb or ""
                ).strip():
                    evidence_url = ""
                    st.warning(
                        "ドライブの表示 URL を取得できませんでした。台帳には証憑URLなしで記録します。"
                    )
            except Exception as e:
                st.error(f"証憑の Google ドライブ保存に失敗しました: {e}")
                return
        else:
            st.warning(
                "GAS_UPLOAD_URL / GAS_API_KEY / GOOGLE_DRIVE_FOLDER_ID が未設定のため、"
                "証憑ファイルは Drive に保存されません（台帳のみ反映）。"
            )
    try:
        ids = allocate_management_ids(ws, total_q)
    except Exception as e:
        st.error(f"管理IDの採番に失敗しました: {e}")
        return
    idx = 0
    try:
        with st.spinner("台帳に書き込んでいます…"):
            _v_dt = jst_now_str()
            batch_rows: list[list[Any]] = []
            for name, q, up, cat in rows_spec:
                memo = f"証憑取込 category={cat}" if cat else "証憑取込"
                incl = price_incl_tax(up, tax_r)
                _v_cat = str(cat).strip()[:80] if cat else ""
                for _ in range(q):
                    batch_rows.append(
                        _inventory_row_values_for_append(
                            _v_dt,
                            rec_dt,
                            "入庫（購入）",
                            name,
                            sup,
                            up,
                            incl,
                            "",
                            ids[idx],
                            memo,
                            inventory_category=_v_cat,
                            consumption_tax_rate=tax_r,
                            voucher_recorded_at=recorded_at,
                            voucher_evidence_url=evidence_url,
                        )
                    )
                    idx += 1
            _append_inventory_data_rows(batch_rows)
    except Exception as e:
        st.error(f"台帳の更新に失敗しました: {e}")
        return
    st.session_state.pop("voucher_preview_df", None)
    st.session_state.pop(VOUCHER_DATA_EDITOR_KEY, None)
    st.session_state.pop("voucher_stash_bytes", None)
    st.session_state.pop("voucher_stash_name", None)
    st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
    _ev = (evidence_url or "").strip()
    _msg = f"証憑取込を記録しました（{total_q} 行・1点1行）。証憑記録日時: {recorded_at}"
    if _ev:
        _msg += " 証憑を Drive に保存済みです。"
    st.session_state["_voucher_import_flash"] = _msg
    st.rerun()


def _render_voucher_inventory_panel() -> None:
    """登録ページ内: 証憑ファイルのアップロード・解析・プレビュー編集・確定反映。"""
    _vflash = st.session_state.pop("_voucher_import_flash", None)
    if _vflash:
        st.success(_vflash)
    with st.expander("証憑から在庫反映（帳票から取込）", expanded=False):
        st.caption(
            "納品書・請求書・領収書を画像・PDF・Excel・Word から読み取り、入庫（購入）として台帳に反映します。"
            "PDF はテキスト優先、空ならページ画像として解析します。解析後は表で修正してから確定してください。"
        )
        if not _secret_str(SECRET_GEMINI_API_KEY):
            st.info(f"{SECRET_GEMINI_API_KEY} が未設定のため使えません。")
            return
        if not _uses_local_inventory_csv() and not _secret_str(SECRET_GOOGLE_SPREADSHEET_ID):
            st.info(
                "在庫の保存先が未設定です。`GOOGLE_SPREADSHEET_ID` を設定するか、"
                "`INVENTORY_SOURCE = \"csv\"` で共有の inventory.csv を使ってください。"
            )
            return

        voucher_up = st.file_uploader(
            "証憑ファイル（画像 / PDF / Excel / Word）",
            type=["jpg", "jpeg", "png", "pdf", "xlsx", "docx"],
            key="voucher_file_uploader",
        )
        if st.button(
            "証憑を解析",
            key="voucher_analyze_btn",
            disabled=voucher_up is None,
        ):
            with st.spinner("証憑を解析しています…"):
                try:
                    raw = analyze_voucher_upload_bytes(
                        voucher_up.getvalue(), voucher_up.name
                    )
                    d = _parse_json_from_model(raw)
                    items = d.get("items")
                    if not isinstance(items, list) or not items:
                        st.error(
                            "items が見つかりませんでした。ファイル内容を確認してください。"
                        )
                    else:
                        merged = _merge_voucher_items_for_preview(
                            [x for x in items if isinstance(x, dict)]
                        )
                        if not merged:
                            st.error("有効な商品行がありません。")
                        else:
                            st.session_state.voucher_supplier_edit = str(
                                d.get("supplier_name") or ""
                            ).strip()
                            st.session_state.voucher_date_edit = str(
                                d.get("purchase_date") or ""
                            ).strip()
                            st.session_state.voucher_preview_df = pd.DataFrame(merged)
                            st.session_state.pop(VOUCHER_DATA_EDITOR_KEY, None)
                            st.session_state["voucher_stash_bytes"] = voucher_up.getvalue()
                            st.session_state["voucher_stash_name"] = voucher_up.name
                            st.success(
                                f"解析しました（{len(merged)} 商品行）。内容を確認して確定してください。"
                            )
                except Exception as e:
                    st.warning(str(e))

        st.text_input(
            "仕入先（証憑・上書き可）",
            key="voucher_supplier_edit",
            placeholder="仕入先・取引先",
        )
        st.text_input(
            "仕入日（YYYY-MM-DD・上書き可）",
            key="voucher_date_edit",
            placeholder="例: 2026-04-15",
        )
        st.radio(
            "証憑取込の消費税（税込計算）",
            options=list(CONSUMPTION_TAX_CHOICE_TO_RATE.keys()),
            horizontal=True,
            key="voucher_consumption_tax_choice",
        )

        base_df = st.session_state.get("voucher_preview_df")
        if base_df is not None and not base_df.empty:
            edited_df = st.data_editor(
                base_df,
                num_rows="dynamic",
                key=VOUCHER_DATA_EDITOR_KEY,
                use_container_width=True,
            )
            if st.button(
                "この内容で台帳に確定反映",
                type="primary",
                key="voucher_confirm_btn",
            ):
                _confirm_voucher_import(
                    edited_df,
                    str(st.session_state.get("voucher_supplier_edit", "") or ""),
                    str(st.session_state.get("voucher_date_edit", "") or ""),
                    str(
                        st.session_state.get("voucher_consumption_tax_choice", "10%")
                        or "10%"
                    ),
                )
            if st.button("プレビューをクリア", key="voucher_clear_preview_btn"):
                st.session_state.pop("voucher_preview_df", None)
                st.session_state.pop(VOUCHER_DATA_EDITOR_KEY, None)
                st.session_state.pop("voucher_stash_bytes", None)
                st.session_state.pop("voucher_stash_name", None)
                st.rerun()


def _open_inventory_workbook():
    sid = _secret_str(SECRET_GOOGLE_SPREADSHEET_ID)
    if not sid:
        return None
    try:
        return _gspread_client().open_by_key(str(sid))
    except Exception:
        return None


def _get_or_create_inventory_worksheet():
    """在庫ワークシートを開く。無ければ十分な行・列で作成する。失敗時は None。"""
    sh = _open_inventory_workbook()
    if sh is None:
        return None
    wname = _secret_str(SECRET_GOOGLE_WORKSHEET_NAME, DEFAULT_WORKSHEET_NAME)
    try:
        try:
            return sh.worksheet(str(wname))
        except gspread.WorksheetNotFound:
            return sh.add_worksheet(
                title=str(wname),
                rows=2000,
                cols=max(20, len(EXPECTED_HEADERS) + 2),
            )
    except Exception:
        return None


def _bump_inventory_sheet_cache_bust() -> None:
    """スプレッドシート由来の get_all_values キャッシュを無効化する（追記・全置換・手動再読込後）。"""
    try:
        st.session_state[SESSION_KEY_INV_SHEET_CACHE_BUST] = int(
            st.session_state.get(SESSION_KEY_INV_SHEET_CACHE_BUST, 0)
        ) + 1
    except Exception:
        pass


@st.cache_data(show_spinner=False)
def _inventory_sheet_get_all_values_cached(
    sheet_id: str, worksheet_title: str, bust: int
) -> list[list[Any]]:
    """同一 bust の間は get_all_values の結果を再利用する（bust は書き込み・再読込で進める）。

    失敗時は **例外を送出**する。``return None`` だと Streamlit の ``@st.cache_data`` に
    失敗結果がキャッシュされ、一時的な通信エラー後も台帳が読めなくなるため。
    """
    _ = bust
    try:
        sh = _gspread_client().open_by_key(str(sheet_id))
    except Exception as e:
        raise RuntimeError(
            "スプレッドシートを開けません。"
            f"{SECRET_GOOGLE_SPREADSHEET_ID}・共有権限・[{SECRET_GOOGLE_SERVICE_ACCOUNT_SECTION}] を確認してください。"
            f" 詳細: {e}"
        ) from e
    try:
        try:
            ws = sh.worksheet(str(worksheet_title))
        except gspread.WorksheetNotFound:
            ws = sh.add_worksheet(
                title=str(worksheet_title),
                rows=2000,
                cols=max(20, len(EXPECTED_HEADERS) + 2),
            )
    except Exception as e:
        raise RuntimeError(
            f"ワークシート「{worksheet_title}」を開けず、新規作成にも失敗しました: {e}"
        ) from e
    try:
        return ws.get_all_values()
    except Exception as e:
        raise RuntimeError(
            f"ワークシート「{worksheet_title}」の get_all_values に失敗しました: {e}"
        ) from e


def ensure_worksheet_header():
    """1行目がヘッダーでなければ作成（初回のみ想定）。secrets 未設定時は None。"""
    ws = _get_or_create_inventory_worksheet()
    if ws is None:
        return None
    try:
        first = ws.row_values(1)
        if not first or first[: len(EXPECTED_HEADERS)] != EXPECTED_HEADERS:
            # ヘッダー差分があると追記列が右にずれて保存されるため、既存データ行も
            # 旧ヘッダー名ベースで EXPECTED_HEADERS 順へ正規化してから書き戻す。
            raw = ws.get_all_values()
            values: list[list[Any]] = [EXPECTED_HEADERS]
            if raw:
                header0 = [("" if c is None else str(c)).strip() for c in raw[0]]
                for rr in raw[1:]:
                    values.append(_sheet_header_row_to_expected_list(header0, list(rr)))
            ws.clear()
            ws.update("A1", values, value_input_option="USER_ENTERED")
            try:
                _apply_inventory_amount_number_formats(ws)
            except Exception:
                pass
            _bump_inventory_sheet_cache_bust()
        return ws
    except Exception:
        return None


def upload_image_to_drive(filename: str, mime: str, data: bytes) -> str:
    """Google Apps Script 経由で Google ドライブに保存し、閲覧用 URL を返す。

    GAS 側は JSON で
    { folderId, fileName, mimeType, base64Data, apiKey } を受け取り、
    { status: \"success\", url: \"...\" } 形式で返す想定。
    """
    gas_url = _secret_str(SECRET_GAS_UPLOAD_URL)
    gas_api_key = _secret_str(SECRET_GAS_API_KEY)
    folder_id = _secret_str(SECRET_GOOGLE_DRIVE_FOLDER_ID)
    if not gas_url or not gas_api_key or not folder_id:
        return _fallback_image_url_when_gas_unconfigured()

    base64_data = base64.b64encode(data).decode("ascii")
    payload = {
        "folderId": folder_id,
        "fileName": filename,
        "mimeType": mime,
        "base64Data": base64_data,
        "apiKey": gas_api_key,
    }

    resp = requests.post(
        gas_url,
        json=payload,
        headers={"Content-Type": "application/json"},
        timeout=_gas_upload_timeout_seconds(),
    )
    resp.raise_for_status()

    try:
        body = resp.json()
    except json.JSONDecodeError as e:
        raise RuntimeError(f"GAS の応答が JSON ではありません: {resp.text[:500]}") from e

    if body.get("status") != "success":
        raise RuntimeError(
            body.get("message") or body.get("error") or f"GAS アップロード失敗: {body!r}"
        )

    url = body.get("url") or body.get("webViewLink") or body.get("fileUrl")
    if not url:
        raise RuntimeError(f"GAS 応答に URL がありません: {body!r}")
    return str(url)


def _optional_amount_cell(yen: int) -> int:
    """0 以下は 0（数値列・スプレッドシートともに空欄相当）。"""
    v = _finite_int(yen, 0)
    return max(0, v)


def _int_from_cell(v: Any) -> int:
    """セル値を有限な int に（計算前の正規化用）。"""
    return _finite_int(v, 0)


def _coerce_money_columns_for_recalc(df: pd.DataFrame) -> pd.DataFrame:
    """数値列を ``pd.to_numeric`` で揃え、inf/NaN を 0 にして int 化する。"""
    out = df.copy()
    money_cols = (
        COL_PRICE_EXCL,
        COL_PRICE_INCL,
        COL_PLANNED_SALE,
        COL_PLANNED_SALE_INCL,
        COL_ACTUAL_SALE,
        COL_ACTUAL_SALE_INCL,
        COL_GROSS_PROFIT,
    )
    for c in money_cols:
        if c not in out.columns:
            continue
        s = (
            _series_to_numeric_loose(out[c])
            .replace([np.inf, -np.inf], np.nan)
            .fillna(0)
        )
        x = pd.to_numeric(s, errors="coerce").to_numpy(dtype=np.float64, copy=False)
        x = np.where(np.isfinite(x), np.rint(x), 0.0)
        out[c] = x.astype(np.int64)
    if COL_QTY in out.columns:
        s = (
            _series_to_numeric_loose(out[COL_QTY])
            .replace([np.inf, -np.inf], np.nan)
            .fillna(1)
        )
        xq = pd.to_numeric(s, errors="coerce").to_numpy(dtype=np.float64, copy=False)
        xq = np.where(np.isfinite(xq), np.rint(xq), 1.0)
        xq = np.maximum(1.0, xq)
        out[COL_QTY] = xq.astype(np.int64)
    return out


def _normalize_stock_status(status: str) -> str:
    s = (status or "").strip()
    return s if s in STOCK_STATUS_OPTIONS else STATUS_IN_STOCK


def _ledger_stocktake_dates_parsed(df: pd.DataFrame) -> pd.Series:
    if COL_LAST_STOCKTAKE not in df.columns:
        return pd.Series(pd.NaT, index=df.index, dtype="datetime64[ns]")
    return pd.to_datetime(df[COL_LAST_STOCKTAKE], errors="coerce")


def _today_jst_date() -> date:
    return datetime.now(TZ_JP).date()


def _mask_ledger_in_stock(df: pd.DataFrame) -> pd.Series:
    if COL_STOCK_STATUS not in df.columns:
        return pd.Series(False, index=df.index)
    return (
        df[COL_STOCK_STATUS].astype(str).str.strip().map(_normalize_stock_status)
        == STATUS_IN_STOCK
    )


def _mask_ledger_sold(df: pd.DataFrame) -> pd.Series:
    """台帳行が **販売済** かどうか（在庫一覧・販売補助用）。"""
    if COL_STOCK_STATUS not in df.columns:
        return pd.Series(False, index=df.index)
    return (
        df[COL_STOCK_STATUS].astype(str).str.strip().map(_normalize_stock_status)
        == STATUS_SOLD
    )


def _mask_ledger_in_stock_outbound_float_loan(df: pd.DataFrame) -> pd.Series:
    """在庫中かつ **出庫種別** が **出庫（浮貸）** の行（出庫（戻入）の検索・候補用）。"""
    if COL_STOCK_STATUS not in df.columns:
        return pd.Series(False, index=df.index)
    m_in = _mask_ledger_in_stock(df)
    if COL_SALE_OUTBOUND_TYPE not in df.columns:
        return pd.Series(False, index=df.index)
    ot = df[COL_SALE_OUTBOUND_TYPE].astype(str).str.strip()
    return m_in & (ot == "出庫（浮貸）")


def _mask_ledger_stocktake_unverified(df: pd.DataFrame) -> pd.Series:
    """在庫中かつ棚卸日が空または解釈不能な行。"""
    m_in = _mask_ledger_in_stock(df)
    dt = _ledger_stocktake_dates_parsed(df)
    return m_in & dt.isna()


def _mask_ledger_stocktake_today_jst(df: pd.DataFrame) -> pd.Series:
    """在庫中かつ棚卸日が今日（JST）の行。"""
    if df.empty or COL_LAST_STOCKTAKE not in df.columns:
        return pd.Series(False, index=df.index)
    m_in = _mask_ledger_in_stock(df)
    today_s = _today_jst_date().isoformat()
    tok = df[COL_LAST_STOCKTAKE].map(_stocktake_date_token_for_compare)
    return m_in & (tok == today_s) & (tok != "")


def _count_stocktake_today_jst_in_management_ids(
    df: pd.DataFrame, mids: set[str]
) -> int:
    """指定した管理 ID のうち、在庫中かつ棚卸日が今日（JST）の行数。"""
    if df.empty or not mids or COL_MANAGEMENT_ID not in df.columns:
        return 0
    m = _mask_ledger_stocktake_today_jst(df) & df[COL_MANAGEMENT_ID].astype(str).str.strip().isin(
        mids
    )
    return int(m.sum())


def _stocktake_date_token_for_compare(val: Any) -> str:
    """棚卸日セルの値を **JST の暦日**（ISO）に正規化。空・解釈不能は空文字。

    サーバーが UTC のときでも、スプレッドシートの日付列と「今日（JST）」の一致判定がずれないようにする。
    """
    if val is None:
        return ""
    if isinstance(val, (float, np.floating)) and (pd.isna(val) or not math.isfinite(float(val))):
        return ""
    dt = pd.to_datetime(val, errors="coerce")
    if pd.isna(dt):
        return ""
    ts = pd.Timestamp(dt)
    try:
        if ts.tzinfo is None:
            ts = ts.tz_localize(TZ_JP, ambiguous="infer", nonexistent="shift_forward")
        else:
            ts = ts.tz_convert(TZ_JP)
    except (TypeError, ValueError):
        try:
            ts = pd.Timestamp(dt).tz_localize("UTC", ambiguous="infer").tz_convert(TZ_JP)
        except Exception:
            return ""
    return ts.date().isoformat()


def _ledger_stocktake_date_token_for_mid(df: pd.DataFrame, mid: str) -> str:
    row = lookup_ledger_row_by_management_id(df, mid.strip())
    if row is None:
        return ""
    return _stocktake_date_token_for_compare(row.get(COL_LAST_STOCKTAKE))


def _management_ids_origin_cleared_session_in_stock(
    df: pd.DataFrame,
    origin_ids: set[str],
    remaining_ids: set[str],
    *,
    limit: int = 18,
) -> list[str]:
    """今回リスト開始時の対象に含まれ、残リストから外れた **在庫中** の管理ID（キャプション用・先頭 limit 件）。"""
    if df.empty or not origin_ids or COL_MANAGEMENT_ID not in df.columns:
        return []
    rem = {str(x).strip() for x in remaining_ids if str(x).strip()}
    out: list[str] = []
    for mid in sorted({str(m).strip() for m in origin_ids if str(m).strip()}):
        if mid in rem:
            continue
        row = lookup_ledger_row_by_management_id(df, mid)
        if row is None:
            continue
        if _normalize_stock_status(str(row.get(COL_STOCK_STATUS, ""))) != STATUS_IN_STOCK:
            continue
        out.append(mid)
        if len(out) >= limit:
            break
    return out


def _all_in_stock_management_ids(df: pd.DataFrame) -> set[str]:
    """在庫中の行の管理ID（空でないもの）の集合。"""
    if df.empty or COL_MANAGEMENT_ID not in df.columns:
        return set()
    m_in = _mask_ledger_in_stock(df)
    ids = df.loc[m_in, COL_MANAGEMENT_ID].astype(str).str.strip()
    return {x for x in ids.tolist() if x}


def _stocktake_work_session_path() -> Path:
    return Path(__file__).resolve().parent / STOCKTAKE_WORK_SESSION_FILENAME


def _inv_stocktake_work_read_disk() -> tuple[bool, set[str], int | None, set[str], dict[str, str]]:
    """ディスク上のセッション（migrate なし）。(有効, 残りID, baseline, 開始時対象の管理ID, 開始時棚卸日トークン)。"""
    p = _stocktake_work_session_path()
    if not p.is_file():
        return (False, set(), None, set(), {})
    try:
        raw = json.loads(p.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, TypeError):
        return (False, set(), None, set(), {})
    if not isinstance(raw, dict) or not raw.get("active"):
        return (False, set(), None, set(), {})
    rem = raw.get("remaining") or []
    if not isinstance(rem, list):
        rem = []
    s = {str(x).strip() for x in rem if str(x).strip()}
    base = raw.get("session_baseline_n")
    bi = int(base) if isinstance(base, int) and base >= 1 else None
    org_raw = raw.get("session_origin_ids") or []
    if not isinstance(org_raw, list):
        org_raw = []
    orig = {str(x).strip() for x in org_raw if str(x).strip()}
    snap_raw = raw.get("session_stocktake_at_start")
    snap: dict[str, str] = {}
    if isinstance(snap_raw, dict):
        for k, v in snap_raw.items():
            ks = str(k).strip()
            if not ks:
                continue
            snap[ks] = (
                _stocktake_date_token_for_compare(v)
                if str(v).strip()
                else ""
            )
    if not orig:
        if snap:
            orig = set(snap.keys())
        elif s:
            orig = set(s)
    return (True, s, bi, orig, snap)


def _inv_stocktake_work_remaining_read_state(
    df: pd.DataFrame | None = None,
) -> tuple[bool, set[str], int, set[str], dict[str, str]]:
    """表示用のセッション状態。無効時 (False, set(), 0, set(), {})。有効時 baseline は最低 1（旧ファイル互換）。"""
    _inv_stocktake_work_remaining_migrate_legacy_from_session_state()
    a, s, b, o, snap = _inv_stocktake_work_read_disk()
    if not a:
        return False, set(), 0, set(), {}
    eff_b = b if b is not None and b >= 1 else max(len(s), 1)
    if df is not None and not df.empty and o and COL_MANAGEMENT_ID in df.columns:
        snap_m = dict(snap)
        dirty = any(m not in snap_m for m in o)
        if dirty:
            for m in o:
                if m not in snap_m:
                    snap_m[m] = _ledger_stocktake_date_token_for_mid(df, m)
            _inv_stocktake_work_remaining_save(
                True,
                s,
                baseline_override=eff_b,
                session_origin=o,
                stocktake_snapshot=snap_m,
            )
            snap = snap_m
    return True, s, eff_b, o, snap


def _inv_stocktake_work_remaining_save(
    active: bool,
    remaining: set[str],
    *,
    baseline_override: int | None = None,
    session_origin: set[str] | None = None,
    stocktake_snapshot: dict[str, str] | None = None,
) -> None:
    """有効時は JSON に保存（baseline・開始時対象 ID・開始時棚卸日スナップショットを維持または上書き）。無効時はファイル削除。"""
    st.session_state.pop(_SESSION_KEY_STOCKTAKE_WORK_REMAINING_LEGACY, None)
    p = _stocktake_work_session_path()
    if not active:
        try:
            if p.is_file():
                p.unlink()
        except OSError:
            pass
        return
    _a, _old_rem, old_base, old_orig, old_snap = _inv_stocktake_work_read_disk()
    eff_base = baseline_override if baseline_override is not None and baseline_override >= 1 else old_base
    if eff_base is None or eff_base < 1:
        eff_base = max(len(remaining), 1)
    if session_origin is not None:
        eff_origin = {str(x).strip() for x in session_origin if str(x).strip()}
    elif old_orig:
        eff_origin = set(old_orig)
    else:
        eff_origin = {str(x).strip() for x in remaining if str(x).strip()}
    if stocktake_snapshot is not None:
        eff_snap = {
            str(k).strip(): str(v).strip() if v is not None else ""
            for k, v in stocktake_snapshot.items()
            if str(k).strip()
        }
    else:
        eff_snap = dict(old_snap) if old_snap else {}
    data = {
        "active": True,
        "remaining": sorted(remaining),
        "session_baseline_n": int(eff_base),
        "session_origin_ids": sorted(eff_origin),
        "session_stocktake_at_start": {k: eff_snap[k] for k in sorted(eff_snap)},
    }
    tmp = p.with_name(p.name + ".tmp")
    try:
        tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(p)
    except OSError:
        try:
            if tmp.is_file():
                tmp.unlink()
        except OSError:
            pass


def _inv_stocktake_work_remaining_finish_if_empty(remaining: set[str]) -> None:
    """残りが空ならセッション終了（全数確認済み）。それ以外は保存。"""
    if not remaining:
        _inv_stocktake_work_remaining_save(False, set())
    else:
        _inv_stocktake_work_remaining_save(True, remaining)


def _inv_stocktake_work_remaining_migrate_legacy_from_session_state() -> None:
    """旧実装の session_state だけに残っているリストを初回ファイルへ移す。"""
    v = st.session_state.get(_SESSION_KEY_STOCKTAKE_WORK_REMAINING_LEGACY)
    if v is None:
        return
    if isinstance(v, set):
        s = set(v)
    elif isinstance(v, (list, tuple)):
        s = {str(x).strip() for x in v if str(x).strip()}
    else:
        s = set()
    st.session_state.pop(_SESSION_KEY_STOCKTAKE_WORK_REMAINING_LEGACY, None)
    if _stocktake_work_session_path().is_file():
        return
    if not s:
        return
    _inv_stocktake_work_remaining_finish_if_empty(s)


def _inv_stocktake_work_remaining_get() -> set[str] | None:
    """棚卸し作業セッションの残り管理 ID。未開始・全件確認済みで終了したあとは None。"""
    _inv_stocktake_work_remaining_migrate_legacy_from_session_state()
    active, rem, _b, _o, _snap = _inv_stocktake_work_read_disk()
    if not active:
        return None
    return rem


def _inv_stocktake_work_remaining_start(df: pd.DataFrame) -> None:
    """在庫中の全管理IDを「今回の作業」の対象にする。"""
    ids = _all_in_stock_management_ids(df)
    if not ids:
        _inv_stocktake_work_remaining_save(False, set())
        return
    n0 = len(ids)
    snap0 = {m: _ledger_stocktake_date_token_for_mid(df, m) for m in ids}
    _inv_stocktake_work_remaining_save(
        True,
        ids,
        baseline_override=n0,
        session_origin=ids,
        stocktake_snapshot=snap0,
    )


def _inv_stocktake_work_remaining_clear() -> None:
    _inv_stocktake_work_remaining_save(False, set())


def _inv_stocktake_work_remaining_prune(df: pd.DataFrame) -> None:
    """棚卸セッションの残り管理IDを同期する。

    - 販売済・削除などで在庫中でなくなった ID は残リストから外す
    - セッション開始後に新規追加された在庫中 ID は残リストへ自動追加する
    - 残りが空ならセッション終了
    """
    cur = _inv_stocktake_work_remaining_get()
    if cur is None:
        return
    valid = _all_in_stock_management_ids(df)
    _, _, _, old_orig, old_snap = _inv_stocktake_work_read_disk()
    base_orig = set(old_orig) if old_orig else set(cur)
    new_added = valid - base_orig
    newrem = ({m for m in cur if m in valid}) | new_added
    new_orig = (base_orig & valid) | new_added
    new_snap: dict[str, str] = {}
    for m in new_orig:
        if m in old_snap:
            new_snap[m] = old_snap[m]
        else:
            new_snap[m] = _ledger_stocktake_date_token_for_mid(df, m)
    if not newrem:
        _inv_stocktake_work_remaining_save(False, set())
    else:
        _inv_stocktake_work_remaining_save(
            True, newrem, session_origin=new_orig, stocktake_snapshot=new_snap
        )


def _inv_stocktake_work_remaining_note_done(mids: set[str] | str) -> None:
    """棚卸日を更新した管理IDを今回の残リストから外す。"""
    if isinstance(mids, str):
        s = {mids.strip()} if mids.strip() else set()
    else:
        s = {str(x).strip() for x in mids if str(x).strip()}
    if not s:
        return
    cur = _inv_stocktake_work_remaining_get()
    if cur is None:
        return
    _inv_stocktake_work_remaining_finish_if_empty(cur - s)


def _management_ids_last_stocktake_changed(
    df_before: pd.DataFrame, df_after: pd.DataFrame
) -> set[str]:
    """同一管理IDについて「最後に確認した日付（棚卸日）」の表記が変わった管理ID。"""
    if (
        COL_MANAGEMENT_ID not in df_before.columns
        or COL_MANAGEMENT_ID not in df_after.columns
        or COL_LAST_STOCKTAKE not in df_before.columns
        or COL_LAST_STOCKTAKE not in df_after.columns
    ):
        return set()
    bcol = "_stk_prev"
    acol = "_stk_new"
    bef = df_before[[COL_MANAGEMENT_ID, COL_LAST_STOCKTAKE]].copy()
    aft = df_after[[COL_MANAGEMENT_ID, COL_LAST_STOCKTAKE]].copy()
    bef[COL_MANAGEMENT_ID] = bef[COL_MANAGEMENT_ID].astype(str).str.strip()
    aft[COL_MANAGEMENT_ID] = aft[COL_MANAGEMENT_ID].astype(str).str.strip()
    merged = bef.rename(columns={COL_LAST_STOCKTAKE: bcol}).merge(
        aft.rename(columns={COL_LAST_STOCKTAKE: acol}),
        on=COL_MANAGEMENT_ID,
        how="inner",
    )
    out: set[str] = set()
    for _, r in merged.iterrows():
        if str(r[bcol] or "").strip() != str(r[acol] or "").strip():
            mid = str(r[COL_MANAGEMENT_ID]).strip()
            if mid:
                out.add(mid)
    return out


def _ledger_in_stock_rows(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty or COL_STOCK_STATUS not in df.columns:
        return df.iloc[:0].copy()
    return df.loc[_mask_ledger_in_stock(df)].copy()


def _ledger_in_stock_outbound_float_loan_rows(df: pd.DataFrame) -> pd.DataFrame:
    """在庫中かつ出庫種別が出庫（浮貸）の行のみ。"""
    if df.empty:
        return df.iloc[:0].copy()
    return df.loc[_mask_ledger_in_stock_outbound_float_loan(df)].copy()


def _ledger_sold_rows(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty or COL_STOCK_STATUS not in df.columns:
        return df.iloc[:0].copy()
    return df.loc[_mask_ledger_sold(df)].copy()


def _iterate_gemini_inventory_rows(
    df: pd.DataFrame,
    *,
    max_lines: int = 400,
    only_in_stock: bool = True,
    only_sold: bool = False,
    management_ids_filter: set[str] | None = None,
    sale_outbound_type_eq: str | None = None,
) -> Iterator[pd.Series]:
    """`_build_gemini_inventory_context` と同一の順序・終了条件でイテレート（参照画像収集用）。"""
    if df.empty:
        return
    if only_sold:
        sub = _ledger_sold_rows(df)
    elif only_in_stock:
        sub = _ledger_in_stock_rows(df)
    else:
        sub = df
    _eq = (sale_outbound_type_eq or "").strip()
    if _eq:
        if COL_SALE_OUTBOUND_TYPE not in sub.columns:
            return
        sub = sub.loc[
            sub[COL_SALE_OUTBOUND_TYPE].astype(str).str.strip() == _eq
        ].copy()
    if sub.empty:
        return
    # 販売・棚卸し照合（在庫中のみ）では最新登録行を優先して
    # コンテキストに含める。台帳件数が多い場合に新規追加行が max_lines で
    # 切り落とされるのを防ぐ。
    _sort_latest = (only_in_stock and not only_sold) or only_sold
    if management_ids_filter is None and _sort_latest and COL_DATETIME in sub.columns:
        _dt = pd.to_datetime(sub[COL_DATETIME], errors="coerce")
        sub = sub.assign(_dt_for_gemini_sort=_dt).sort_values(
            "_dt_for_gemini_sort", ascending=False, na_position="last"
        )
        sub = sub.drop(columns=["_dt_for_gemini_sort"], errors="ignore")
    eff_max_lines = int(max_lines)
    if management_ids_filter is not None:
        filt = {str(x).strip() for x in management_ids_filter if str(x).strip()}
        if not filt:
            return
        if COL_MANAGEMENT_ID not in sub.columns:
            return
        sub = sub.loc[
            sub[COL_MANAGEMENT_ID].astype(str).str.strip().isin(filt)
        ].copy()
        if sub.empty:
            return
        eff_max_lines = max(eff_max_lines, 800)
    n_lines = 0
    for _, row in sub.iterrows():
        if n_lines >= eff_max_lines:
            break
        mid = str(row.get(COL_MANAGEMENT_ID, "") or "").strip()
        if not mid:
            continue
        n_lines += 1
        yield row


def _inventory_line_text_for_gemini_prompt(row: pd.Series) -> str:
    """Gemini 向け台帳1行。原価・金額・メモ・連絡先等は API に送らない（照合に不要な機密の隔離）。"""
    mid = str(row.get(COL_MANAGEMENT_ID, "") or "").strip()
    pn = str(row.get(COL_NAME, "") or "").strip().replace("\n", " ")
    su = str(row.get(COL_SUPPLIER, "") or "").strip().replace("\n", " ")
    st_lbl = ""
    if COL_STOCK_STATUS in row.index:
        st_lbl = _normalize_stock_status(str(row.get(COL_STOCK_STATUS, "") or ""))
    st_seg = f" 状態={json.dumps(st_lbl, ensure_ascii=False)}" if st_lbl else ""
    cat_seg = ""
    if COL_CATEGORY in row.index:
        cat = str(row.get(COL_CATEGORY, "") or "").strip().replace("\n", " ")
        if cat:
            cat_seg = (
                f" {COL_CATEGORY}={json.dumps(cat, ensure_ascii=False)}"
            )
    return (
        f"- 管理ID={json.dumps(mid, ensure_ascii=False)} "
        f"商品名={json.dumps(pn, ensure_ascii=False)} "
        f"仕入先={json.dumps(su, ensure_ascii=False)}"
        f"{st_seg}{cat_seg}"
    )


def _build_gemini_inventory_context(
    df: pd.DataFrame,
    *,
    max_lines: int = 400,
    only_in_stock: bool = True,
    only_sold: bool = False,
    management_ids_filter: set[str] | None = None,
    sale_outbound_type_eq: str | None = None,
) -> str:
    """台帳行を短い箇条書きにし、画像照合用プロンプトへ埋め込む（**金額・原価・メモ等は送らない**）。

    ``only_in_stock=True`` … 販売照合・棚卸し用（在庫中のみ）。
    ``only_sold=True`` … **販売済** のみ（出庫（返品）の写真照合用）。``only_in_stock`` と同時には使わない。
    ``only_in_stock=False`` … 登録画面の AI 解析用（在庫中・販売済など全行を最大 max_lines 件）。
    ``management_ids_filter`` … 指定時はその管理 ID に含まれる行だけ（棚卸し「今回の残リスト」向け）。
    ``sale_outbound_type_eq`` … 指定時は **出庫種別** がその文字列と一致する行に限定（出庫（戻入）の **出庫（浮貸）** 絞り込み等）。
    **管理IDが空の行はスキップ** し、行数上限を無駄に使わない。
    """
    lines: list[str] = []
    for row in _iterate_gemini_inventory_rows(
        df,
        max_lines=max_lines,
        only_in_stock=only_in_stock and not only_sold,
        only_sold=only_sold,
        management_ids_filter=management_ids_filter,
        sale_outbound_type_eq=sale_outbound_type_eq,
    ):
        lines.append(_inventory_line_text_for_gemini_prompt(row))
    return "\n".join(lines)


def _fuzzy_ledger_match_rows(
    df: pd.DataFrame,
    product_name: str,
    supplier: str,
    *,
    limit: int | None = 8,
) -> pd.DataFrame:
    """台帳候補を返す（優先順: 名前+仕入先 → 名前 → 仕入先）。"""
    if df is None or df.empty:
        return pd.DataFrame()
    pn = (product_name or "").strip().casefold()
    su = (supplier or "").strip().casefold()
    if not pn and not su:
        return df.iloc[:0]

    scored: list[tuple[int, float, Any]] = []
    for i, row in df.iterrows():
        rpn = str(row.get(COL_NAME, "") or "").strip()
        rsu = str(row.get(COL_SUPPLIER, "") or "").strip()
        rpn_l = rpn.casefold()
        rsu_l = rsu.casefold()
        both_ok = bool(pn and su and (pn in rpn_l) and (su in rsu_l))
        name_ok = bool(pn and (pn in rpn_l))
        sup_ok = bool(su and (su in rsu_l))
        if not (both_ok or name_ok or sup_ok):
            continue
        if both_ok:
            prio = 1
        elif name_ok:
            prio = 2
        else:
            prio = 3
        sim = 0.0
        if pn:
            sim += 0.6 * difflib.SequenceMatcher(None, rpn_l, pn).ratio()
        if su:
            sim += 0.4 * difflib.SequenceMatcher(None, rsu_l, su).ratio()
        scored.append((prio, -sim, i))

    scored.sort(key=lambda x: (x[0], x[1], _management_id_sort_key(str(df.loc[x[2]].get(COL_MANAGEMENT_ID, "") or "").strip())))
    if limit is None:
        picked = [i for _, _, i in scored]
    else:
        picked = [i for _, _, i in scored[: max(1, int(limit))]]
    if not picked:
        return df.iloc[:0]
    return df.loc[picked]


def _single_row_fuzzy_ledger_match(
    df: pd.DataFrame | None,
    product_name: str,
    supplier: str,
    *,
    only_in_stock: bool = False,
    only_sold: bool = False,
    only_float_loan_outbound: bool = False,
    limit: int = 12,
) -> pd.Series | None:
    """商品名・仕入先の近い台帳行が **1件に一意に** 定まるときだけその行を返す（AI の match 補完用）。"""
    if df is None or df.empty:
        return None
    if only_float_loan_outbound:
        base = _ledger_in_stock_outbound_float_loan_rows(df)
    elif only_sold:
        base = _ledger_sold_rows(df)
    elif only_in_stock:
        base = _ledger_in_stock_rows(df)
    else:
        base = df
    if base.empty:
        return None
    pn = (product_name or "").strip()
    su = (supplier or "").strip()
    if not pn and not su:
        return None
    cand = _fuzzy_ledger_match_rows(base, pn, su, limit=limit)
    if cand.shape[0] != 1:
        return None
    row = cand.iloc[0]
    mid = str(row.get(COL_MANAGEMENT_ID, "") or "").strip()
    return row if mid else None


def _apply_purchase_ledger_match_supplement(
    result: dict[str, Any],
    df_ledger: pd.DataFrame | None,
) -> dict[str, Any]:
    """Gemini が match.management_id を返さないとき、商品名+特徴を優先して一意照合し match を補う。"""
    if not isinstance(result, dict) or df_ledger is None or df_ledger.empty:
        return result
    m0 = result.get("match")
    m = m0 if isinstance(m0, dict) else {}
    if str(m.get("management_id") or m.get("管理ID") or "").strip():
        return result
    pn = str(
        result.get("product_name")
        or result.get("商品名")
        or m.get("product_name")
        or ""
    ).strip()
    su = str(
        result.get("supplier")
        or result.get("仕入先・取引先")
        or result.get("仕入先")
        or result.get("取引先")
        or m.get("supplier")
        or ""
    ).strip()
    row = _single_row_fuzzy_ledger_match(
        df_ledger, pn, su, only_in_stock=False, limit=12
    )
    if row is None:
        return result
    mid0 = str(row.get(COL_MANAGEMENT_ID, "") or "").strip()
    if not mid0:
        return result
    mm = dict(m) if isinstance(m0, dict) else {}
    mm["management_id"] = mid0
    mm["confidence"] = max(float(mm.get("confidence") or 0), 0.78)
    ly = _finite_int(row.get(COL_PRICE_EXCL), 0)
    if ly > 0:
        mm["line_price_excl"] = ly
    if COL_CATEGORY in row.index:
        ic = str(row.get(COL_CATEGORY, "") or "").strip()
        if ic:
            mm["inventory_category"] = ic
    result["match"] = mm
    return result


def _infer_tax_rate_from_lines_vectorized(
    excl: np.ndarray, incl: np.ndarray
) -> np.ndarray:
    """各行の仕入税抜・税込から税率を推定（:func:`_infer_tax_rate_from_main_line` と同じ優先）。"""
    excl = np.asarray(excl, dtype=np.int64)
    incl = np.asarray(incl, dtype=np.int64)
    default_r = float(CONSUMPTION_TAX_RATE)
    out = np.full(excl.shape[0], default_r, dtype=np.float64)
    excl_le0 = excl <= 0
    out[excl_le0] = default_r
    incl_le = (~excl_le0) & (incl <= excl)
    out[incl_le] = 0.0
    cand = (~excl_le0) & (~incl_le)
    if not np.any(cand):
        return out
    ec = excl[cand].astype(np.float64)
    ic = incl[cand]
    idx = np.flatnonzero(cand)
    rates_found = np.full(len(idx), default_r, dtype=np.float64)
    matched = np.zeros(len(idx), dtype=bool)
    for _label, rate in CONSUMPTION_TAX_CHOICE_TO_RATE.items():
        rr = float(rate)
        pred = np.rint(ec * (1.0 + rr)).astype(np.int64)
        m_new = (~matched) & (pred == ic)
        rates_found[m_new] = rr
        matched |= m_new
    out[idx] = rates_found
    return out


def _compute_gross_profit_row(
    cogs_line_excl: int,
    planned_line_excl: int,
    actual_line_excl: int,
    status: str,
) -> int | None:
    """税抜ベース（行計）。販売済は実売行計−原価、在庫中は販売予定行計−原価。算出不可時は None。"""
    st = _normalize_stock_status(status)
    cg = _finite_int(cogs_line_excl, 0)
    plex = _finite_int(planned_line_excl, 0)
    aex = _finite_int(actual_line_excl, 0)
    if st == STATUS_SOLD:
        if aex > 0:
            return int(aex - cg)
        return None
    if st == STATUS_IN_STOCK:
        if plex > 0:
            return int(plex - cg)
        return None
    return None


def _recalc_gross_profit_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """販売予定/実売の税込総額列と粗利を再計算する（数値列は int で統一）。

    行ループではなく NumPy ベクトル演算で処理し、大行数でも応答を速くする。
    """
    need = (
        COL_GROSS_PROFIT,
        COL_STOCK_STATUS,
        COL_PRICE_EXCL,
        COL_PRICE_INCL,
        COL_PLANNED_SALE,
        COL_ACTUAL_SALE,
        COL_PLANNED_SALE_INCL,
        COL_ACTUAL_SALE_INCL,
    )
    if not all(c in df.columns for c in need):
        return df.copy()
    out = _coerce_money_columns_for_recalc(df)
    n = len(out.index)
    if n == 0:
        return out

    def _i64_col(c: str) -> np.ndarray:
        x = pd.to_numeric(out[c], errors="coerce").fillna(0).to_numpy(
            dtype=np.float64, copy=False
        )
        return np.where(np.isfinite(x), np.rint(x), 0.0).astype(np.int64, copy=False)

    cogs = _i64_col(COL_PRICE_EXCL)
    line_in = _i64_col(COL_PRICE_INCL)
    if COL_QTY in out.columns:
        qv = _i64_col(COL_QTY)
        qty = np.maximum(1, qv).astype(np.int64, copy=False)
    else:
        qty = np.ones(n, dtype=np.int64)
    pu = _i64_col(COL_PLANNED_SALE)
    au = _i64_col(COL_ACTUAL_SALE)

    st_raw = out[COL_STOCK_STATUS].astype(str).str.strip().to_numpy()
    st = np.where(np.isin(st_raw, list(STOCK_STATUS_OPTIONS)), st_raw, STATUS_IN_STOCK)
    out[COL_STOCK_STATUS] = st

    tax_r = _infer_tax_rate_from_lines_vectorized(cogs, line_in)

    plex = np.where(pu > 0, pu * qty, 0).astype(np.int64, copy=False)
    fplex = plex.astype(np.float64, copy=False)
    pincl = np.zeros(n, dtype=np.int64)
    m_plex = plex > 0
    pincl[m_plex] = np.rint(
        fplex[m_plex] * (1.0 + tax_r[m_plex].astype(np.float64))
    ).astype(np.int64)

    is_sold = st == STATUS_SOLD
    aex = np.where(is_sold & (au > 0), au * qty, 0).astype(np.int64, copy=False)
    faex = aex.astype(np.float64, copy=False)
    aincl = np.zeros(n, dtype=np.int64)
    m_aex = aex > 0
    aincl[m_aex] = np.rint(
        faex[m_aex] * (1.0 + tax_r[m_aex].astype(np.float64))
    ).astype(np.int64)

    gp = np.zeros(n, dtype=np.int64)
    m_sold = is_sold & (aex > 0)
    gp[m_sold] = aex[m_sold] - cogs[m_sold]
    is_in = st == STATUS_IN_STOCK
    m_in = is_in & (plex > 0)
    gp[m_in] = plex[m_in] - cogs[m_in]

    out[COL_PLANNED_SALE_INCL] = pincl
    out[COL_ACTUAL_SALE_INCL] = aincl
    out[COL_GROSS_PROFIT] = gp
    return out


def _max_management_serial_from_dataframe(df: pd.DataFrame) -> int:
    """DataFrame の管理ID列から最大シリアルを返す。"""
    mx = 0
    if df is None or df.empty or COL_MANAGEMENT_ID not in df.columns:
        return mx
    for v in df[COL_MANAGEMENT_ID].astype(str):
        s = v.strip()
        if not s:
            continue
        m = re.fullmatch(r"(?i)G(\d+)", s)
        if m:
            mx = max(mx, int(m.group(1)))
            continue
        if s.isdigit():
            mx = max(mx, int(s))
    return mx


def allocate_management_ids(ws: Any, count: int) -> list[str]:
    """管理ID（G########）を count 件、CSV またはシート現状から連番で採番する。"""
    if count <= 0:
        return []
    if _uses_local_inventory_csv():
        df = _inventory_csv_read_df()
        mx = _max_management_serial_from_dataframe(df)
        return [f"G{mx + i + 1:08d}" for i in range(count)]
    if ws is None:
        return []
    try:
        df_exist = load_inventory_dataframe()
    except Exception:
        df_exist = None
    mx = _max_management_serial_from_dataframe(df_exist) if df_exist is not None else 0
    return [f"G{mx + i + 1:08d}" for i in range(count)]


def _inventory_row_values_for_append(
    dt_a: str,
    dt_purchase: str,
    purchase_movement: str,
    product_name: str,
    supplier: str,
    line_price_excl_yen: int,
    line_price_incl_yen: int,
    image_url: str,
    management_id: str,
    memo: str,
    *,
    quantity: int = 1,
    inventory_category: str = "",
    planned_sale_unit_excl_yen: int = 0,
    actual_sale_unit_excl_yen: int = 0,
    stock_status: str = STATUS_IN_STOCK,
    consumption_tax_rate: float | None = None,
    loan_datetime: str = "",
    voucher_recorded_at: str = "",
    voucher_evidence_url: str = "",
) -> list[Any]:
    """台帳 EXPECTED_HEADERS 順の1行分セル値を組み立てる（追記用）。"""
    cogs = _finite_int(line_price_excl_yen, 0)
    qty_i = max(1, _finite_int(quantity, 1))
    pl_u = _finite_int(planned_sale_unit_excl_yen, 0)
    ac_u = _finite_int(actual_sale_unit_excl_yen, 0)
    stt = _normalize_stock_status(str(stock_status))
    tax_r = (
        float(consumption_tax_rate)
        if consumption_tax_rate is not None
        and math.isfinite(float(consumption_tax_rate))
        else _infer_tax_rate_from_main_line(
            _finite_int(line_price_excl_yen, 0),
            _finite_int(line_price_incl_yen, 0),
        )
    )
    plex, pincl, aex, aincl = _planned_actual_line_amounts(
        qty_i, pl_u, ac_u, stt, tax_r
    )
    planned_unit_cell = _optional_amount_cell(pl_u)
    planned_incl_cell = _optional_amount_cell(pincl)
    actual_unit_cell = (
        _optional_amount_cell(ac_u) if stt == STATUS_SOLD else 0
    )
    actual_incl_cell = (
        _optional_amount_cell(aincl) if stt == STATUS_SOLD else 0
    )
    gp = _compute_gross_profit_row(
        cogs,
        plex,
        aex if stt == STATUS_SOLD else 0,
        stt,
    )
    gross_cell = 0 if gp is None else _finite_int(gp, 0)
    pm = (purchase_movement or "").strip()
    row_map: dict[str, Any] = {
        COL_DATETIME: dt_a,
        COL_NAME: product_name,
        COL_SUPPLIER: supplier,
        COL_QTY: qty_i,
        COL_PRICE_EXCL: line_price_excl_yen,
        COL_PRICE_INCL: line_price_incl_yen,
        COL_PLANNED_SALE: planned_unit_cell,
        COL_PLANNED_SALE_INCL: planned_incl_cell,
        COL_ACTUAL_SALE: actual_unit_cell,
        COL_ACTUAL_SALE_INCL: actual_incl_cell,
        COL_GROSS_PROFIT: gross_cell,
        COL_STOCK_STATUS: stt,
        COL_MEMO: memo,
        COL_CATEGORY: (inventory_category or "").strip(),
        COL_IMAGE_URL: image_url,
        COL_SALE_IMAGE_URL: "",
        COL_MANAGEMENT_ID: management_id,
        COL_LAST_STOCKTAKE: "",
        COL_VOUCHER_RECORDED_AT: (voucher_recorded_at or "").strip(),
        COL_VOUCHER_EVIDENCE_URL: (voucher_evidence_url or "").strip(),
        COL_PURCHASE_DATETIME: dt_purchase,
        COL_PURCHASE_MOVEMENT: pm,
        COL_LOAN_DATETIME: (loan_datetime or "").strip(),
        COL_SALE_DATETIME: (dt_a if stt == STATUS_SOLD else ""),
        COL_SALE_OUTBOUND_TYPE: (
            pm
            if stt == STATUS_SOLD and _movement_is_outbound(pm)
            else ("出庫（販売）" if stt == STATUS_SOLD else "")
        ),
    }
    return [row_map.get(c, "") for c in EXPECTED_HEADERS]


def _append_inventory_data_rows(rows: list[list[Any]]) -> None:
    """台帳に複数行をまとめて追記する（CSV は1回の再計算・書き込み、シートは append_rows 1回）。

    連続で数十行追記するとき、行ごとの append + 書式 API を繰り返すと失敗や取りこぼしの原因になるため一括にする。
    """
    if not rows:
        return
    if _uses_local_inventory_csv():
        df = _inventory_csv_read_df()
        for rv in rows:
            new_row = dict(zip(EXPECTED_HEADERS, rv))
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        df = _recalc_gross_profit_dataframe(df)
        _inventory_csv_write_df(df.reindex(columns=EXPECTED_HEADERS))
        return
    ws = ensure_worksheet_header()
    if ws is None:
        raise RuntimeError(
            "スプレッドシートに接続できません。"
            f"{SECRET_GOOGLE_SPREADSHEET_ID} とサービスアカウント権限を確認してください。"
        )
    try:
        try:
            ws.append_rows(rows, value_input_option="USER_ENTERED")
        except Exception:
            try:
                ws.resize(rows=int(ws.row_count) + len(rows) + 200)
            except Exception:
                pass
            ws.append_rows(rows, value_input_option="USER_ENTERED")
        try:
            _apply_inventory_amount_number_formats(ws)
        except Exception:
            pass
    except Exception as e:
        raise RuntimeError(f"スプレッドシート追記に失敗しました: {e}") from e
    _bump_inventory_sheet_cache_bust()


def append_sheet_row(
    purchase_movement: str,
    product_name: str,
    supplier: str,
    line_price_excl_yen: int,
    line_price_incl_yen: int,
    image_url: str,
    management_id: str,
    memo: str = "",
    record_datetime: str | None = None,
    *,
    quantity: int = 1,
    inventory_category: str = "",
    planned_sale_unit_excl_yen: int = 0,
    actual_sale_unit_excl_yen: int = 0,
    stock_status: str = STATUS_IN_STOCK,
    consumption_tax_rate: float | None = None,
    loan_datetime: str = "",
    voucher_recorded_at: str = "",
    voucher_evidence_url: str = "",
):
    """1点1行で台帳に追記する（仕入単価列は持たない）。

    A列「日時」は **追記実行の JST**。仕入の暦は「仕入日時」に ``record_datetime``（EXIF 等）を渡す。
    """
    dt_a = jst_now_str()
    dt_purchase = ((record_datetime or "").strip() or dt_a)
    row_vals = _inventory_row_values_for_append(
        dt_a,
        dt_purchase,
        purchase_movement,
        product_name,
        supplier,
        line_price_excl_yen,
        line_price_incl_yen,
        image_url,
        management_id,
        memo,
        quantity=quantity,
        inventory_category=inventory_category,
        planned_sale_unit_excl_yen=planned_sale_unit_excl_yen,
        actual_sale_unit_excl_yen=actual_sale_unit_excl_yen,
        stock_status=stock_status,
        consumption_tax_rate=consumption_tax_rate,
        loan_datetime=loan_datetime,
        voucher_recorded_at=voucher_recorded_at,
        voucher_evidence_url=voucher_evidence_url,
    )
    _append_inventory_data_rows([row_vals])


def _sheet_header_row_to_expected_list(header: list[str], row: list[Any]) -> list[str]:
    """ヘッダー名でセルを対応付け、EXPECTED_HEADERS 順の1行にする。旧「入出庫種別」は入庫種別へ移す。"""
    h = [("" if x is None else str(x)).strip() for x in header]
    rlist = [("" if x is None else str(x)) for x in list(row)]
    h2: list[str] = []
    r2: list[str] = []
    for i, nm in enumerate(h):
        if nm == LEGACY_COL_UNIT_PRICE:
            continue
        if nm == LEGACY_COL_SALE_SOURCE_MGMT_ID:
            continue
        h2.append(nm)
        r2.append(rlist[i] if i < len(rlist) else "")
    dd = dict(zip(h2, r2))
    out: dict[str, str] = {c: str(dd.get(c, "") or "") for c in EXPECTED_HEADERS}
    if LEGACY_COL_MOVEMENT_TYPE in dd and not (out.get(COL_PURCHASE_MOVEMENT) or "").strip():
        out[COL_PURCHASE_MOVEMENT] = str(dd.get(LEGACY_COL_MOVEMENT_TYPE, "") or "")
    return [out[c] for c in EXPECTED_HEADERS]


def _repair_rows_shifted_by_sale_image_column(df: pd.DataFrame) -> pd.DataFrame:
    """販売画像URL列追加前の旧列順で追記された行を補正する。"""
    if df is None or df.empty:
        return df
    req_cols = (
        COL_SALE_IMAGE_URL,
        COL_MANAGEMENT_ID,
        COL_LAST_STOCKTAKE,
        COL_VOUCHER_RECORDED_AT,
        COL_VOUCHER_EVIDENCE_URL,
        COL_PURCHASE_DATETIME,
        COL_PURCHASE_MOVEMENT,
        COL_LOAN_DATETIME,
        COL_SALE_DATETIME,
        COL_SALE_OUTBOUND_TYPE,
    )
    if any(c not in df.columns for c in req_cols):
        return df
    out = df.copy()
    sale_img = out[COL_SALE_IMAGE_URL].astype(str).str.strip()
    mid = out[COL_MANAGEMENT_ID].astype(str).str.strip()
    mask = mid.eq("") & sale_img.str.fullmatch(r"G\d{8,}", na=False)
    if not mask.any():
        return out
    for idx in out.index[mask]:
        v16 = str(out.at[idx, COL_SALE_IMAGE_URL] or "").strip()
        v17 = str(out.at[idx, COL_MANAGEMENT_ID] or "").strip()
        v18 = str(out.at[idx, COL_LAST_STOCKTAKE] or "").strip()
        v19 = str(out.at[idx, COL_VOUCHER_RECORDED_AT] or "").strip()
        v20 = str(out.at[idx, COL_VOUCHER_EVIDENCE_URL] or "").strip()
        v21 = str(out.at[idx, COL_PURCHASE_DATETIME] or "").strip()
        v22 = str(out.at[idx, COL_PURCHASE_MOVEMENT] or "").strip()
        v23 = str(out.at[idx, COL_LOAN_DATETIME] or "").strip()
        v24 = str(out.at[idx, COL_SALE_DATETIME] or "").strip()
        out.at[idx, COL_SALE_IMAGE_URL] = ""
        out.at[idx, COL_MANAGEMENT_ID] = v16
        out.at[idx, COL_LAST_STOCKTAKE] = v17
        out.at[idx, COL_VOUCHER_RECORDED_AT] = v18
        out.at[idx, COL_VOUCHER_EVIDENCE_URL] = v19
        out.at[idx, COL_PURCHASE_DATETIME] = v20
        out.at[idx, COL_PURCHASE_MOVEMENT] = v21
        out.at[idx, COL_LOAN_DATETIME] = v22
        out.at[idx, COL_SALE_DATETIME] = v23
        out.at[idx, COL_SALE_OUTBOUND_TYPE] = v24
    return out


def load_inventory_dataframe() -> pd.DataFrame | None:
    """1行目をヘッダー、2行目以降をデータとして読み込み、列は EXPECTED_HEADERS に揃える。"""
    if _uses_local_inventory_csv():
        return _repair_rows_shifted_by_sale_image_column(_inventory_csv_read_df())
    sid = _secret_str(SECRET_GOOGLE_SPREADSHEET_ID)
    if not sid:
        return None
    wname = _secret_str(SECRET_GOOGLE_WORKSHEET_NAME, DEFAULT_WORKSHEET_NAME)
    if SESSION_KEY_INV_SHEET_CACHE_BUST not in st.session_state:
        # 新しいブラウザセッションが古い cache_data(=bust 0) を拾わないよう初期値を現在時刻にする
        st.session_state[SESSION_KEY_INV_SHEET_CACHE_BUST] = int(
            datetime.now().timestamp()
        )
    bust = int(st.session_state.get(SESSION_KEY_INV_SHEET_CACHE_BUST, 0))
    try:
        raw = _inventory_sheet_get_all_values_cached(str(sid), str(wname), bust)
    except Exception as e:
        st.session_state["_inventory_sheet_load_error"] = str(e)
        return None
    if not raw:
        return pd.DataFrame(columns=EXPECTED_HEADERS)
    header0 = [("" if c is None else str(c)).strip() for c in raw[0]]
    rows = raw[1:]
    data_rows = [_sheet_header_row_to_expected_list(header0, list(r)) for r in rows]
    return _repair_rows_shifted_by_sale_image_column(
        pd.DataFrame(data_rows, columns=EXPECTED_HEADERS)
    )


def _ledger_hint_dataframe() -> pd.DataFrame | None:
    """登録画面の台帳照合用に在庫を読む（失敗時は None）。"""
    if not _uses_local_inventory_csv() and not _secret_str(SECRET_GOOGLE_SPREADSHEET_ID):
        return None
    try:
        return load_inventory_dataframe()
    except Exception as e:
        st.session_state["_ledger_hint_load_error"] = str(e)
        return None


def _refresh_ledger_quick_search_candidates(df_ledger: pd.DataFrame | None) -> None:
    """写真解析・手入力の商品名／仕入先から、在庫中の近い行を session_state に格納する。"""
    if df_ledger is None or df_ledger.empty:
        st.session_state.pop("ledger_quick_candidates", None)
        return
    pn = str(st.session_state.get("field_product_name", "") or "").strip()
    su = str(st.session_state.get("field_supplier", "") or "").strip()
    cand = _fuzzy_ledger_match_rows(df_ledger, pn, su, limit=None)
    if cand.empty:
        st.session_state.pop("ledger_quick_candidates", None)
    else:
        st.session_state["ledger_quick_candidates"] = cand


def _ledger_unique_col_values(df: pd.DataFrame, col: str, *, max_n: int = 800) -> list[str]:
    """台帳 DataFrame から列のユニーク値（空除く）を昇順で返す。"""
    if df is None or df.empty or col not in df.columns:
        return []
    s = df[col].astype(str).str.strip()
    s = s[s != ""]
    return sorted(set(s.tolist()), key=lambda x: (x.casefold(), x))[:max_n]


def _assist_field_keys(prefix: str) -> dict[str, str]:
    """台帳入力補助の session / widget キー（prefix 無しなら仕入タブと同じ名前）。"""
    if prefix:
        return {
            "fp_filter": f"{prefix}hint_filter_product_name",
            "fs_filter": f"{prefix}hint_filter_supplier",
            "fc_filter": f"{prefix}hint_filter_inventory_category",
            "fm_filter": f"{prefix}hint_filter_management_id",
            "pick_p": f"{prefix}ledger_pick_product_name",
            "pick_s": f"{prefix}ledger_pick_supplier",
            "pick_c": f"{prefix}ledger_pick_inventory_category",
            "pick_m": f"{prefix}ledger_pick_management_id",
            "seen_fp": f"{prefix}_hint_seen_product",
            "seen_fs": f"{prefix}_hint_seen_supplier",
            "seen_fc": f"{prefix}_hint_seen_category",
            "seen_fm": f"{prefix}_hint_seen_management_id",
        }
    return {
        "fp_filter": "hint_filter_product_name",
        "fs_filter": "hint_filter_supplier",
        "fc_filter": "hint_filter_inventory_category",
        "fm_filter": "hint_filter_management_id",
        "pick_p": "ledger_pick_product_name",
        "pick_s": "ledger_pick_supplier",
        "pick_c": "ledger_pick_inventory_category",
        "pick_m": "ledger_pick_management_id",
        "seen_fp": "_hint_fp_seen",
        "seen_fs": "_hint_fs_seen",
        "seen_fc": "_hint_cat_seen",
        "seen_fm": "_hint_mid_seen",
    }


def _render_ledger_pick_assist_three_columns(
    df: pd.DataFrame,
    *,
    key_prefix: str,
    body_caption: str,
    on_pick_product_name: Any,
    on_pick_supplier: Any,
    on_pick_inventory_category: Any,
    on_pick_management_id: Any,
    empty_message: str = "台帳が空か読み込めないため、入力補助の候補は表示できません。",
    sales_restrict_to_sold: bool = False,
    sales_restrict_to_float_loan_outbound: bool = False,
) -> None:
    """商品名／仕入先／カテゴリー／管理IDの絞り込みと台帳プルダウン（仕入・販売・棚卸で共用）。"""
    if df is None or df.empty:
        st.caption(empty_message)
        return
    k = _assist_field_keys(key_prefix)
    st.caption(body_caption)
    df_mid_scope = df
    if key_prefix == "sales_" and COL_MANAGEMENT_ID in df.columns:
        if sales_restrict_to_sold:
            df_mid_scope = df.loc[_mask_ledger_sold(df)]
        elif sales_restrict_to_float_loan_outbound:
            df_mid_scope = df.loc[_mask_ledger_in_stock_outbound_float_loan(df)]
        else:
            df_mid_scope = df.loc[_mask_ledger_in_stock(df)]
    df_opts = df_mid_scope if key_prefix == "sales_" else df
    hc1, hc2 = st.columns(2)
    with hc1:
        st.text_input(
            "商品名の絞り込み（部分一致）",
            key=k["fp_filter"],
            placeholder="例: 帯",
        )
        fp = str(st.session_state.get(k["fp_filter"], "") or "").strip()
        if st.session_state.get(k["seen_fp"], "") != fp:
            st.session_state[k["seen_fp"]] = fp
            st.session_state[k["pick_p"]] = LEDGER_PICK_PLACEHOLDER
        opts_p = _ledger_unique_col_values(df_opts, COL_NAME)
        if fp.casefold():
            q = fp.casefold()
            opts_p = [x for x in opts_p if q in x.casefold()][:400]
        st.selectbox(
            "台帳に登録済みの商品名から選ぶ",
            options=[LEDGER_PICK_PLACEHOLDER] + opts_p,
            key=k["pick_p"],
            on_change=on_pick_product_name,
        )
    with hc2:
        st.text_input(
            "仕入先・取引先の絞り込み（部分一致）",
            key=k["fs_filter"],
            placeholder="例: ⚫︎⚫︎会社",
        )
        fs = str(st.session_state.get(k["fs_filter"], "") or "").strip()
        if st.session_state.get(k["seen_fs"], "") != fs:
            st.session_state[k["seen_fs"]] = fs
            st.session_state[k["pick_s"]] = LEDGER_PICK_PLACEHOLDER
        opts_s = _ledger_unique_col_values(df_opts, COL_SUPPLIER)
        if fs.casefold():
            q = fs.casefold()
            opts_s = [x for x in opts_s if q in x.casefold()][:400]
        st.selectbox(
            "台帳に登録済みの仕入先・取引先から選ぶ",
            options=[LEDGER_PICK_PLACEHOLDER] + opts_s,
            key=k["pick_s"],
            on_change=on_pick_supplier,
        )

    hc3, hc4 = st.columns(2)
    with hc3:
        if COL_CATEGORY in df.columns:
            st.text_input(
                "在庫カテゴリーの絞り込み（部分一致）",
                key=k["fc_filter"],
                placeholder="例: 帯",
            )
            fc = str(st.session_state.get(k["fc_filter"], "") or "").strip()
            if st.session_state.get(k["seen_fc"], "") != fc:
                st.session_state[k["seen_fc"]] = fc
                st.session_state[k["pick_c"]] = LEDGER_PICK_PLACEHOLDER
            opts_c = _ledger_unique_col_values(df_opts, COL_CATEGORY)
            if fc.casefold():
                q = fc.casefold()
                opts_c = [x for x in opts_c if q in x.casefold()][:400]
            st.selectbox(
                "台帳の在庫カテゴリーから選ぶ",
                options=[LEDGER_PICK_PLACEHOLDER] + opts_c,
                key=k["pick_c"],
                on_change=on_pick_inventory_category,
            )
        else:
            st.caption("台帳に在庫カテゴリー列がありません。")
    with hc4:
        if COL_MANAGEMENT_ID not in df.columns:
            st.caption("台帳に管理ID列がありません。")
        elif df_mid_scope is None or df_mid_scope.empty:
            if key_prefix == "sales_" and sales_restrict_to_sold:
                _mid_empty_msg = "販売済で **管理ID** のある行がありません。"
            elif key_prefix == "sales_" and sales_restrict_to_float_loan_outbound:
                _mid_empty_msg = (
                    "在庫中かつ **出庫種別** が **出庫（浮貸）** で **管理ID** のある行がありません。"
                )
            elif key_prefix == "sales_":
                _mid_empty_msg = (
                    "在庫中で **管理ID** のある行がありません（すべて販売済みの可能性があります）。"
                )
            else:
                _mid_empty_msg = "管理IDが付いた対象行がありません。"
            st.caption(_mid_empty_msg)
        else:
            st.text_input(
                "管理IDの絞り込み（部分一致）",
                key=k["fm_filter"],
                placeholder="例: G00042",
            )
            fm = str(st.session_state.get(k["fm_filter"], "") or "").strip()
            if st.session_state.get(k["seen_fm"], "") != fm:
                st.session_state[k["seen_fm"]] = fm
                st.session_state[k["pick_m"]] = LEDGER_PICK_PLACEHOLDER
            opts_ids = sorted(
                {
                    str(x).strip()
                    for x in df_mid_scope[COL_MANAGEMENT_ID].tolist()
                    if str(x).strip()
                },
                key=_management_id_sort_key,
            )[:600]
            if fm.casefold():
                qm = fm.casefold()
                opts_ids = [
                    x for x in opts_ids if qm in x.casefold()
                ][:400]
            st.selectbox(
                (
                    (
                        "入力補助：販売済の管理IDから選ぶ"
                        if sales_restrict_to_sold
                        else (
                            "入力補助：出庫（浮貸）の在庫の管理IDから選ぶ"
                            if sales_restrict_to_float_loan_outbound
                            else "入力補助：在庫中の管理IDから選ぶ"
                        )
                    )
                    if key_prefix == "sales_"
                    else "入力補助：対象リストの管理IDから選ぶ"
                    if key_prefix == "stocktake_"
                    else "入力補助：台帳の管理IDから選ぶ"
                ),
                options=[LEDGER_PICK_PLACEHOLDER] + opts_ids,
                key=k["pick_m"],
                on_change=on_pick_management_id,
            )


def _stocktake_assist_scope_dataframe(
    df: pd.DataFrame | None, remaining_management_ids: set[str] | None
) -> pd.DataFrame | None:
    """棚卸入力補助の対象行（在庫中。remaining付きならその管理IDのみ）。"""
    if df is None or df.empty:
        return None
    sub = df.loc[_mask_ledger_in_stock(df)]
    if COL_MANAGEMENT_ID not in sub.columns:
        return sub
    if remaining_management_ids is not None:
        filt = {str(x).strip() for x in remaining_management_ids if str(x).strip()}
        if not filt:
            return sub.iloc[0:0].copy()
        m = sub[COL_MANAGEMENT_ID].astype(str).str.strip().isin(filt)
        sub = sub.loc[m].copy()
    return sub


def _sales_rows_matching_assist_buffers() -> tuple[pd.DataFrame, list[str]]:
    """販売用: 入力補助バッファに一致する台帳行と管理 ID 一覧。

    **出庫（返品）** のときは **販売済** の行のみ。**出庫（戻入）** のときは **在庫中かつ出庫種別が出庫（浮貸）** の行のみ。それ以外は **在庫中**。
    """
    try:
        df = load_inventory_dataframe()
    except Exception:
        return pd.DataFrame(), []
    if df is None or df.empty:
        return pd.DataFrame(), []
    _sold_scope = (
        str(st.session_state.get("sales_tab_outbound_kind", "") or "").strip()
        == "出庫（返品）"
    )
    _receipt_scope = (
        str(st.session_state.get("sales_tab_outbound_kind", "") or "").strip()
        == "出庫（戻入）"
    )
    if _sold_scope:
        sub = df.loc[_mask_ledger_sold(df)]
    elif _receipt_scope:
        sub = df.loc[_mask_ledger_in_stock_outbound_float_loan(df)]
    else:
        sub = df.loc[_mask_ledger_in_stock(df)]
    pn = str(st.session_state.get("sales_assist_buf_product_name", "") or "").strip()
    su = str(st.session_state.get("sales_assist_buf_supplier", "") or "").strip()
    cat = str(st.session_state.get("sales_assist_buf_inventory_category", "") or "").strip()
    midb = str(st.session_state.get("sales_assist_buf_management_id", "") or "").strip()
    if midb and COL_MANAGEMENT_ID in sub.columns:
        sub = sub.loc[sub[COL_MANAGEMENT_ID].astype(str).str.strip() == midb]
    if pn and COL_NAME in sub.columns:
        sub = sub.loc[sub[COL_NAME].astype(str).str.strip() == pn]
    if su and COL_SUPPLIER in sub.columns:
        sub = sub.loc[sub[COL_SUPPLIER].astype(str).str.strip() == su]
    if cat and COL_CATEGORY in sub.columns:
        sub = sub.loc[sub[COL_CATEGORY].astype(str).str.strip() == cat]
    if sub.empty or COL_MANAGEMENT_ID not in sub.columns:
        return sub, []
    mids = sorted(
        {str(x).strip() for x in sub[COL_MANAGEMENT_ID].tolist() if str(x).strip()},
        key=_management_id_sort_key,
    )
    return sub, mids


def _sales_apply_mgmt_id_after_assist_pick() -> None:
    _, mids = _sales_rows_matching_assist_buffers()
    st.session_state.sales_assist_last_n_matching_mids = len(mids)
    if len(mids) == 1:
        st.session_state.field_sale_source_mgmt_id = mids[0]


def _on_sales_assist_pick_product_name() -> None:
    v = str(st.session_state.get("sales_ledger_pick_product_name", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.sales_assist_buf_product_name = v
        st.session_state.sales_ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
    _sales_apply_mgmt_id_after_assist_pick()


def _on_sales_assist_pick_supplier() -> None:
    v = str(st.session_state.get("sales_ledger_pick_supplier", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.sales_assist_buf_supplier = v
        st.session_state.sales_ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
    _sales_apply_mgmt_id_after_assist_pick()


def _on_sales_assist_pick_inventory_category() -> None:
    v = str(st.session_state.get("sales_ledger_pick_inventory_category", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.sales_assist_buf_inventory_category = v
        st.session_state.sales_ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER
    _sales_apply_mgmt_id_after_assist_pick()


def _on_sales_assist_pick_management_id() -> None:
    v = str(st.session_state.get("sales_ledger_pick_management_id", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.sales_assist_buf_management_id = v
        st.session_state.sales_ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
    _sales_apply_mgmt_id_after_assist_pick()


def _refresh_sales_assist_quick_candidates(df_hint: pd.DataFrame | None) -> None:
    """販売タブ・入力補助バッファから近い行を一覧用に格納する（返品＝販売済、戻入＝出庫浮貸の在庫中、他＝在庫中）。"""
    if df_hint is None or df_hint.empty:
        st.session_state.pop("sales_assist_quick_candidates", None)
        return
    pn = str(st.session_state.get("sales_assist_buf_product_name", "") or "").strip()
    su = str(st.session_state.get("sales_assist_buf_supplier", "") or "").strip()
    _sold_scope = (
        str(st.session_state.get("sales_tab_outbound_kind", "") or "").strip()
        == "出庫（返品）"
    )
    _receipt_scope = (
        str(st.session_state.get("sales_tab_outbound_kind", "") or "").strip()
        == "出庫（戻入）"
    )
    if _sold_scope:
        sub = df_hint.loc[_mask_ledger_sold(df_hint)]
    elif _receipt_scope:
        sub = df_hint.loc[_mask_ledger_in_stock_outbound_float_loan(df_hint)]
    else:
        sub = df_hint.loc[_mask_ledger_in_stock(df_hint)]
    if sub.empty:
        st.session_state.pop("sales_assist_quick_candidates", None)
        return
    cand = _fuzzy_ledger_match_rows(sub, pn, su, limit=None)
    if cand.empty:
        st.session_state.pop("sales_assist_quick_candidates", None)
    else:
        st.session_state["sales_assist_quick_candidates"] = cand


def _stocktake_rows_matching_assist_buffers(
    df_hint: pd.DataFrame | None, remaining: set[str] | None
) -> tuple[pd.DataFrame, list[str]]:
    base = _stocktake_assist_scope_dataframe(df_hint, remaining)
    if base is None or base.empty:
        return base or pd.DataFrame(), []
    sub = base
    pn = str(st.session_state.get("stocktake_assist_buf_product_name", "") or "").strip()
    su = str(st.session_state.get("stocktake_assist_buf_supplier", "") or "").strip()
    cat = str(st.session_state.get("stocktake_assist_buf_inventory_category", "") or "").strip()
    midb = str(st.session_state.get("stocktake_assist_buf_management_id", "") or "").strip()
    if midb and COL_MANAGEMENT_ID in sub.columns:
        sub = sub.loc[sub[COL_MANAGEMENT_ID].astype(str).str.strip() == midb]
    if pn and COL_NAME in sub.columns:
        sub = sub.loc[sub[COL_NAME].astype(str).str.strip() == pn]
    if su and COL_SUPPLIER in sub.columns:
        sub = sub.loc[sub[COL_SUPPLIER].astype(str).str.strip() == su]
    if cat and COL_CATEGORY in sub.columns:
        sub = sub.loc[sub[COL_CATEGORY].astype(str).str.strip() == cat]
    if sub.empty or COL_MANAGEMENT_ID not in sub.columns:
        return sub, []
    mids = sorted(
        {str(x).strip() for x in sub[COL_MANAGEMENT_ID].tolist() if str(x).strip()},
        key=_management_id_sort_key,
    )
    return sub, mids


def _stocktake_apply_selected_mid_after_assist_pick() -> None:
    st_rem = _inv_stocktake_work_remaining_get()
    df = _ledger_hint_dataframe()
    _, mids = _stocktake_rows_matching_assist_buffers(df, st_rem)
    st.session_state.stocktake_assist_last_n_matching_mids = len(mids)
    if len(mids) == 1:
        st.session_state["_stocktake_selected_mid"] = mids[0]


def _on_stocktake_assist_pick_product_name() -> None:
    v = str(st.session_state.get("stocktake_ledger_pick_product_name", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.stocktake_assist_buf_product_name = v
        st.session_state.stocktake_ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
    _stocktake_apply_selected_mid_after_assist_pick()


def _on_stocktake_assist_pick_supplier() -> None:
    v = str(st.session_state.get("stocktake_ledger_pick_supplier", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.stocktake_assist_buf_supplier = v
        st.session_state.stocktake_ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
    _stocktake_apply_selected_mid_after_assist_pick()


def _on_stocktake_assist_pick_inventory_category() -> None:
    v = str(st.session_state.get("stocktake_ledger_pick_inventory_category", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.stocktake_assist_buf_inventory_category = v
        st.session_state.stocktake_ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER
    _stocktake_apply_selected_mid_after_assist_pick()


def _on_stocktake_assist_pick_management_id() -> None:
    v = str(st.session_state.get("stocktake_ledger_pick_management_id", "") or "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.stocktake_assist_buf_management_id = v
        st.session_state.stocktake_ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
    _stocktake_apply_selected_mid_after_assist_pick()


def _refresh_stocktake_assist_quick_candidates(
    df_hint: pd.DataFrame | None, remaining: set[str] | None
) -> None:
    if df_hint is None or df_hint.empty:
        st.session_state.pop("stocktake_assist_quick_candidates", None)
        return
    base = _stocktake_assist_scope_dataframe(df_hint, remaining)
    if base is None or base.empty:
        st.session_state.pop("stocktake_assist_quick_candidates", None)
        return
    pn = str(st.session_state.get("stocktake_assist_buf_product_name", "") or "").strip()
    su = str(st.session_state.get("stocktake_assist_buf_supplier", "") or "").strip()
    cand = _fuzzy_ledger_match_rows(base, pn, su, limit=None)
    if cand.empty:
        st.session_state.pop("stocktake_assist_quick_candidates", None)
    else:
        st.session_state["stocktake_assist_quick_candidates"] = cand


def _ledger_in_stock_management_ids(df: pd.DataFrame, *, max_n: int = 600) -> list[str]:
    """在庫中の行の管理ID一覧（販売元のプルダウン用）。"""
    if df is None or df.empty or COL_MANAGEMENT_ID not in df.columns:
        return []
    sub = df.loc[_mask_ledger_in_stock(df)]
    s = sub[COL_MANAGEMENT_ID].astype(str).str.strip()
    s = s[s != ""]
    return sorted(set(s.tolist()), key=lambda x: (x.casefold(), x))[:max_n]


def _ledger_sold_management_ids(df: pd.DataFrame, *, max_n: int = 600) -> list[str]:
    """販売済の行の管理ID一覧（出庫（返品）のプルダウン用）。"""
    if df is None or df.empty or COL_MANAGEMENT_ID not in df.columns:
        return []
    sub = df.loc[_mask_ledger_sold(df)]
    s = sub[COL_MANAGEMENT_ID].astype(str).str.strip()
    s = s[s != ""]
    return sorted(set(s.tolist()), key=lambda x: (x.casefold(), x))[:max_n]


def _ledger_in_stock_outbound_float_loan_management_ids(
    df: pd.DataFrame, *, max_n: int = 600
) -> list[str]:
    """在庫中かつ出庫種別が出庫（浮貸）の管理ID一覧（出庫（戻入）のプルダウン用）。"""
    if df is None or df.empty or COL_MANAGEMENT_ID not in df.columns:
        return []
    sub = df.loc[_mask_ledger_in_stock_outbound_float_loan(df)]
    s = sub[COL_MANAGEMENT_ID].astype(str).str.strip()
    s = s[s != ""]
    return sorted(set(s.tolist()), key=lambda x: (x.casefold(), x))[:max_n]


def _sync_planned_sale_from_ledger_in_stock_match() -> None:
    """在庫中の台帳行で、入力中の商品名（＋仕入先が入っていれば一致）に合わせ販売予定（税抜・1点）を反映する。"""
    try:
        df = load_inventory_dataframe()
    except Exception:
        return
    if df is None or df.empty or COL_PLANNED_SALE not in df.columns:
        return
    pn = str(st.session_state.get("field_product_name", "") or "").strip()
    if not pn:
        return
    sub = df.loc[_mask_ledger_in_stock(df)]
    if sub.empty:
        return
    m = sub[COL_NAME].astype(str).str.strip() == pn
    su = str(st.session_state.get("field_supplier", "") or "").strip()
    if su and COL_SUPPLIER in sub.columns:
        m = m & (sub[COL_SUPPLIER].astype(str).str.strip() == su)
    hit = sub.loc[m]
    if hit.empty:
        return
    if COL_MANAGEMENT_ID in hit.columns:
        hit = hit.sort_values(
            COL_MANAGEMENT_ID,
            key=lambda s: s.astype(str).str.strip(),
            na_position="last",
        )
    pl = _finite_int(hit.iloc[0].get(COL_PLANNED_SALE), 0)
    st.session_state.field_planned_sale_excl = max(0, int(pl))
    if COL_CATEGORY in hit.columns:
        cat0 = str(hit.iloc[0].get(COL_CATEGORY, "") or "").strip()
        if cat0:
            st.session_state.field_inventory_category = cat0


def _on_ledger_pick_product_name() -> None:
    v = st.session_state.get("ledger_pick_product_name", "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.field_product_name = v
        st.session_state.ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
        _sync_planned_sale_from_ledger_in_stock_match()


def _on_ledger_pick_supplier() -> None:
    v = st.session_state.get("ledger_pick_supplier", "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.field_supplier = v
        st.session_state.ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
        _sync_planned_sale_from_ledger_in_stock_match()


def _on_ledger_pick_inventory_category() -> None:
    v = st.session_state.get("ledger_pick_inventory_category", "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.field_inventory_category = v
        st.session_state.ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER


def _apply_ledger_row_to_purchase_session(row: pd.Series) -> None:
    """入力補助で選んだ台帳1行から、仕入フォームの主項目へ反映する。"""
    pn = str(row.get(COL_NAME, "") or "").strip()
    su = str(row.get(COL_SUPPLIER, "") or "").strip()
    if pn:
        st.session_state.field_product_name = pn
    if su:
        st.session_state.field_supplier = su
    if COL_CATEGORY in row.index:
        rc = str(row.get(COL_CATEGORY, "") or "").strip()
        if rc:
            st.session_state.field_inventory_category = rc
    ly = _finite_int(row.get(COL_PRICE_EXCL), 0)
    if ly > 0:
        st.session_state.field_line_excl_yen = ly
    pl = _finite_int(row.get(COL_PLANNED_SALE), 0)
    st.session_state.field_planned_sale_excl = max(0, int(pl))


def _on_ledger_pick_management_id() -> None:
    v = str(st.session_state.get("ledger_pick_management_id", "") or "")
    if not v or v == LEDGER_PICK_PLACEHOLDER:
        return
    try:
        df = load_inventory_dataframe()
    except Exception:
        st.session_state.ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
        return
    if df is None or df.empty or COL_MANAGEMENT_ID not in df.columns:
        st.session_state.ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
        return
    m = df[COL_MANAGEMENT_ID].astype(str).str.strip() == v
    hits = df.loc[m]
    if len(hits) != 1:
        st.session_state.ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
        return
    _apply_ledger_row_to_purchase_session(hits.iloc[0])
    st.session_state["_gemini_match_management_id"] = v.strip()
    st.session_state.ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER


def _on_sale_pick_source_id() -> None:
    v = st.session_state.get("sale_pick_source_id", "")
    if v and v != LEDGER_PICK_PLACEHOLDER:
        st.session_state.field_sale_source_mgmt_id = v
        st.session_state.sale_pick_source_id = LEDGER_PICK_PLACEHOLDER


def _ledger_row_content_differs(
    a: pd.Series, b: pd.Series, *, skip_cols: frozenset[str]
) -> bool:
    """同一行の表示内容に差があるか（日時列などは比較から除外）。"""
    for c in EXPECTED_HEADERS:
        if c in skip_cols:
            continue
        if c not in a.index or c not in b.index:
            continue
        if str(a.get(c, "")).strip() != str(b.get(c, "")).strip():
            return True
    return False


def _stamp_row_datetime_on_changes(
    before: pd.DataFrame, after: pd.DataFrame
) -> pd.DataFrame:
    """行順・行数が一致する前提で、内容が変わった行の「日時」を JST 現在にする。"""
    out = after.copy()
    if COL_DATETIME not in out.columns or before is None or before.empty or after.empty:
        return out
    ts = jst_now_str()
    skip = frozenset({COL_DATETIME})
    n = min(len(before), len(after))
    loc_dt = out.columns.get_loc(COL_DATETIME)
    for i in range(n):
        if _ledger_row_content_differs(before.iloc[i], after.iloc[i], skip_cols=skip):
            out.iloc[i, loc_dt] = ts
    for i in range(n, len(after)):
        out.iloc[i, loc_dt] = ts
    return out


def _cell_value_for_sheet(v: Any) -> Any:
    """スプレッドシート1セル向けに欠損を正規化する。数値の NaN は 0、それ以外の欠損は空文字。"""
    try:
        if pd.api.types.is_scalar(v) and pd.isna(v):
            if isinstance(v, (float, np.floating)):
                return 0
            return ""
        if isinstance(v, (float, np.floating)) and (
            not math.isfinite(float(v)) or pd.isna(v)
        ):
            return 0
    except Exception:
        pass
    return v


def overwrite_inventory_worksheet_from_dataframe(
    df: pd.DataFrame, *, previous_df: pd.DataFrame | None = None
) -> None:
    """編集後の DataFrame で inventory.csv またはワークシートを全置換する。

    ``previous_df`` を渡したとき、同一行位置で内容が変わった行の **日時** を保存実行の JST に更新する。
    """
    work = df
    if previous_df is not None and COL_DATETIME in df.columns:
        work = _stamp_row_datetime_on_changes(
            previous_df.reset_index(drop=True),
            df.reset_index(drop=True),
        )
    out = _recalc_gross_profit_dataframe(
        work.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    )
    if _uses_local_inventory_csv():
        _inventory_csv_write_df(out)
        return
    ws = _get_or_create_inventory_worksheet()
    if ws is None:
        raise RuntimeError(
            f"スプレッドシートに接続できません。{SECRET_GOOGLE_SPREADSHEET_ID} とサービスアカウントを確認してください。"
        )
    values: list[list[Any]] = [EXPECTED_HEADERS]
    if not out.empty:
        for row in out[EXPECTED_HEADERS].to_numpy(dtype=object):
            values.append([_cell_value_for_sheet(x) for x in row.tolist()])
    try:
        ws.clear()
        ws.update("A1", values, value_input_option="USER_ENTERED")
        try:
            _apply_inventory_amount_number_formats(ws)
        except Exception:
            pass
    except Exception as e:
        raise RuntimeError(f"スプレッドシートの上書きに失敗しました: {e}") from e
    _bump_inventory_sheet_cache_bust()


def _ledger_df_loosen_numeric_columns_for_assignment(df: pd.DataFrame) -> None:
    """StringDtype 等の厳格な列に int を代入すると失敗するため、金額列を object に揃える（原地変更）。"""
    for _c in (
        COL_QTY,
        COL_PRICE_EXCL,
        COL_PRICE_INCL,
        COL_PLANNED_SALE,
        COL_PLANNED_SALE_INCL,
        COL_ACTUAL_SALE,
        COL_ACTUAL_SALE_INCL,
        COL_GROSS_PROFIT,
    ):
        if _c in df.columns:
            df[_c] = df[_c].astype(object)


def lookup_ledger_row_by_management_id(
    df: pd.DataFrame | None, management_id: str
) -> pd.Series | None:
    """管理IDに一致する台帳行を1件返す。無ければ None。"""
    if df is None or df.empty or COL_MANAGEMENT_ID not in df.columns:
        return None
    sid = (management_id or "").strip()
    if not sid:
        return None
    m = df[COL_MANAGEMENT_ID].astype(str).str.strip() == sid
    if not m.any():
        return None
    return df.loc[m].iloc[0]


def _split_management_ids_from_field(raw: str) -> list[str]:
    """管理IDの列挙文字列をカンマ・読点・区切り文字・空白・改行で分割し、空を除いたリストを返す。"""
    parts = re.split(r"[,、;；\s\n]+", (raw or "").strip())
    return [p.strip() for p in parts if p.strip()]


def apply_outbound_disposal_excluded_by_management_id(
    source_management_id: str,
    *,
    new_image_url: str = "",
    memo_suffix: str = "",
) -> None:
    """出庫（除外）かつ **対象外**: 在庫中の1行をステータス **対象外** に更新（新規行なし）。

    販売日時・出庫種別・日時は確定実行の JST。実売・販売画像はクリアして粗利を再計算する。
    """
    sid = (source_management_id or "").strip()
    if not sid:
        raise ValueError("管理IDが空です。")
    df_src = load_inventory_dataframe()
    if df_src is None or df_src.empty:
        raise RuntimeError("台帳を読み込めませんでした。")
    df_src = df_src.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    _ledger_df_loosen_numeric_columns_for_assignment(df_src)
    msk = df_src[COL_MANAGEMENT_ID].astype(str).str.strip() == sid
    if not msk.any():
        raise RuntimeError(f"管理ID {sid} の行が台帳に見つかりません。")
    if int(msk.sum()) != 1:
        raise RuntimeError(f"管理ID {sid} が複数行に重複しています。")
    cur_st = _normalize_stock_status(
        str(df_src.loc[msk, COL_STOCK_STATUS].iloc[0])
    )
    if cur_st != STATUS_IN_STOCK:
        raise RuntimeError(
            f"管理ID {sid} は「{cur_st}」のため、除外（対象外）の対象外です（在庫中の行のみ）。"
        )
    now_exec = jst_now_str()
    df_src.loc[msk, COL_DATETIME] = now_exec
    df_src.loc[msk, COL_STOCK_STATUS] = STATUS_EXCLUDED
    if COL_SALE_DATETIME in df_src.columns:
        df_src.loc[msk, COL_SALE_DATETIME] = now_exec
    if COL_SALE_OUTBOUND_TYPE in df_src.columns:
        df_src.loc[msk, COL_SALE_OUTBOUND_TYPE] = OUTBOUND_KIND_EXCLUDE
    df_src.loc[msk, COL_ACTUAL_SALE] = 0
    if COL_ACTUAL_SALE_INCL in df_src.columns:
        df_src.loc[msk, COL_ACTUAL_SALE_INCL] = 0
    if COL_SALE_IMAGE_URL in df_src.columns:
        df_src.loc[msk, COL_SALE_IMAGE_URL] = ""
    nu = (new_image_url or "").strip()
    if nu:
        df_src.loc[msk, COL_IMAGE_URL] = nu
    if (memo_suffix or "").strip():
        old_memo = str(df_src.loc[msk, COL_MEMO].iloc[0] or "").strip()
        tag = (memo_suffix or "").strip()
        df_src.loc[msk, COL_MEMO] = (old_memo + "\n" if old_memo else "") + tag
    df_src = _recalc_gross_profit_dataframe(df_src)
    overwrite_inventory_worksheet_from_dataframe(df_src)


def apply_outbound_sale_to_ledger_by_management_id(
    source_management_id: str,
    *,
    actual_sale_unit_excl_yen: int,
    new_image_url: str = "",
    memo_suffix: str = "",
    update_sale_voucher: bool = False,
    sale_voucher_recorded_at: str = "",
    sale_voucher_evidence_url: str = "",
    sale_outbound_type: str = "出庫（販売）",
    loan_datetime_jst: str | None = None,
) -> None:
    """在庫中の1行を販売済に更新（新規行なし）。A列「日時」は確定実行の JST。出庫種別は ``sale_outbound_type``
    （例: 出庫（販売）／出庫（除外）／出庫（浮貸）／出庫（戻入））。
    ``loan_datetime_jst`` を渡したときは **浮貸日時** 列にも記録する（出庫（戻入）で販売済にするときなど）。
    """
    sid = (source_management_id or "").strip()
    if not sid:
        raise ValueError("管理IDが空です。")
    df_src = load_inventory_dataframe()
    if df_src is None or df_src.empty:
        raise RuntimeError("台帳を読み込めませんでした。")
    df_src = df_src.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    _ledger_df_loosen_numeric_columns_for_assignment(df_src)
    msk = df_src[COL_MANAGEMENT_ID].astype(str).str.strip() == sid
    if not msk.any():
        raise RuntimeError(f"管理ID {sid} の行が台帳に見つかりません。")
    if int(msk.sum()) != 1:
        raise RuntimeError(f"管理ID {sid} が複数行に重複しています。")
    cur_st = _normalize_stock_status(
        str(df_src.loc[msk, COL_STOCK_STATUS].iloc[0])
    )
    if cur_st != STATUS_IN_STOCK:
        raise RuntimeError(
            f"管理ID {sid} は「{cur_st}」のため、販売反映の対象外です（在庫中の行のみ更新します）。"
        )
    av = _finite_int(actual_sale_unit_excl_yen, 0)
    if av < 0:
        raise RuntimeError("実売金額（税抜）は0円以上にしてください。")

    now_exec = jst_now_str()
    _prev_row_dt = str(df_src.loc[msk, COL_DATETIME].iloc[0] or "")
    if COL_PURCHASE_DATETIME in df_src.columns:
        cp = str(df_src.loc[msk, COL_PURCHASE_DATETIME].iloc[0] or "").strip()
        if not cp:
            df_src.loc[msk, COL_PURCHASE_DATETIME] = _prev_row_dt
    if COL_PURCHASE_MOVEMENT in df_src.columns:
        cm = str(df_src.loc[msk, COL_PURCHASE_MOVEMENT].iloc[0] or "").strip()
        if not cm:
            df_src.loc[msk, COL_PURCHASE_MOVEMENT] = "入庫（購入）"
    if COL_SALE_DATETIME in df_src.columns:
        df_src.loc[msk, COL_SALE_DATETIME] = now_exec
    if COL_SALE_OUTBOUND_TYPE in df_src.columns:
        _ot = (sale_outbound_type or "").strip() or "出庫（販売）"
        df_src.loc[msk, COL_SALE_OUTBOUND_TYPE] = _ot
    if loan_datetime_jst is not None and COL_LOAN_DATETIME in df_src.columns:
        _ld = (loan_datetime_jst or "").strip()
        if _ld:
            df_src.loc[msk, COL_LOAN_DATETIME] = _ld
    df_src.loc[msk, COL_DATETIME] = now_exec
    df_src.loc[msk, COL_STOCK_STATUS] = STATUS_SOLD
    df_src.loc[msk, COL_ACTUAL_SALE] = av
    cur_img = str(df_src.loc[msk, COL_IMAGE_URL].iloc[0] or "").strip()
    nu = (new_image_url or "").strip()
    if not cur_img and nu:
        df_src.loc[msk, COL_IMAGE_URL] = nu
    if COL_SALE_IMAGE_URL in df_src.columns and nu:
        df_src.loc[msk, COL_SALE_IMAGE_URL] = nu
    if (memo_suffix or "").strip():
        old_memo = str(df_src.loc[msk, COL_MEMO].iloc[0] or "").strip()
        tag = (memo_suffix or "").strip()
        df_src.loc[msk, COL_MEMO] = (old_memo + "\n" if old_memo else "") + tag
    if update_sale_voucher and (sale_voucher_evidence_url or "").strip():
        df_src.loc[msk, COL_VOUCHER_RECORDED_AT] = (
            sale_voucher_recorded_at or ""
        ).strip()
        df_src.loc[msk, COL_VOUCHER_EVIDENCE_URL] = sale_voucher_evidence_url.strip()
    df_src = _recalc_gross_profit_dataframe(df_src)
    overwrite_inventory_worksheet_from_dataframe(df_src)


def apply_outbound_loan_in_stock_datetime_by_management_id(
    source_management_id: str,
    *,
    loan_datetime_jst: str | None = None,
    new_image_url: str = "",
    memo_suffix: str = "",
) -> None:
    """出庫（浮貸）かつ在庫中のまま: 該当行の **浮貸日時** を記録し、出庫種別を記録。新規行は追加しない。"""
    sid = (source_management_id or "").strip()
    if not sid:
        raise ValueError("管理IDが空です。")
    df_src = load_inventory_dataframe()
    if df_src is None or df_src.empty:
        raise RuntimeError("台帳を読み込めませんでした。")
    df_src = df_src.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    _ledger_df_loosen_numeric_columns_for_assignment(df_src)
    msk = df_src[COL_MANAGEMENT_ID].astype(str).str.strip() == sid
    if not msk.any():
        raise RuntimeError(f"管理ID {sid} の行が台帳に見つかりません。")
    if int(msk.sum()) != 1:
        raise RuntimeError(f"管理ID {sid} が複数行に重複しています。")
    cur_st = _normalize_stock_status(
        str(df_src.loc[msk, COL_STOCK_STATUS].iloc[0])
    )
    if cur_st != STATUS_IN_STOCK:
        raise RuntimeError(
            f"管理ID {sid} は「{cur_st}」のため、浮貸（在庫中）の記録対象外です（在庫中の行のみ）。"
        )
    loan_dt = (loan_datetime_jst or "").strip() or jst_now_str()
    now_exec = jst_now_str()
    if COL_LOAN_DATETIME in df_src.columns:
        df_src.loc[msk, COL_LOAN_DATETIME] = loan_dt
    if COL_SALE_OUTBOUND_TYPE in df_src.columns:
        df_src.loc[msk, COL_SALE_OUTBOUND_TYPE] = "出庫（浮貸）"
    df_src.loc[msk, COL_DATETIME] = now_exec
    cur_img = str(df_src.loc[msk, COL_IMAGE_URL].iloc[0] or "").strip()
    nu = (new_image_url or "").strip()
    if not cur_img and nu:
        df_src.loc[msk, COL_IMAGE_URL] = nu
    if (memo_suffix or "").strip():
        old_memo = str(df_src.loc[msk, COL_MEMO].iloc[0] or "").strip()
        tag = (memo_suffix or "").strip()
        df_src.loc[msk, COL_MEMO] = (old_memo + "\n" if old_memo else "") + tag
    df_src = _recalc_gross_profit_dataframe(df_src)
    overwrite_inventory_worksheet_from_dataframe(df_src)


def apply_outbound_sale_return_to_in_stock_by_management_id(
    source_management_id: str,
    *,
    new_image_url: str = "",
    memo_suffix: str = "",
) -> None:
    """出庫（返品）: **販売済** の行を **在庫中** に戻す（新規行なし）。

    **販売日時** に確定実行の JST を記録し、**出庫種別** に「出庫（返品）」を入れる。
    実売・販売画像URLはクリアし、粗利を再計算する。
    """
    sid = (source_management_id or "").strip()
    if not sid:
        raise ValueError("管理IDが空です。")
    df_src = load_inventory_dataframe()
    if df_src is None or df_src.empty:
        raise RuntimeError("台帳を読み込めませんでした。")
    df_src = df_src.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    _ledger_df_loosen_numeric_columns_for_assignment(df_src)
    msk = df_src[COL_MANAGEMENT_ID].astype(str).str.strip() == sid
    if not msk.any():
        raise RuntimeError(f"管理ID {sid} の行が台帳に見つかりません。")
    if int(msk.sum()) != 1:
        raise RuntimeError(f"管理ID {sid} が複数行に重複しています。")
    cur_st = _normalize_stock_status(
        str(df_src.loc[msk, COL_STOCK_STATUS].iloc[0])
    )
    if cur_st != STATUS_SOLD:
        raise RuntimeError(
            f"管理ID {sid} は「{cur_st}」のため、出庫（返品）の対象外です（販売済の行のみ）。"
        )
    now_exec = jst_now_str()
    df_src.loc[msk, COL_DATETIME] = now_exec
    df_src.loc[msk, COL_STOCK_STATUS] = STATUS_IN_STOCK
    if COL_SALE_DATETIME in df_src.columns:
        df_src.loc[msk, COL_SALE_DATETIME] = now_exec
    if COL_SALE_OUTBOUND_TYPE in df_src.columns:
        df_src.loc[msk, COL_SALE_OUTBOUND_TYPE] = "出庫（返品）"
    df_src.loc[msk, COL_ACTUAL_SALE] = 0
    if COL_ACTUAL_SALE_INCL in df_src.columns:
        df_src.loc[msk, COL_ACTUAL_SALE_INCL] = 0
    if COL_SALE_IMAGE_URL in df_src.columns:
        df_src.loc[msk, COL_SALE_IMAGE_URL] = ""
    if COL_LOAN_DATETIME in df_src.columns:
        df_src.loc[msk, COL_LOAN_DATETIME] = ""
    nu = (new_image_url or "").strip()
    if nu:
        df_src.loc[msk, COL_IMAGE_URL] = nu
    if (memo_suffix or "").strip():
        old_memo = str(df_src.loc[msk, COL_MEMO].iloc[0] or "").strip()
        tag = (memo_suffix or "").strip()
        df_src.loc[msk, COL_MEMO] = (old_memo + "\n" if old_memo else "") + tag
    df_src = _recalc_gross_profit_dataframe(df_src)
    overwrite_inventory_worksheet_from_dataframe(df_src)


def apply_outbound_receipt_in_stock_by_management_id(
    source_management_id: str,
    *,
    receipt_datetime_jst: str | None = None,
    new_image_url: str = "",
    memo_suffix: str = "",
) -> None:
    """出庫（戻入）かつ在庫中のまま: **浮貸日時** に記録し、出庫種別を「出庫（戻入）」にする。"""
    sid = (source_management_id or "").strip()
    if not sid:
        raise ValueError("管理IDが空です。")
    df_src = load_inventory_dataframe()
    if df_src is None or df_src.empty:
        raise RuntimeError("台帳を読み込めませんでした。")
    df_src = df_src.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    _ledger_df_loosen_numeric_columns_for_assignment(df_src)
    msk = df_src[COL_MANAGEMENT_ID].astype(str).str.strip() == sid
    if not msk.any():
        raise RuntimeError(f"管理ID {sid} の行が台帳に見つかりません。")
    if int(msk.sum()) != 1:
        raise RuntimeError(f"管理ID {sid} が複数行に重複しています。")
    cur_st = _normalize_stock_status(
        str(df_src.loc[msk, COL_STOCK_STATUS].iloc[0])
    )
    if cur_st != STATUS_IN_STOCK:
        raise RuntimeError(
            f"管理ID {sid} は「{cur_st}」のため、出庫（戻入）・在庫中の対象外です。"
        )
    loan_dt = (receipt_datetime_jst or "").strip() or jst_now_str()
    now_exec = jst_now_str()
    if COL_LOAN_DATETIME in df_src.columns:
        df_src.loc[msk, COL_LOAN_DATETIME] = loan_dt
    if COL_SALE_OUTBOUND_TYPE in df_src.columns:
        df_src.loc[msk, COL_SALE_OUTBOUND_TYPE] = "出庫（戻入）"
    df_src.loc[msk, COL_DATETIME] = now_exec
    nu = (new_image_url or "").strip()
    if nu:
        df_src.loc[msk, COL_IMAGE_URL] = nu
    if (memo_suffix or "").strip():
        old_memo = str(df_src.loc[msk, COL_MEMO].iloc[0] or "").strip()
        tag = (memo_suffix or "").strip()
        df_src.loc[msk, COL_MEMO] = (old_memo + "\n" if old_memo else "") + tag
    df_src = _recalc_gross_profit_dataframe(df_src)
    overwrite_inventory_worksheet_from_dataframe(df_src)


def apply_last_stocktake_jst_for_management_ids(
    management_ids: Iterable[str],
) -> tuple[int, list[str]]:
    """複数の在庫中の行について棚卸日を本日（JST）にし、読込・保存で反映する。

    件数が多いときは分割して保存し、スプレッドシートAPIの制限・タイムアウトを避ける。

    戻り値: (更新に成功した件数, スキップ理由の短文リスト。重複・不在・在庫外などはスキップ)
    """
    ids_sorted = sorted({str(x).strip() for x in management_ids if str(x).strip()})
    if not ids_sorted:
        raise ValueError("管理IDが1件以上必要です。")
    if len(ids_sorted) <= STOCKTAKE_SHEET_SAVE_MAX_IDS:
        return _apply_last_stocktake_jst_for_management_ids_one_save(ids_sorted)
    total_ok = 0
    all_skips: list[str] = []
    for i in range(0, len(ids_sorted), STOCKTAKE_SHEET_SAVE_MAX_IDS):
        chunk = ids_sorted[i : i + STOCKTAKE_SHEET_SAVE_MAX_IDS]
        n_ok, sk = _apply_last_stocktake_jst_for_management_ids_one_save(chunk)
        total_ok += n_ok
        all_skips.extend(sk)
    return total_ok, all_skips


def _apply_last_stocktake_jst_for_management_ids_one_save(
    ids_sorted: list[str],
) -> tuple[int, list[str]]:
    ids = set(ids_sorted)
    df_src = load_inventory_dataframe()
    if df_src is None or df_src.empty:
        raise RuntimeError("台帳を読み込めませんでした。")
    df_src = df_src.reindex(columns=EXPECTED_HEADERS, fill_value="").copy()
    _ledger_df_loosen_numeric_columns_for_assignment(df_src)
    today_s = _today_jst_date().isoformat()
    now_exec = jst_now_str()
    updated: set[str] = set()
    skips: list[str] = []
    for sid in sorted(ids):
        msk = df_src[COL_MANAGEMENT_ID].astype(str).str.strip() == sid
        if not msk.any():
            skips.append(f"{sid}: 台帳に見つかりません")
            continue
        if int(msk.sum()) != 1:
            skips.append(f"{sid}: 管理IDが重複しています")
            continue
        cur_st = _normalize_stock_status(
            str(df_src.loc[msk, COL_STOCK_STATUS].iloc[0])
        )
        if cur_st != STATUS_IN_STOCK:
            skips.append(f"{sid}: 在庫中ではない（{cur_st}）")
            continue
        if COL_LAST_STOCKTAKE in df_src.columns:
            df_src.loc[msk, COL_LAST_STOCKTAKE] = today_s
        df_src.loc[msk, COL_DATETIME] = now_exec
        updated.add(sid)
    if not updated:
        raise RuntimeError(
            "更新できる在庫中の行がありませんでした。"
            + (f" 詳細: {'; '.join(skips[:6])}" if skips else "")
        )
    df_src = _recalc_gross_profit_dataframe(df_src)
    overwrite_inventory_worksheet_from_dataframe(df_src)
    _inv_stocktake_work_remaining_note_done(updated)
    return len(updated), skips


def apply_last_stocktake_jst_for_management_id(management_id: str) -> None:
    """在庫中の1行について「最後に確認した日付（棚卸日）」を本日（JST）にし、日時を更新して保存する。"""
    apply_last_stocktake_jst_for_management_ids([management_id])


def _apply_ledger_sort(
    df: pd.DataFrame,
    primary: str,
    primary_asc: bool,
    secondary: str,
    secondary_asc: bool,
    tertiary: str,
    tertiary_asc: bool,
) -> pd.DataFrame:
    """在庫一覧の表示用ソート（最大3キー、コピーを返す）。"""
    if df.empty:
        return df
    col_map = {
        "日時": COL_DATETIME,
        "仕入先・取引先": COL_SUPPLIER,
        "管理ID": COL_MANAGEMENT_ID,
        "仕入日時": COL_PURCHASE_DATETIME,
        "販売日時": COL_SALE_DATETIME,
    }
    pairs: list[tuple[str, bool]] = []
    for label, asc in (
        (primary, primary_asc),
        (secondary, secondary_asc),
        (tertiary, tertiary_asc),
    ):
        if label == "なし" or label not in col_map:
            continue
        col = col_map[label]
        if any(p[0] == col for p in pairs):
            continue
        pairs.append((col, asc))
    if not pairs:
        return df.copy()

    out = df.copy()
    sort_cols: list[str] = []
    ascending: list[bool] = []
    for col, asc in pairs:
        if col not in out.columns:
            continue
        if col in (COL_DATETIME, COL_PURCHASE_DATETIME, COL_SALE_DATETIME):
            tmp = "_sort_dt_internal"
            if col != COL_DATETIME:
                tmp = f"_sort_dt_internal_{col}"
            out[tmp] = pd.to_datetime(out[col], errors="coerce")
            sort_cols.append(tmp)
        else:
            sort_cols.append(col)
        ascending.append(asc)
    if not sort_cols:
        return out
    out = out.sort_values(by=sort_cols, ascending=ascending, na_position="last")
    return out.drop(columns=[c for c in out.columns if c.startswith("_sort_dt")], errors="ignore")


def _evidence_url_column_all_http_or_blank(s: pd.Series) -> bool:
    """証憑URL列を LinkColumn にできるか（空・欠損・http(s) のみ）。"""
    t = s.fillna("").astype(str).str.strip().str.lower()
    t = t.replace({"nan": "", "none": "", "<na>": "", "nat": ""}, regex=False)
    nonempty = t[t != ""]
    if nonempty.empty:
        return True
    ok = nonempty.str.startswith("http://") | nonempty.str.startswith("https://")
    return bool(ok.all())


def _normalize_evidence_urls_for_link_editor(df: pd.DataFrame, col: str) -> bool:
    """証憑URL列を LinkColumn 向けに正規化（空相当→None）。LinkColumn を使うなら True。"""
    if col not in df.columns:
        return False
    if not _evidence_url_column_all_http_or_blank(df[col]):
        return False
    v = df[col]
    strv = v.fillna("").astype(str).str.strip()
    low = strv.str.lower()
    isblank = v.isna() | low.isin(("", "nan", "none", "<na>", "nat"))
    nb = isblank.fillna(False).to_numpy(dtype=bool, copy=False)
    arr = np.empty(len(strv), dtype=object)
    arr[nb] = None
    arr[~nb] = strv[~isblank].to_numpy(dtype=object, copy=False)
    df[col] = arr
    return True


def _ledger_dashboard_axis_datetime(df: pd.DataFrame) -> pd.Series:
    """ダッシュボードの期間・月次軸に使う日時。販売済は **販売日時**（空なら日時）を優先。"""
    base = pd.to_datetime(df[COL_DATETIME], errors="coerce")
    if COL_SALE_DATETIME not in df.columns or COL_STOCK_STATUS not in df.columns:
        return base
    st = df[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
    sold = st == STATUS_SOLD
    sd = pd.to_datetime(df[COL_SALE_DATETIME], errors="coerce")
    alt = sd.combine_first(base)
    return base.where(~sold, alt)


def _prepare_ledger_analysis(df: pd.DataFrame) -> pd.DataFrame:
    """入出庫判定・行金額・年月などの派生列を付与したコピーを返す。

    1点1行ライフサイクル前提: **在庫中** かつ **入庫** だけを仕入（入庫）側に計上し、
    **販売済** かつ実売がある行は **出庫側の金額を実売ではなく当行の仕入金額（税抜・税込）** で計上する
    （「入庫種別」が空の在庫中は入庫扱い。数量の出庫計上は販売済＋実売で行い、金額の入出庫は仕入列で統一）。
    **在庫中** のまま **出庫**（浮貸など）の行は仕入列ベースで出庫に含める。
    """
    d = df.copy()
    if d.empty:
        return d
    d[COL_DATETIME] = pd.to_datetime(d[COL_DATETIME], errors="coerce")
    if COL_QTY in d.columns:
        qty = (
            pd.to_numeric(d[COL_QTY], errors="coerce")
            .fillna(1.0)
            .clip(lower=1.0)
            .astype(float)
        )
    else:
        qty = pd.Series(1.0, index=d.index, dtype=float)
    line_stored_ex = _series_to_numeric_loose(d[COL_PRICE_EXCL]).fillna(0)
    line_stored_in = _series_to_numeric_loose(d[COL_PRICE_INCL]).fillna(0)
    if COL_STOCK_STATUS in d.columns:
        st_col = d[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
    else:
        st_col = pd.Series(STATUS_IN_STOCK, index=d.index)
    pur_mv = (
        d[COL_PURCHASE_MOVEMENT].astype(str).str.strip()
        if COL_PURCHASE_MOVEMENT in d.columns
        else pd.Series("", index=d.index, dtype=str)
    )
    is_in_mv = pur_mv.str.startswith("入庫") | (
        (pur_mv == "") & (st_col == STATUS_IN_STOCK)
    )
    is_out_pur = pur_mv.str.startswith("出庫")
    # 税抜・税込はいずれも行合計として解釈（仕入単価列は廃止）
    line_ex = line_stored_ex.astype(float)
    mask_ex_from_incl = (
        (line_ex.fillna(0) <= 0)
        & (qty.fillna(0) > 0)
        & (line_stored_in.fillna(0) > 0)
    )
    if bool(mask_ex_from_incl.any()):
        fill_ex = line_stored_in.map(
            lambda v: float(_estimate_excl_yen_from_incl_yen(_finite_int(v, 0)))
        )
        line_ex = line_ex.where(~mask_ex_from_incl, fill_ex)
    line_in = line_stored_in.astype(float)
    line_needs_derive_in = line_stored_in.fillna(0) <= 0
    _ex_int = line_ex.fillna(0).round().clip(lower=0).astype(int)
    line_in = line_in.mask(line_needs_derive_in, _ex_int.map(price_incl_tax).astype(float))
    line_ex = line_ex.fillna(0).replace([np.inf, -np.inf], 0)
    line_in = line_in.fillna(0).replace([np.inf, -np.inf], 0)

    ac_u = _series_to_numeric_loose(
        d[COL_ACTUAL_SALE] if COL_ACTUAL_SALE in d.columns else 0
    ).fillna(0)
    rev_ex = (ac_u * qty).clip(lower=0)
    m_stock_in = is_in_mv & (st_col == STATUS_IN_STOCK)
    # 販売済は実売ベースで売上計上（区分がまだ入庫の旧行も、実売があればここに含める）
    m_sold_rev = (st_col == STATUS_SOLD) & (rev_ex > 0)
    m_float_out = is_out_pur & (st_col == STATUS_IN_STOCK)

    d["_qty_in"] = qty.where(m_stock_in, 0.0).fillna(0).astype(float)
    d["_amt_ex_in"] = line_ex.where(m_stock_in, 0.0).fillna(0).astype(float)
    d["_amt_in_in"] = line_in.where(m_stock_in, 0.0).fillna(0).astype(float)

    d["_qty_out"] = qty.where(m_sold_rev | m_float_out, 0.0).fillna(0).astype(float)
    # 金額の入出庫はいずれも当行の仕入金額（税抜・税込）で統一（販売済の実売は金額集計に使わない）
    d["_amt_ex_out"] = (
        line_ex.where(m_sold_rev, 0.0).fillna(0)
        + line_ex.where(m_float_out, 0.0).fillna(0)
    ).astype(float)
    d["_amt_in_out"] = (
        line_in.where(m_sold_rev, 0.0).fillna(0)
        + line_in.where(m_float_out, 0.0).fillna(0)
    ).astype(float)

    axis_dt = _ledger_dashboard_axis_datetime(d)
    d["_ym"] = axis_dt.dt.to_period("M").astype(str)
    d["_year"] = axis_dt.dt.year
    d["_month"] = axis_dt.dt.month
    return d


def _dash_ratio_pct(numer: int, denom: int) -> str:
    """入庫÷合計の比率表示（分母0はダッシュ）。"""
    if denom <= 0:
        return "—"
    return f"{100.0 * float(numer) / float(denom):.1f}%"


def _ledger_dashboard_date_bounds(df: pd.DataFrame) -> tuple[date, date]:
    """ダッシュボード用の日付範囲既定値（販売済は販売日時軸を含む）。"""
    s = _ledger_dashboard_axis_datetime(df).dropna()
    if s.empty:
        t = jst_now().date()
        return t, t
    return s.min().date(), s.max().date()


def _altair_y_scale_positive(s: pd.Series) -> alt.Scale:
    """金額がすべて 0 のときでも棒グラフが潰れないよう Y 軸上限を確保する。"""
    v = pd.to_numeric(s, errors="coerce").replace([np.inf, -np.inf], np.nan).fillna(0.0)
    hi = float(v.max())
    if not math.isfinite(hi):
        hi = 0.0
    top = max(hi * 1.08, 1.0)
    return alt.Scale(domain=[0.0, top])


def render_ledger_dashboard(df: pd.DataFrame) -> None:
    """在庫台帳 DataFrame から入出庫集計・仕入先・取引先別・グラフを表示する。"""
    st.subheader("集計")
    st.caption(
        "上の表の現在の内容（未保存の編集を含む）を集計します。"
        f"入出庫の **金額** は「{COL_PRICE_EXCL}」「{COL_PRICE_INCL}」（仕入・税抜・税込）を行合計として集計します（販売済行の出庫側も実売ではなく当行の仕入金額で計上）。"
        f"仕入先・取引先別の粗利は「{COL_GROSS_PROFIT}」列を合算しています（税抜・台帳保存時の値）。"
        f"入出庫の **数量** は「{COL_QTY}」列（空・未入力は **1**）を在庫中の入庫行・販売済で実売がある行で合算します。"
        "期間フィルタと月次の軸は **販売済は販売日時**（未入力の旧行は日時にフォールバック）、在庫中は **日時** です。"
        "（税抜の仕入金額が空で税込だけある行は、10%/8%/非課税のいずれかに税込が一致する税抜を逆算します。"
        "カンマ区切り・円記号付きの数値も読み取ります。）"
    )
    if df.empty:
        st.info("集計する行がありません。")
        return

    df_in = df.copy()
    if COL_QTY in df_in.columns:
        df_in[COL_QTY] = (
            _series_to_numeric_loose(df_in[COL_QTY])
            .replace([np.inf, -np.inf], np.nan)
            .fillna(1)
            .clip(lower=1)
        )
    for _col in (COL_PRICE_EXCL, COL_PRICE_INCL, COL_GROSS_PROFIT):
        if _col in df_in.columns:
            df_in[_col] = (
                _series_to_numeric_loose(df_in[_col])
                .replace([np.inf, -np.inf], np.nan)
                .fillna(0)
            )
    ad = _prepare_ledger_analysis(df_in)
    ad_f = ad.dropna(subset=[COL_DATETIME], how="all")

    d_lo, d_hi = _ledger_dashboard_date_bounds(ad_f)
    p1, p2, p3 = st.columns([1, 1, 2])
    with p1:
        date_from = st.date_input(
            "開始日（From）",
            value=d_lo,
            min_value=date(1970, 1, 1),
            max_value=date(2100, 12, 31),
            key="dash_date_from",
        )
    with p2:
        date_to = st.date_input(
            "終了日（To）",
            value=d_hi,
            min_value=date(1970, 1, 1),
            max_value=date(2100, 12, 31),
            key="dash_date_to",
        )
    with p3:
        supplier_filter = st.multiselect(
            "仕入先・取引先で絞り込み（未選択は全件）",
            options=sorted(ad_f[COL_SUPPLIER].fillna("").astype(str).unique().tolist()),
            key="dash_supplier_filter",
        )

    dfb = date_from
    dtb = date_to
    if dfb > dtb:
        st.warning("開始日が終了日より後です。入れ替えて集計します。")
        dfb, dtb = dtb, dfb

    row_ts = _ledger_dashboard_axis_datetime(ad_f)
    row_day = row_ts.dt.normalize()
    from_ts = pd.Timestamp(datetime.combine(dfb, datetime.min.time()))
    to_ts = pd.Timestamp(datetime.combine(dtb, datetime.min.time()))
    flt = ad_f[(row_day >= from_ts) & (row_day <= to_ts)]
    if supplier_filter:
        flt = flt[flt[COL_SUPPLIER].astype(str).isin(supplier_filter)]

    period_cogs_ex = 0
    if COL_PRICE_EXCL in flt.columns:
        period_cogs_ex = _finite_int(
            _series_to_numeric_loose(flt[COL_PRICE_EXCL]).fillna(0).sum(),
            0,
        )
    period_cogs_in = 0
    if COL_PRICE_INCL in flt.columns:
        period_cogs_in = _finite_int(
            _series_to_numeric_loose(flt[COL_PRICE_INCL]).fillna(0).sum(),
            0,
        )
    gp_sold_period = 0
    if not flt.empty and COL_STOCK_STATUS in flt.columns:
        _msf = (
            flt[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
            == STATUS_SOLD
        )
        gp_sold_period = _finite_int(
            _series_to_numeric_loose(flt.loc[_msf, COL_GROSS_PROFIT])
            .fillna(0)
            .sum(),
            0,
        )
    period_sale_excl = 0
    period_sale_incl = 0
    if not flt.empty and COL_STOCK_STATUS in flt.columns:
        _sold_sub = flt.loc[
            flt[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
            == STATUS_SOLD
        ]
        if (
            not _sold_sub.empty
            and COL_ACTUAL_SALE in _sold_sub.columns
        ):
            _au = _series_to_numeric_loose(_sold_sub[COL_ACTUAL_SALE]).fillna(0)
            if COL_QTY in _sold_sub.columns:
                _qv = _sold_sub[COL_QTY].map(lambda x: max(1, _finite_int(x, 1)))
            else:
                _qv = pd.Series(1, index=_sold_sub.index, dtype=int)
            _m_au = _au > 0
            period_sale_excl = _finite_int(
                (_au * _qv.astype(np.int64)).loc[_m_au].sum(),
                0,
            )
            if COL_ACTUAL_SALE_INCL in _sold_sub.columns:
                period_sale_incl = _finite_int(
                    _series_to_numeric_loose(
                        _sold_sub.loc[_m_au, COL_ACTUAL_SALE_INCL]
                    )
                    .fillna(0)
                    .sum(),
                    0,
                )
    st.markdown("##### ライフサイクル指標")
    st.caption(
        "次の **5つ** は **From〜To・仕入先フィルタ** に合致する行のみを対象にします（数量列による追加の絞り込みはありません）。"
        f"原価は台帳の **{COL_PRICE_EXCL}** / **{COL_PRICE_INCL}** を合算した税抜・税込の総額です。"
        "実売は **販売済** 行の **実売金額（税抜・税込）** を行計で合算したものです（実売単価が 0 の行は含みません）。"
        "確定粗利は **販売済** 行の粗利列（税抜）の合計です。"
    )
    _lc1, _lc2, _lc3, _lc4, _lc5 = st.columns(5)
    with _lc1:
        st.metric("税抜原価総額（期間内）", f"¥{period_cogs_ex:,}")
    with _lc2:
        st.metric("税込原価総額（期間内）", f"¥{period_cogs_in:,}")
    with _lc3:
        st.metric(
            "実売金額合計（税抜・期間内・販売済）",
            f"¥{period_sale_excl:,}",
        )
    with _lc4:
        st.metric(
            "実売金額合計（税込・期間内・販売済）",
            f"¥{period_sale_incl:,}",
        )
    with _lc5:
        st.metric("確定粗利（期間内・販売済・税抜）", f"¥{gp_sold_period:,}")

    if flt.empty:
        st.warning(
            "条件に一致するデータがありません。"
            "From〜To の日付範囲または仕入先・取引先の絞り込みを見直してください。"
        )
        return

    q_in = _finite_int(flt["_qty_in"].sum(), 0)
    q_out = _finite_int(flt["_qty_out"].sum(), 0)
    q_sum = q_in + q_out
    ex_in = _finite_int(flt["_amt_ex_in"].sum(), 0)
    ex_out = _finite_int(flt["_amt_ex_out"].sum(), 0)
    in_in = _finite_int(flt["_amt_in_in"].sum(), 0)
    in_out = _finite_int(flt["_amt_in_out"].sum(), 0)

    st.markdown("##### 期間内の数量")
    q_ratio = _dash_ratio_pct(q_in, q_sum)
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("入庫 合計数量", f"{q_in:,}")
    m2.metric("出庫 合計数量", f"{q_out:,}")
    m3.metric("合計数量（入+出）", f"{q_sum:,}")
    m4.metric("入庫÷合計（数量比率）", q_ratio)

    ex_sum = ex_in + ex_out
    in_sum = in_in + in_out
    ex_ratio = _dash_ratio_pct(ex_in, ex_sum)
    in_ratio = _dash_ratio_pct(in_in, in_sum)

    st.markdown("##### 期間内の金額（仕入ベース・税抜）")
    m5, m6, m10, m11 = st.columns(4)
    m5.metric("入庫 合計金額（税抜）", f"¥{ex_in:,}")
    m6.metric("出庫 合計金額（税抜）", f"¥{ex_out:,}")
    m10.metric("合計金額（入+出・税抜）", f"¥{ex_sum:,}")
    m11.metric("入庫÷合計（税抜比率）", ex_ratio)
    if COL_GROSS_PROFIT in flt.columns:
        gp_tot = _finite_int(
            _series_to_numeric_loose(flt[COL_GROSS_PROFIT]).fillna(0).sum(), 0
        )
        st.metric("粗利合計（税抜）", f"¥{gp_tot:,}")
    else:
        st.metric("粗利合計（税抜）", "—")

    st.markdown("##### 期間内の金額（仕入ベース・税込）")
    m7, m8, m12, m13 = st.columns(4)
    m7.metric("入庫 合計金額（税込）", f"¥{in_in:,}")
    m8.metric("出庫 合計金額（税込）", f"¥{in_out:,}")
    m12.metric("合計金額（入+出・税込）", f"¥{in_sum:,}")
    m13.metric("入庫÷合計（税込比率）", in_ratio)

    st.markdown("##### 仕入先・取引先別サマリー（税抜金額・数量・粗利）")
    sup_col = "仕入先・取引先"

    def _supplier_grp(part: pd.DataFrame) -> pd.DataFrame:
        if part.empty:
            cols = [
                sup_col,
                "入庫数量",
                "出庫数量",
                "入庫金額税抜",
                "出庫金額税抜",
            ]
            if COL_GROSS_PROFIT in flt.columns:
                cols.append("粗利合計")
            return pd.DataFrame(columns=cols)
        g = part.assign(**{sup_col: part[COL_SUPPLIER].fillna("(未設定)").astype(str)})
        _agg_sup: dict[str, tuple[str, str]] = {
            "入庫数量": ("_qty_in", "sum"),
            "出庫数量": ("_qty_out", "sum"),
            "入庫金額税抜": ("_amt_ex_in", "sum"),
            "出庫金額税抜": ("_amt_ex_out", "sum"),
        }
        if COL_GROSS_PROFIT in g.columns:
            _agg_sup["粗利合計"] = (COL_GROSS_PROFIT, "sum")
        out = g.groupby(sup_col, dropna=False).agg(**_agg_sup).reset_index()
        for _gc in ("入庫数量", "出庫数量", "入庫金額税抜", "出庫金額税抜", "粗利合計"):
            if _gc in out.columns:
                out[_gc] = (
                    pd.to_numeric(out[_gc], errors="coerce")
                    .replace([np.inf, -np.inf], np.nan)
                    .fillna(0.0)
                )
        return out

    grp = _supplier_grp(flt)
    if COL_STOCK_STATUS not in flt.columns:
        st.caption(
            f"「{COL_STOCK_STATUS}」列がないため、在庫中／販売済に分けず期間内の全体を表示します。"
        )
        st.dataframe(
            grp.sort_values(sup_col),
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.caption(
            f"「{COL_STOCK_STATUS}」に従い、在庫中と販売済の行をそれぞれ集計した表です。"
            "下の「仕入先・取引先別 税抜金額」「粗利」グラフは従来どおり期間内の全体です。"
        )
        sn = flt[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
        flt_in = flt.loc[sn == STATUS_IN_STOCK]
        flt_so = flt.loc[sn == STATUS_SOLD]
        st.markdown("###### 在庫中")
        st.dataframe(
            _supplier_grp(flt_in).sort_values(sup_col),
            use_container_width=True,
            hide_index=True,
        )
        st.markdown("###### 販売済")
        st.dataframe(
            _supplier_grp(flt_so).sort_values(sup_col),
            use_container_width=True,
            hide_index=True,
        )

    st.markdown("##### 月次推移（数量）")
    st.caption("各月で入庫・出庫を並べた棒グラフ（積み上げではありません）。")
    monthly = (
        flt.dropna(subset=[COL_DATETIME])
        .groupby("_ym", as_index=False)
        .agg(
            入庫数量=("_qty_in", "sum"),
            出庫数量=("_qty_out", "sum"),
            入庫金額税抜=("_amt_ex_in", "sum"),
            出庫金額税抜=("_amt_ex_out", "sum"),
        )
        .sort_values("_ym")
    )
    if not monthly.empty:
        for _mc in (
            "入庫数量",
            "出庫数量",
            "入庫金額税抜",
            "出庫金額税抜",
        ):
            if _mc in monthly.columns:
                monthly[_mc] = (
                    pd.to_numeric(monthly[_mc], errors="coerce")
                    .replace([np.inf, -np.inf], np.nan)
                    .fillna(0.0)
                )
        month_order = monthly["_ym"].astype(str).tolist()
        mdf_qty = monthly.rename(
            columns={"_ym": "月", "入庫数量": "入庫", "出庫数量": "出庫"}
        )
        qty_long = pd.melt(
            mdf_qty,
            id_vars=["月"],
            value_vars=["入庫", "出庫"],
            var_name="区分",
            value_name="数量",
        )
        qty_long["数量"] = pd.to_numeric(
            qty_long["数量"], errors="coerce"
        ).fillna(0.0)
        chart_qty = (
            alt.Chart(qty_long)
            .mark_bar()
            .encode(
                x=alt.X("月:N", sort=month_order, axis=alt.Axis(title="月", labelAngle=-45)),
                y=alt.Y(
                    "数量:Q",
                    title="数量",
                    axis=alt.Axis(format=",.0f"),
                    scale=_altair_y_scale_positive(qty_long["数量"]),
                ),
                xOffset=alt.XOffset("区分:N"),
                color=alt.Color(
                    "区分:N",
                    scale=alt.Scale(domain=["入庫", "出庫"], range=["#1f77b4", "#ff7f0e"]),
                    legend=alt.Legend(title=None),
                ),
                tooltip=[
                    alt.Tooltip("月:N", title="月"),
                    "区分",
                    alt.Tooltip("数量:Q", title="数量", format=",.0f"),
                ],
            )
            .properties(height=340)
        )
        st.altair_chart(chart_qty, use_container_width=True)
    else:
        st.caption("月次グラフを表示できる日付がありません。")

    st.markdown("##### 月次推移（金額・税抜）")
    st.caption(
        "数量グラフと同じ月次集計で、入庫・出庫の税抜金額（いずれも **仕入金額ベース**）を並べた棒グラフです（積み上げではありません）。"
    )
    if not monthly.empty:
        month_order = monthly["_ym"].astype(str).tolist()
        mdf_amt = monthly.rename(
            columns={"_ym": "月", "入庫金額税抜": "入庫", "出庫金額税抜": "出庫"}
        )
        amt_bar_long = pd.melt(
            mdf_amt,
            id_vars=["月"],
            value_vars=["入庫", "出庫"],
            var_name="区分",
            value_name="金額",
        )
        amt_bar_long["金額"] = pd.to_numeric(
            amt_bar_long["金額"], errors="coerce"
        ).fillna(0.0)
        chart_amt_bar = (
            alt.Chart(amt_bar_long)
            .mark_bar()
            .encode(
                x=alt.X("月:N", sort=month_order, axis=alt.Axis(title="月", labelAngle=-45)),
                y=alt.Y(
                    "金額:Q",
                    title="金額（税抜・円）",
                    axis=alt.Axis(format=",.0f"),
                    scale=_altair_y_scale_positive(amt_bar_long["金額"]),
                ),
                xOffset=alt.XOffset("区分:N"),
                color=alt.Color(
                    "区分:N",
                    scale=alt.Scale(domain=["入庫", "出庫"], range=["#1f77b4", "#ff7f0e"]),
                    legend=alt.Legend(title=None),
                ),
                tooltip=[
                    alt.Tooltip("月:N", title="月"),
                    "区分",
                    alt.Tooltip("金額:Q", title="金額（円）", format=",.0f"),
                ],
            )
            .properties(height=340)
        )
        st.altair_chart(chart_amt_bar, use_container_width=True)
    else:
        st.caption("月次の金額グラフを表示できる日付がありません。")

    st.markdown("##### 金額推移（税抜）")
    st.caption(
        "その月までの入庫・出庫それぞれの税抜金額（**仕入金額ベース**）の**累計**（いわゆる累積）を、"
        "各月で入庫・出庫の2本の棒として並べたグラフです。"
        "上の From〜To・仕入先・取引先の絞り込みに従います。"
    )
    if not monthly.empty:
        month_order_c = monthly["_ym"].astype(str).tolist()
        mc = monthly.sort_values("_ym").copy()
        mc["入庫累積"] = mc["入庫金額税抜"].cumsum()
        mc["出庫累積"] = mc["出庫金額税抜"].cumsum()
        mdf_cum = mc.rename(columns={"_ym": "月", "入庫累積": "入庫", "出庫累積": "出庫"})
        cum_long = pd.melt(
            mdf_cum,
            id_vars=["月"],
            value_vars=["入庫", "出庫"],
            var_name="区分",
            value_name="累積金額",
        )
        cum_long["累積金額"] = pd.to_numeric(
            cum_long["累積金額"], errors="coerce"
        ).fillna(0.0)
        chart_cum_bar = (
            alt.Chart(cum_long)
            .mark_bar()
            .encode(
                x=alt.X("月:N", sort=month_order_c, axis=alt.Axis(title="月", labelAngle=-45)),
                y=alt.Y(
                    "累積金額:Q",
                    title="累積金額（税抜・円）",
                    axis=alt.Axis(format=",.0f"),
                    scale=_altair_y_scale_positive(cum_long["累積金額"]),
                ),
                xOffset=alt.XOffset("区分:N"),
                color=alt.Color(
                    "区分:N",
                    scale=alt.Scale(domain=["入庫", "出庫"], range=["#1f77b4", "#ff7f0e"]),
                    legend=alt.Legend(title=None),
                ),
                tooltip=[
                    alt.Tooltip("月:N", title="月"),
                    "区分",
                    alt.Tooltip("累積金額:Q", title="累積（円）", format=",.0f"),
                ],
            )
            .properties(height=340)
        )
        st.altair_chart(chart_cum_bar, use_container_width=True)
    else:
        st.caption("累積金額グラフを表示できる月次データがありません。")

    st.markdown("##### 仕入先・取引先別 税抜金額（変動幅の大きい順・上位15件）")
    st.caption("各仕入先・取引先で入庫・出庫の税抜金額（**仕入金額ベース**）を並べた棒グラフです。")
    chart_src = (
        grp.set_index(sup_col)[["入庫金額税抜", "出庫金額税抜"]]
        .assign(_abs=lambda x: (x["入庫金額税抜"] - x["出庫金額税抜"]).abs())
        .sort_values("_abs", ascending=False)
        .drop(columns=["_abs"])
        .head(15)
    )
    if not chart_src.empty:
        top_src = chart_src.reset_index()
        top_order = top_src[sup_col].astype(str).tolist()
        top_long = pd.melt(
            top_src.rename(columns={"入庫金額税抜": "入庫", "出庫金額税抜": "出庫"}),
            id_vars=[sup_col],
            value_vars=["入庫", "出庫"],
            var_name="区分",
            value_name="金額",
        )
        top_long["金額"] = pd.to_numeric(
            top_long["金額"], errors="coerce"
        ).fillna(0.0)
        sup_chart = (
            alt.Chart(top_long)
            .mark_bar()
            .encode(
                x=alt.X(
                    f"{sup_col}:N",
                    sort=top_order,
                    axis=alt.Axis(title=sup_col, labelAngle=-45),
                ),
                y=alt.Y(
                    "金額:Q",
                    title="金額（税抜・円）",
                    axis=alt.Axis(format=",.0f"),
                    scale=_altair_y_scale_positive(top_long["金額"]),
                ),
                xOffset=alt.XOffset("区分:N"),
                color=alt.Color(
                    "区分:N",
                    scale=alt.Scale(domain=["入庫", "出庫"], range=["#1f77b4", "#ff7f0e"]),
                    legend=alt.Legend(title=None),
                ),
                tooltip=[
                    alt.Tooltip(f"{sup_col}:N", title=sup_col),
                    "区分",
                    alt.Tooltip("金額:Q", title="金額（円）", format=",.0f"),
                ],
            )
            .properties(height=380)
        )
        st.altair_chart(sup_chart, use_container_width=True)

    if "粗利合計" in grp.columns:
        st.markdown("##### 仕入先・取引先別 粗利（税抜・上位15件）")
        st.caption(
            f"台帳の「{COL_GROSS_PROFIT}」列を仕入先・取引先ごとに合算しています。"
            "未販売は「在庫中」、販売済は「販売済」の行に限定します。"
            "各グラフの並びは粗利の絶対値が大きい順です（マイナスも含みます）。"
        )

        def _gp_supplier_top15_bar(sub: pd.DataFrame) -> None:
            if sub.empty:
                st.caption("該当する行がありません。")
                return
            gsub = sub.assign(**{sup_col: sub[COL_SUPPLIER].fillna("(未設定)").astype(str)})
            gg = (
                gsub.groupby(sup_col, dropna=False)
                .agg(粗利合計=(COL_GROSS_PROFIT, "sum"))
                .reset_index()
            )
            gg["粗利合計"] = (
                pd.to_numeric(gg["粗利合計"], errors="coerce")
                .replace([np.inf, -np.inf], np.nan)
                .fillna(0.0)
            )
            gp_ch = (
                gg.assign(_abs=lambda x: x["粗利合計"].abs())
                .sort_values("_abs", ascending=False)
                .drop(columns=["_abs"])
                .head(15)
            )
            if gp_ch.empty:
                st.caption("該当するデータがありません。")
                return
            g_order = gp_ch[sup_col].astype(str).tolist()
            gp_bar = (
                alt.Chart(gp_ch)
                .mark_bar()
                .encode(
                    x=alt.X(
                        f"{sup_col}:N",
                        sort=g_order,
                        axis=alt.Axis(title=sup_col, labelAngle=-45),
                    ),
                    y=alt.Y(
                        "粗利合計:Q",
                        title="粗利（税抜・円）",
                        axis=alt.Axis(format=",.0f"),
                    ),
                    tooltip=[
                        alt.Tooltip(f"{sup_col}:N", title=sup_col),
                        alt.Tooltip("粗利合計:Q", title="粗利（円）", format=",.0f"),
                    ],
                )
                .properties(height=380)
            )
            st.altair_chart(gp_bar, use_container_width=True)

        if COL_STOCK_STATUS not in flt.columns:
            st.caption(
                f"「{COL_STOCK_STATUS}」列がないため、未販売／販売済に分けず期間内の全体を表示します。"
            )
            st.markdown("###### 全体")
            _gp_supplier_top15_bar(flt)
        else:
            sn = flt[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
            flt_unsold = flt.loc[sn == STATUS_IN_STOCK]
            flt_sold = flt.loc[sn == STATUS_SOLD]
            st.markdown("###### 未販売（在庫中）")
            _gp_supplier_top15_bar(flt_unsold)
            st.markdown("###### 販売済")
            _gp_supplier_top15_bar(flt_sold)


def _render_inventory_price_summary(df: pd.DataFrame) -> None:
    """在庫中の行について、合計原価・販売予定（税抜行計・税込総額）・想定粗利を表示する。"""
    st.markdown("##### 価格管理サマリー（在庫中）")
    st.caption(
        "上の表のうち、ステータスが「在庫中」の行だけを合算しています（未保存の編集を含みます）。"
        f"「販売予定金額（税抜）」は1点あたりとして **{COL_QTY}** を掛けた行計の合計、税込列は再計算後の値の合計です。"
    )
    if df is None or df.empty:
        return
    if COL_STOCK_STATUS not in df.columns:
        return
    calc = _recalc_gross_profit_dataframe(df.copy())
    mask = calc[COL_STOCK_STATUS].astype(str).str.strip() == STATUS_IN_STOCK
    sub = calc.loc[mask]
    if sub.empty:
        st.info("「在庫中」の行がありません。")
        return
    cg = sub[COL_PRICE_EXCL].map(_int_from_cell)
    pl_u = sub[COL_PLANNED_SALE].map(_int_from_cell)
    if COL_QTY in sub.columns:
        rq = sub[COL_QTY].map(lambda x: max(1, _finite_int(x, 1)))
    else:
        rq = pd.Series(1, index=sub.index, dtype=int)
    pl_line_ex = pl_u * rq.astype(np.int64, copy=False)
    pl_in = sub[COL_PLANNED_SALE_INCL].map(_int_from_cell)
    total_cogs = int(cg.sum())
    total_planned_excl = int(pl_line_ex.sum())
    total_planned_incl = int(pl_in.sum())
    m = pl_line_ex > 0
    total_margin = int((pl_line_ex.loc[m] - cg.loc[m]).sum())
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("合計原価（税抜）", f"¥{total_cogs:,}")
    m2.metric("合計販売予定（税抜・行計）", f"¥{total_planned_excl:,}")
    m3.metric("合計販売予定（税込）", f"¥{total_planned_incl:,}")
    m4.metric("想定粗利（税抜・合計）", f"¥{total_margin:,}")


def _product_keyword_category(name: str) -> str:
    """ダッシュボード用の簡易カテゴリ（商品名のキーワードから推定。該当なしはその他）。"""
    s = str(name or "")
    if not s.strip():
        return "その他"
    # 帯本体より先に付属小物を判定（「帯」単独マッチの取りこぼし防止）
    if re.search(r"(帯締|帯揚|帯留|シャール|ストール|草履|下駄|バッグ|巾着)", s):
        return "小物・装飾"
    if re.search(r"(名古屋帯|袋帯|半幅帯|角帯|丸帯|兵児帯|博多帯|献上帯)", s):
        return "帯"
    if re.search(r"帯", s):
        return "帯"
    if re.search(r"(長襦袢|半襦袢|襦袢|肌着)", s):
        return "長襦袢・肌着類"
    if re.search(r"振袖", s):
        return "振袖"
    if re.search(r"訪問着", s):
        return "訪問着"
    if re.search(r"(留袖|黒留)", s):
        return "留袖"
    if re.search(r"(色無地|紬|小紋|絽|江戸褞袍|付け下げ)", s):
        return "着物（色無地・紬等）"
    if re.search(r"(男|紳士|羽織袴)", s):
        return "紳士・男物"
    return "その他"


def _inventory_category_cache_path() -> Path:
    return Path(__file__).resolve().parent / INVENTORY_CATEGORY_CACHE_FILENAME


def _inventory_category_cache_key(product_name: str, supplier: str = "") -> str:
    """商品名＋仕入先でキャッシュ参照用のキー（大小・前後空白は正規化）。"""
    n = (product_name or "").strip()
    if not n:
        return ""
    s = (supplier or "").strip()
    return f"{n.casefold()}\t{s.casefold()}" if s else n.casefold()


def _inventory_category_cache_load() -> dict[str, str]:
    p = _inventory_category_cache_path()
    if not p.is_file():
        return {}
    try:
        raw = json.loads(p.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, TypeError, ValueError):
        return {}
    if not isinstance(raw, dict):
        return {}
    out: dict[str, str] = {}
    for k, v in raw.items():
        ks = str(k).strip()
        if not ks or not isinstance(v, str):
            continue
        vv = v.strip()
        if vv:
            out[ks] = vv[:80]
    return out


def _inventory_category_cache_write(mapping: dict[str, str]) -> None:
    p = _inventory_category_cache_path()
    data = {k: mapping[k] for k in sorted(mapping.keys())}
    tmp = p.with_name(p.name + ".tmp")
    try:
        tmp.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        tmp.replace(p)
    except OSError:
        try:
            if tmp.is_file():
                tmp.unlink()
        except OSError:
            pass


def _inventory_category_cache_merge(updates: dict[str, str]) -> int:
    """キャッシュに反映したキー数（新規または値が変わったもののみカウント）。"""
    if not updates:
        return 0
    cur = _inventory_category_cache_load()
    changed = 0
    for k, v in updates.items():
        ks = str(k).strip()
        if not ks or not isinstance(v, str):
            continue
        vv = v.strip()[:80]
        if not vv:
            continue
        if cur.get(ks) != vv:
            changed += 1
        cur[ks] = vv
    _inventory_category_cache_write(cur)
    return changed


def _resolve_inventory_category_label(
    row: pd.Series,
    cache: dict[str, str],
) -> str:
    """構成比用ラベル: 台帳列 → ローカルキャッシュ → 和装向けキーワード。"""
    if COL_CATEGORY in row.index:
        cell = str(row.get(COL_CATEGORY, "") or "").strip()
        if cell:
            return cell[:80]
    name = str(row.get(COL_NAME, "") or "").strip()
    sup = str(row.get(COL_SUPPLIER, "") or "").strip()
    ck = _inventory_category_cache_key(name, sup)
    if ck and ck in cache:
        return str(cache[ck]).strip()[:80]
    return _product_keyword_category(name)


def _apply_inventory_category_map_to_dataframe(
    df: pd.DataFrame,
    mapping: dict[str, str],
    *,
    only_in_stock: bool,
    only_empty: bool,
    overwrite: bool,
) -> tuple[pd.DataFrame, int]:
    """mapping のキー（_inventory_category_cache_key）に一致する行の COL_CATEGORY を更新。更新行数を返す。"""
    if COL_CATEGORY not in df.columns or COL_NAME not in df.columns:
        return df, 0
    out = df.copy()
    if COL_STOCK_STATUS in out.columns:
        stt = out[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
        m_in = stt == STATUS_IN_STOCK
    else:
        m_in = pd.Series(True, index=out.index)
    n_up = 0
    for idx in out.index:
        if only_in_stock and not bool(m_in.loc[idx]):
            continue
        cur = str(out.loc[idx, COL_CATEGORY] or "").strip()
        if only_empty and cur:
            continue
        if not overwrite and cur:
            continue
        pn = str(out.loc[idx, COL_NAME] or "").strip()
        sp = (
            str(out.loc[idx, COL_SUPPLIER] or "").strip()
            if COL_SUPPLIER in out.columns
            else ""
        )
        ck = _inventory_category_cache_key(pn, sp)
        if ck in mapping and str(mapping[ck]).strip():
            out.loc[idx, COL_CATEGORY] = str(mapping[ck]).strip()[:80]
            n_up += 1
    return out, n_up


def infer_inventory_categories_with_gemini(
    pairs: list[tuple[str, str]],
) -> dict[str, str]:
    """(商品名, 仕入先) の組ごとに在庫カテゴリーラベルを推定し、キャッシュキー → ラベルを返す。"""
    if not pairs:
        return {}
    api_key = _secret_str(SECRET_GEMINI_API_KEY)
    if not api_key:
        raise RuntimeError(
            f"{SECRET_GEMINI_API_KEY} が設定されていません。`.streamlit/secrets.toml` を確認してください。"
        )
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(_gemini_model_name())
    lines: list[str] = []
    for i, (pn, sp) in enumerate(pairs, start=1):
        lines.append(
            f"{i}. product: {json.dumps(pn, ensure_ascii=False)}, "
            f"supplier: {json.dumps(sp, ensure_ascii=False)}"
        )
    block = "\n".join(lines)
    prompt = f"""あなたは小売・卸の在庫分析のための分類アシスタントです（業種は問いません。呉服以外も同様に扱う）。
次の各行は「商品名」「仕入先・取引先」の組です。円グラフ向けに、**短い在庫カテゴリー名**（全角含め20文字以内・同種商品は同じラベルに統一）を付けてください。

ルール:
- ラベルは可能な限り少ない種類にまとめ、在庫ポートフォリオの把握に役立つ粒度にする。
- 出力は **純粋な JSON オブジェクト1個だけ**（説明・Markdown・コードフェンス禁止）。
- キーは "items" のみ。値は配列で、各要素は "product"（入力と同一文字列）, "supplier"（入力と同一。無ければ ""）, "category"（ラベル文字列）。
- 与えられた全行について必ず1要素ずつ返す（欠落禁止）。迷ったら "その他"。

--- 対象行 ---
{block}
"""
    response = model.generate_content(prompt)
    text = (response.text or "").strip()
    data = _parse_json_from_model(text)
    items = data.get("items")
    if not isinstance(items, list):
        return {}
    out: dict[str, str] = {}
    for it in items:
        if not isinstance(it, dict):
            continue
        pn = str(it.get("product", "") or "").strip()
        if not pn:
            continue
        sp = str(it.get("supplier", "") or "").strip()
        cat = str(it.get("category", "") or "").strip()[:80]
        if not cat:
            cat = "その他"
        out[_inventory_category_cache_key(pn, sp)] = cat
    return out


def _render_inventory_category_pie_altair(
    pie_df: pd.DataFrame, *, chart_title: str | None = None
) -> None:
    """Plotly 未導入時のカテゴリ構成比（円グラフ相当）。"""
    ttl = chart_title or "原価シェア（在庫カテゴリー・Altair）"
    chart = (
        alt.Chart(pie_df)
        .mark_arc(innerRadius=55)
        .encode(
            theta=alt.Theta("金額税抜:Q", stack=True),
            color=alt.Color("カテゴリー:N"),
            tooltip=[
                alt.Tooltip("カテゴリー:N"),
                alt.Tooltip("金額税抜:Q", title="金額（税抜）", format=","),
            ],
        )
        .properties(
            title=ttl,
            height=400,
        )
    )
    st.altair_chart(chart, use_container_width=True)


def _render_inventory_category_pie(
    pie_df: pd.DataFrame, *, chart_title: str | None = None
) -> None:
    """plotly が入っていれば ``st.plotly_chart``、なければ Altair にフォールバック。"""
    if float(pie_df["金額税抜"].sum()) <= 0:
        st.caption("対象ステータスで原価が入っている行がありません。")
        return
    ttl = chart_title or "原価シェア（在庫カテゴリー）"
    try:
        import plotly.express as px
    except ImportError:
        _render_inventory_category_pie_altair(pie_df, chart_title=ttl)
        return
    fig = px.pie(
        pie_df,
        names="カテゴリー",
        values="金額税抜",
        hole=0.35,
        title=ttl,
    )
    fig.update_traces(textposition="inside", textinfo="percent+label")
    st.plotly_chart(fig, use_container_width=True)


def _analytics_table_multisort(
    df: pd.DataFrame, specs: list[tuple[str, bool]]
) -> pd.DataFrame:
    """分析の対象行一覧用。複数列を mergesort で安定ソート（日時・金額列は型変換して比較）。"""
    out = df.copy()
    keys: list[str] = []
    ascending: list[bool] = []
    drops: list[str] = []
    for i, (col, asc) in enumerate(specs):
        if not col or col not in out.columns:
            continue
        if col == COL_DATETIME:
            k = f"___an_sort_dt_{i}"
            out[k] = pd.to_datetime(out[col], errors="coerce")
            drops.append(k)
            keys.append(k)
            ascending.append(asc)
        elif col in (COL_PRICE_EXCL, COL_GROSS_PROFIT, COL_QTY):
            k = f"___an_sort_num_{i}"
            out[k] = pd.to_numeric(out[col], errors="coerce")
            drops.append(k)
            keys.append(k)
            ascending.append(asc)
        else:
            keys.append(col)
            ascending.append(asc)
    if not keys:
        return out
    out = out.sort_values(
        by=keys,
        ascending=ascending,
        na_position="last",
        kind="mergesort",
    )
    return out.drop(columns=drops, errors="ignore")


def _sold_actual_sale_line_totals(df: pd.DataFrame) -> tuple[int, int]:
    """販売済行について、実売の行計を税抜・税込で合算（単価0の行は除く）。"""
    if df is None or df.empty:
        return 0, 0
    need = (COL_STOCK_STATUS, COL_ACTUAL_SALE, COL_ACTUAL_SALE_INCL)
    if not all(c in df.columns for c in need):
        return 0, 0
    sn = df[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status)
    sold = df.loc[sn == STATUS_SOLD]
    if sold.empty:
        return 0, 0
    au = _series_to_numeric_loose(sold[COL_ACTUAL_SALE]).fillna(0)
    if COL_QTY in sold.columns:
        qv = sold[COL_QTY].map(lambda x: max(1, _finite_int(x, 1)))
    else:
        qv = pd.Series(1, index=sold.index, dtype=int)
    m = au > 0
    if not bool(m.any()):
        return 0, 0
    total_excl = int((au * qv.astype(np.int64)).loc[m].sum())
    inc = _series_to_numeric_loose(sold.loc[m, COL_ACTUAL_SALE_INCL]).fillna(0)
    total_incl = int(inc.sum())
    return total_excl, total_incl


def _purchase_age_days_series(df: pd.DataFrame) -> pd.Series:
    """仕入日時からの経過日数（JST 今日との差）。列が無い・日時が無効な行は NaN。"""
    if df is None or df.empty or COL_PURCHASE_DATETIME not in df.columns:
        return pd.Series(np.nan, index=df.index, dtype=np.float64)
    p = pd.to_datetime(df[COL_PURCHASE_DATETIME], errors="coerce")
    if getattr(p.dtype, "tz", None) is not None:
        p = p.dt.tz_convert(TZ_JP).dt.tz_localize(None)
    pur = p.dt.normalize()
    today = pd.Timestamp.combine(_today_jst_date(), datetime.min.time())
    td = today - pur
    return td.dt.days


def _mask_stagnant_inventory(df: pd.DataFrame, *, min_days: int = 90) -> pd.Series:
    """在庫中かつ仕入日時から min_days 日超が経過した行。"""
    if df is None or df.empty:
        return pd.Series(False, index=df.index)
    m_in = _mask_ledger_in_stock(df)
    age = _purchase_age_days_series(df)
    return m_in & age.notna() & (age > float(min_days))


def render_stagnant_inventory_section(calc: pd.DataFrame) -> None:
    """滞留在庫: 在庫中で仕入から90日超の行をカテゴリー別・仕入先別に可視化。"""
    st.subheader("滞留在庫")
    st.caption(
        f"**{STATUS_IN_STOCK}** かつ **{COL_PURCHASE_DATETIME}**（仕入が表示される日時）からの経過が "
        "**90日を超える** 行を対象にしています。日時が空・不正な行は含みません。"
    )
    if COL_PURCHASE_DATETIME not in calc.columns:
        st.warning(
            f"「{COL_PURCHASE_DATETIME}」列がないため、滞留在庫を集計できません。"
        )
        return
    if COL_STOCK_STATUS not in calc.columns:
        st.warning(
            f"「{COL_STOCK_STATUS}」列がないため、滞留在庫を集計できません。"
        )
        return
    m_st = _mask_stagnant_inventory(calc, min_days=90)
    stg = calc.loc[m_st].copy()
    n_st = int(len(stg))
    m1, m2 = st.columns(2)
    m1.metric("滞留在庫 件数", f"{n_st:,}")
    if stg.empty:
        m2.metric("滞留在庫 原価合計（税抜・行計）", "—")
        st.info("条件に合う行がありません（在庫の仕入日時を確認してください）。")
        return
    cg = _series_to_numeric_loose(stg[COL_PRICE_EXCL]).fillna(0)
    if COL_QTY in stg.columns:
        rq = stg[COL_QTY].map(lambda x: max(1, _finite_int(x, 1)))
        line_cost = cg * rq.astype(np.int64, copy=False)
    else:
        line_cost = cg
    total_cogs = int(line_cost.sum())
    m2.metric("滞留在庫 原価合計（税抜・行計）", f"¥{total_cogs:,}")

    _cat_cache = _inventory_category_cache_load()
    stg["_cat_label"] = stg.apply(
        lambda r: _resolve_inventory_category_label(r, _cat_cache), axis=1
    )
    stg["_age_days"] = _purchase_age_days_series(stg)
    stg["_line_cost"] = line_cost

    cat_agg = (
        stg.groupby("_cat_label", dropna=False)
        .agg(件数=("_line_cost", "count"), 原価税抜=("_line_cost", "sum"))
        .reset_index()
        .rename(columns={"_cat_label": "カテゴリー"})
    )
    cat_agg["原価税抜"] = pd.to_numeric(cat_agg["原価税抜"], errors="coerce").fillna(0)
    cat_agg = cat_agg.sort_values("原価税抜", ascending=False, kind="mergesort")

    st.markdown("##### 商品カテゴリー別（滞留在庫）")
    st.caption(
        "カテゴリーは分析と同じ優先（台帳の在庫カテゴリー → キャッシュ → キーワード → その他）。"
        "グラフ・一覧は **原価（仕入・税抜）の行計** で集計しています。"
    )
    if not cat_agg.empty and float(cat_agg["原価税抜"].sum()) > 0:
        cat_chart = (
            alt.Chart(cat_agg)
            .mark_bar()
            .encode(
                x=alt.X(
                    "原価税抜:Q",
                    title="原価合計（税抜・円）",
                    axis=alt.Axis(format=",.0f"),
                ),
                y=alt.Y("カテゴリー:N", sort="-x", title=None),
                tooltip=[
                    alt.Tooltip("カテゴリー:N"),
                    alt.Tooltip("件数:Q", title="件数", format=",.0f"),
                    alt.Tooltip("原価税抜:Q", title="原価（税抜）", format=",.0f"),
                ],
            )
            .properties(height=min(420, max(220, 28 * len(cat_agg))))
        )
        st.altair_chart(cat_chart, use_container_width=True)
    else:
        st.caption("グラフ用の原価が付いた行がありません。")
    _disp_cat = cat_agg.copy()
    _disp_cat["原価税抜"] = _disp_cat["原価税抜"].map(lambda x: f"¥{int(x):,}")
    st.dataframe(_disp_cat, use_container_width=True, hide_index=True)

    sup_col = COL_SUPPLIER
    st.markdown("##### 仕入先・取引先別（滞留在庫）")
    if sup_col in stg.columns:
        stg["_sup"] = stg[sup_col].fillna("(未設定)").astype(str)
    else:
        stg["_sup"] = pd.Series("(未設定)", index=stg.index, dtype=str)

    sup_agg = (
        stg.groupby("_sup", dropna=False)
        .agg(件数=("_line_cost", "count"), 原価税抜=("_line_cost", "sum"))
        .reset_index()
        .rename(columns={"_sup": sup_col})
    )
    sup_agg["原価税抜"] = pd.to_numeric(sup_agg["原価税抜"], errors="coerce").fillna(0)
    sup_agg = sup_agg.sort_values("原価税抜", ascending=False, kind="mergesort")

    if not sup_agg.empty and float(sup_agg["原価税抜"].sum()) > 0:
        sup_chart = (
            alt.Chart(sup_agg.head(25))
            .mark_bar()
            .encode(
                x=alt.X(
                    "原価税抜:Q",
                    title="原価合計（税抜・円）",
                    axis=alt.Axis(format=",.0f"),
                ),
                y=alt.Y(f"{sup_col}:N", sort="-x", title=None),
                tooltip=[
                    alt.Tooltip(f"{sup_col}:N", title=sup_col),
                    alt.Tooltip("件数:Q", title="件数", format=",.0f"),
                    alt.Tooltip("原価税抜:Q", title="原価（税抜）", format=",.0f"),
                ],
            )
            .properties(height=min(520, max(240, 22 * min(25, len(sup_agg)))))
        )
        st.altair_chart(sup_chart, use_container_width=True)
        if len(sup_agg) > 25:
            st.caption("グラフは原価合計の上位 **25** 件まで表示しています。一覧は全件です。")
    else:
        st.caption("グラフ用のデータがありません。")
    _disp_sup = sup_agg.copy()
    _disp_sup["原価税抜"] = _disp_sup["原価税抜"].map(lambda x: f"¥{int(x):,}")
    st.dataframe(_disp_sup, use_container_width=True, hide_index=True)

    _tbl_cols = [
        c
        for c in (
            COL_MANAGEMENT_ID,
            COL_NAME,
            COL_SUPPLIER,
            COL_CATEGORY,
            "_cat_label",
            COL_PURCHASE_DATETIME,
            "_age_days",
            COL_PRICE_EXCL,
            COL_QTY,
        )
        if c in stg.columns
    ]
    if _tbl_cols:
        with st.expander("滞留在庫の行一覧", expanded=False):
            _show = stg[_tbl_cols].copy()
            if "_cat_label" in _show.columns:
                _show.rename(columns={"_cat_label": "カテゴリー（解析用）"}, inplace=True)
            if "_age_days" in _show.columns:
                _show.rename(columns={"_age_days": "経過日数"}, inplace=True)
            _sort_cols = [
                c for c in ("経過日数", COL_PURCHASE_DATETIME) if c in _show.columns
            ]
            if _sort_cols:
                _show = _show.sort_values(
                    _sort_cols,
                    ascending=[False] * len(_sort_cols),
                    kind="mergesort",
                )
            st.dataframe(_show, use_container_width=True, hide_index=True)


def render_analytics_dashboard_page() -> None:
    """集計・分析: メトリクス・Plotly・既存の月次ダッシュボード。"""
    st.subheader("分析")
    st.caption(
        "共有の **inventory.csv**（`INVENTORY_SOURCE=csv`）または **Google スプレッドシート**から読み込んだ最新データを集計します。"
    )
    if not _uses_local_inventory_csv() and not _secret_str(SECRET_GOOGLE_SPREADSHEET_ID):
        st.info(
            "`GOOGLE_SPREADSHEET_ID` を設定するか、`INVENTORY_SOURCE = \"csv\"` で CSV を使ってください。"
        )
        return
    try:
        df_raw = load_inventory_dataframe()
    except Exception as e:
        st.error(f"読み込みに失敗しました: {e}")
        return
    if df_raw is None:
        st.warning("台帳を開けませんでした。サービスアカウントと共有設定を確認してください。")
        return
    if df_raw.empty:
        st.info("集計する行がありません。")
        return

    calc = _recalc_gross_profit_dataframe(df_raw.copy())
    if COL_STOCK_STATUS not in calc.columns:
        st.warning("ステータス列がないため在庫中の集計ができません。")
        render_ledger_dashboard(calc)
        return

    st.markdown("##### 分析の対象・構成比の並び")
    st.caption(
        "上段の **ステータス** で、このページのメトリクス・円グラフ・AI 推定の対象行を絞り込みます（在庫中だけ／販売済だけ／両方など）。"
        "下段の **構成比の並び** は円グラフの凡例・扇の順に反映されます。"
    )
    _ac1, _ac2 = st.columns(2)
    with _ac1:
        _sel_st = st.multiselect(
            "ステータス（対象に含める）",
            options=list(STOCK_STATUS_OPTIONS),
            default=[STATUS_IN_STOCK],
            key="analytics_status_include",
        )
    with _ac2:
        _pie_sort = st.selectbox(
            "構成比（円グラフ）のカテゴリー並び",
            options=["金額の多い順", "金額の少ない順", "カテゴリー名（昇順）"],
            index=0,
            key="analytics_pie_category_sort",
        )
    if not _sel_st:
        st.warning("ステータスを1つ以上選んでください（未選択のときは在庫中のみにします）。")
        _sel_st = [STATUS_IN_STOCK]
    _sel_norm = [_normalize_stock_status(str(s)) for s in _sel_st]
    _status_lbl = "・".join(_sel_norm)
    _mask_status = (
        calc[COL_STOCK_STATUS].astype(str).map(_normalize_stock_status).isin(_sel_norm)
    )
    sub = calc.loc[_mask_status].copy()
    cg = _series_to_numeric_loose(sub[COL_PRICE_EXCL]).fillna(0).clip(lower=0)
    total_inv = int(cg.sum())
    n_lines = int(len(sub))
    gp_s = _series_to_numeric_loose(sub[COL_GROSS_PROFIT]).fillna(0)
    gp_total = int(gp_s.sum())

    total_sale_excl, total_sale_incl = _sold_actual_sale_line_totals(calc)

    k1, k2, k3 = st.columns(3)
    k1.metric(f"対象（{_status_lbl}）総額（仕入・税抜）", f"¥{total_inv:,}")
    k2.metric(f"対象（{_status_lbl}）行数", f"{n_lines:,}")
    k3.metric(
        f"対象（{_status_lbl}）粗利合計（税抜）",
        f"¥{gp_total:,}",
    )

    st.caption(
        "販売済の実売は、次の2つは **ステータスが販売済の行（台帳全体）** の合計です（上のステータス絞り込みとは独立）。"
        "単価0の販売済行は含みません。"
    )
    s1, s2 = st.columns(2)
    s1.metric(
        "販売済・実売金額合計（税抜・行計）",
        f"¥{total_sale_excl:,}",
    )
    s2.metric(
        "販売済・実売金額合計（税込・行計）",
        f"¥{total_sale_incl:,}",
    )

    st.markdown("##### カテゴリー別 在庫原価（税抜）の構成比")
    st.caption(
        f"対象ステータス: **{_status_lbl}**。優先順: ①台帳の **{COL_CATEGORY}** 列 ②"
        f"**{INVENTORY_CATEGORY_CACHE_FILENAME}**（在庫一覧の AI 一括などで更新） ③和装向けキーワード ④**その他**。"
        "金額は仕入金額（税抜）を行で合算（数量列があるときは原価×数量）。"
    )
    _cat_cache = _inventory_category_cache_load()
    sub["_category"] = sub.apply(
        lambda r: _resolve_inventory_category_label(r, _cat_cache), axis=1
    )
    sub["_px"] = _series_to_numeric_loose(sub[COL_PRICE_EXCL]).fillna(0)
    if COL_QTY in sub.columns:
        _rq_pie = sub[COL_QTY].map(lambda x: max(1, _finite_int(x, 1)))
        sub["_px"] = sub["_px"] * _rq_pie
    pie_df = sub.groupby("_category", dropna=False)["_px"].sum().reset_index()
    pie_df.columns = ["カテゴリー", "金額税抜"]
    if _pie_sort == "金額の多い順":
        pie_df = pie_df.sort_values("金額税抜", ascending=False, kind="mergesort")
    elif _pie_sort == "金額の少ない順":
        pie_df = pie_df.sort_values("金額税抜", ascending=True, kind="mergesort")
    else:
        pie_df = pie_df.sort_values("カテゴリー", ascending=True, kind="mergesort")
    _pie_chart_title = f"原価シェア（在庫カテゴリー）— {_status_lbl}"
    _render_inventory_category_pie(pie_df, chart_title=_pie_chart_title)

    with st.expander("分析: 対象行一覧（ソート）", expanded=False):
        st.caption(
            "クリックで開閉します。**第1ソート** にステータス、**第2ソート** に仕入先を指定すると、"
            "ステータス別にまとめたうえで仕入先順に並べ替えられます（第2は「なし」で1段だけにもできます）。"
        )
        if sub.empty:
            st.caption("選んだステータスに該当する行がありません。")
        else:
            _tbl_opts = [
                COL_DATETIME,
                COL_NAME,
                COL_SUPPLIER,
                COL_MANAGEMENT_ID,
                COL_STOCK_STATUS,
                COL_QTY,
                COL_PRICE_EXCL,
                COL_GROSS_PROFIT,
                COL_CATEGORY,
            ]
            _tbl_sort_choices = [c for c in _tbl_opts if c in sub.columns]
            if not _tbl_sort_choices:
                _tbl_sort_choices = list(sub.columns)[: min(8, len(sub.columns))]
            _sec_opts = ["なし"] + _tbl_sort_choices
            _def_p = (
                COL_STOCK_STATUS
                if COL_STOCK_STATUS in _tbl_sort_choices
                else _tbl_sort_choices[0]
            )
            _def_s = (
                COL_SUPPLIER
                if COL_SUPPLIER in _tbl_sort_choices
                else "なし"
            )
            _ix_p = (
                _tbl_sort_choices.index(_def_p)
                if _def_p in _tbl_sort_choices
                else 0
            )
            _ix_s = _sec_opts.index(_def_s) if _def_s in _sec_opts else 0
            _a1, _a2, _a3, _a4 = st.columns([2, 1, 2, 1])
            with _a1:
                _p_col = st.selectbox(
                    "第1ソート",
                    options=_tbl_sort_choices,
                    index=_ix_p,
                    key="analytics_tbl_sort_p",
                )
            with _a2:
                _p_ord = st.radio(
                    "第1の順序",
                    ["昇順", "降順"],
                    horizontal=True,
                    key="analytics_tbl_sort_p_ord",
                )
            with _a3:
                _s_col = st.selectbox(
                    "第2ソート",
                    options=_sec_opts,
                    index=_ix_s,
                    key="analytics_tbl_sort_s",
                )
            with _a4:
                _s_ord = st.radio(
                    "第2の順序",
                    ["昇順", "降順"],
                    horizontal=True,
                    key="analytics_tbl_sort_s_ord",
                    disabled=(_s_col == "なし"),
                )
            _cols_show = [c for c in _tbl_opts if c in sub.columns]
            if _cols_show:
                _view = sub[_cols_show].copy()
                _specs: list[tuple[str, bool]] = [
                    (_p_col, _p_ord == "昇順"),
                ]
                if _s_col != "なし" and _s_col in _view.columns:
                    if _s_col != _p_col:
                        _specs.append((_s_col, _s_ord == "昇順"))
                _view = _analytics_table_multisort(_view, _specs)
                st.dataframe(_view, use_container_width=True, hide_index=True)
            else:
                st.caption("表示できる列がありません。")

    st.divider()
    render_stagnant_inventory_section(calc)

    st.divider()
    render_ledger_dashboard(calc)


def _inject_prominent_main_tabs_style() -> None:
    """メインエリアの ``st.tabs`` ラベルを大きく太字にする（登録画面のタブ用）。"""
    st.markdown(
        """
        <style>
        section.main [data-testid="stTabs"] button {
            font-size: clamp(1.1rem, 2.2vw, 1.45rem) !important;
            font-weight: 700 !important;
            line-height: 1.35 !important;
            padding-top: 0.65rem !important;
            padding-bottom: 0.65rem !important;
            letter-spacing: 0.02em !important;
        }
        section.main [data-testid="stTabs"] [role="tablist"] {
            gap: 0.35rem !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _google_drive_file_id_from_url(url: str) -> str | None:
    u = (url or "").strip()
    if not u:
        return None
    m = re.search(r"/file/d/([a-zA-Z0-9_-]{10,})", u, re.I)
    if m:
        return m.group(1)
    m2 = re.search(r"[?&]id=([a-zA-Z0-9_-]{10,})", u, re.I)
    if m2:
        return m2.group(1)
    m3 = re.search(r"googleusercontent\.com/d/([a-zA-Z0-9_-]{10,})", u, re.I)
    if m3:
        return m3.group(1)
    return None


def _render_http_product_image_from_url(
    image_url: str,
    *,
    pixel_width: int | None = None,
    use_container_width: bool = False,
) -> None:
    """HTTP(S) 画像を取得して表示。Google Drive 直リンクは ``st.image(URL)`` が効かないことが多いため中継取得する。

    ``pixel_width`` … タイル等の固定幅（ピクセル）。``None`` のときは長辺をおおむね 2000px までに抑えて **コンテナ幅** で表示（詳細ダイアログ用）。
    """
    iu = (image_url or "").strip()
    if not (iu.startswith("http://") or iu.startswith("https://")):
        st.caption("（画像なし）")
        return
    if pixel_width is not None:
        _w = max(120, min(480, int(pixel_width)))
        sz = min(1200, _w * 4)
        thumb_max = (_w * 2, _w * 2)
    else:
        _w = None
        sz = 2000
        thumb_max = (2000, 2000)
    fid = _google_drive_file_id_from_url(iu)
    candidates: list[str] = []
    if fid:
        candidates.append(f"https://drive.google.com/thumbnail?id={fid}&sz=w{sz}")
        candidates.append(
            f"https://drive.usercontent.google.com/download?id={fid}&export=view"
        )
    candidates.append(iu)

    ua = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
        )
    }
    timeout = 20.0
    for cand in candidates:
        try:
            r = requests.get(cand, timeout=timeout, headers=ua, allow_redirects=True)
        except Exception:
            continue
        if r.status_code != 200 or not r.content or len(r.content) < 32:
            continue
        low512 = r.content[:512].lower()
        if b"<html" in low512 or b"<!doctype" in low512:
            continue
        try:
            im = Image.open(io.BytesIO(r.content))
            im = ImageOps.exif_transpose(im)
            im.thumbnail(thumb_max, Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            im.convert("RGB").save(buf, format="JPEG", quality=85, optimize=True)
            out_b = buf.getvalue()
        except Exception:
            out_b = r.content
        try:
            if use_container_width:
                st.image(out_b, use_container_width=True)
            else:
                st.image(out_b, width=_w, use_container_width=False)
            return
        except Exception:
            continue
    st.caption("（ブラウザで開くと表示できる画像です）")
    st.link_button("画像を開く", iu, use_container_width=True)


def _render_inventory_gallery_thumbnail(image_url: str, *, width: int, sold: bool) -> None:
    """ギャラリー用サムネイル。``sold`` は呼び出し側の表示スタイル用（互換のため残す）。"""
    _ = sold
    _render_http_product_image_from_url(
        image_url, pixel_width=width, use_container_width=False
    )


def _sale_card_hit_from_series(
    row: pd.Series,
    *,
    confidence: float | None = None,
    extra_caption: str | None = None,
) -> dict[str, Any]:
    """販売候補カード1件分の表示用 dict。"""
    return {
        "management_id": str(row.get(COL_MANAGEMENT_ID, "") or "").strip(),
        "product_name": str(row.get(COL_NAME, "") or "").strip(),
        "supplier": str(row.get(COL_SUPPLIER, "") or "").strip(),
        "inventory_category": str(row.get(COL_CATEGORY, "") or "").strip(),
        "line_price_excl": _finite_int(row.get(COL_PRICE_EXCL), 0),
        "planned_sale_excl": _finite_int(row.get(COL_PLANNED_SALE), 0),
        "memo": str(row.get(COL_MEMO, "") or "").strip(),
        "image_url": str(row.get(COL_IMAGE_URL, "") or "").strip(),
        "confidence": confidence,
        "last_stocktake": str(row.get(COL_LAST_STOCKTAKE, "") or "").strip()
        if COL_LAST_STOCKTAKE in row.index
        else "",
        "extra_caption": extra_caption or "",
    }


def _render_mid_pick_candidate_cards(
    hits: list[dict[str, Any]],
    *,
    widget_key_namespace: str,
    sold: bool = False,
    pick_mode: str = "sale",
    page_size: int = 5,
    pager: str = "internal",
    caption_override: str | None = None,
    show_action_button: bool = True,
    sale_merge_selection: bool = False,
) -> None:
    """管理IDを選ぶ候補をカード表示（5件ページング付き）。

    ``pager`` が ``"hidden"`` のときは ``hits`` をそのまま1ページとして表示し、
    内蔵の前後ページボタンは出しません（外側でページングするとき用）。
    """
    valid_hits = [
        h for h in (hits or []) if str(h.get("management_id") or "").strip()
    ]
    if not valid_hits:
        return
    psize = max(1, int(page_size))
    page_key = f"_{widget_key_namespace}_page"
    total = len(valid_hits)
    if pager == "hidden":
        page_hits = valid_hits
        start_i = 0
        if caption_override:
            st.caption(caption_override)
        else:
            st.caption(f"候補 **{total}** 件")
    else:
        n_pages = max(1, (total + psize - 1) // psize)
        cur = int(st.session_state.get(page_key, 0) or 0)
        cur = max(0, min(n_pages - 1, cur))
        st.session_state[page_key] = cur
        start_i = cur * psize
        end_i = min(total, start_i + psize)
        page_hits = valid_hits[start_i:end_i]
        if caption_override:
            st.caption(caption_override)
        else:
            st.caption(
                f"候補 **{total}** 件（1ページ最大 **{psize}** 件） / "
                f"ページ **{cur + 1} / {n_pages}**"
            )
        if n_pages > 1:
            p1, p2, p3 = st.columns([1, 2, 1])
            with p1:
                if st.button("◀ 前へ", disabled=cur <= 0, key=f"{widget_key_namespace}_prev"):
                    st.session_state[page_key] = max(0, cur - 1)
                    st.rerun()
            with p2:
                st.caption(f"{start_i + 1}〜{end_i} 件を表示")
            with p3:
                if st.button("次へ ▶", disabled=cur >= n_pages - 1, key=f"{widget_key_namespace}_next"):
                    st.session_state[page_key] = min(n_pages - 1, cur + 1)
                    st.rerun()

    if pick_mode in ("stocktake", "stocktake_merge"):
        btn_label = "この候補を選ぶ"
    elif pick_mode == "purchase":
        btn_label = "この候補を仕入入力へ反映"
    else:
        btn_label = "この候補を販売元にする"

    for j, hit in enumerate(page_hits):
        mid = str(hit.get("management_id") or "").strip()
        abs_idx = start_i + j
        with st.container(border=True):
            h1, h2 = st.columns([1, 2])
            with h1:
                _render_inventory_gallery_thumbnail(
                    str(hit.get("image_url") or ""),
                    width=200,
                    sold=sold,
                )
            with h2:
                st.markdown(f"**管理ID:** `{mid}`")
                st.write(f"**商品名:** {hit.get('product_name') or '—'}")
                st.write(f"**仕入先:** {hit.get('supplier') or '—'}")
                lst = str(hit.get("last_stocktake") or "").strip()
                if lst:
                    st.write(f"**前回の棚卸日:** {lst}")
                memo_text = str(hit.get("memo") or "").strip()
                if memo_text:
                    st.markdown("**メモ:**")
                    for ln in [x.strip(" ・-\t") for x in memo_text.splitlines()]:
                        if ln:
                            st.write(f"- {ln}")
                xc = str(hit.get("extra_caption") or "").strip()
                if xc:
                    st.caption(xc)
                cf = hit.get("confidence")
                try:
                    cfn = float(cf) if cf is not None else None
                except (TypeError, ValueError):
                    cfn = None
                if cfn is not None and math.isfinite(cfn):
                    st.caption(f"AI 確信度（参考）: {cfn:.2f}")
                if show_action_button and st.button(
                    btn_label,
                    key=f"{widget_key_namespace}_mid_card_{abs_idx}_{mid}",
                    type="secondary",
                ):
                    if pick_mode == "stocktake":
                        st.session_state["_stocktake_selected_mid"] = mid
                    elif pick_mode == "stocktake_merge":
                        _cur_m = {
                            str(x).strip()
                            for x in (st.session_state.get("stocktake_assist_batch_mids") or [])
                            if str(x).strip()
                        }
                        _cur_m.add(mid)
                        st.session_state[_PENDING_STOCKTAKE_ASSIST_BATCH_MIDS] = sorted(
                            _cur_m, key=_management_id_sort_key
                        )
                    elif pick_mode == "purchase":
                        st.session_state.field_product_name = str(
                            hit.get("product_name") or ""
                        ).strip()
                        st.session_state.field_supplier = str(
                            hit.get("supplier") or ""
                        ).strip()
                        cat = str(hit.get("inventory_category") or "").strip()
                        if cat:
                            st.session_state.field_inventory_category = cat
                        lp = _finite_int(hit.get("line_price_excl"), 0)
                        if lp > 0:
                            st.session_state.field_line_excl_yen = lp
                        ps = _finite_int(hit.get("planned_sale_excl"), 0)
                        if ps > 0:
                            st.session_state.field_planned_sale_excl = ps
                        st.session_state["_gemini_match_management_id"] = mid
                    elif sale_merge_selection:
                        _cur_s = set(
                            _split_management_ids_from_field(
                                str(
                                    st.session_state.get("field_sale_source_mgmt_id", "")
                                    or ""
                                )
                            )
                        )
                        _cur_s.add(mid)
                        st.session_state.field_sale_source_mgmt_id = ", ".join(
                            sorted(_cur_s, key=_management_id_sort_key)
                        )
                    else:
                        st.session_state.field_sale_source_mgmt_id = mid
                    st.rerun()


def _sales_photo_match_card_hits_from_result(
    result: dict[str, Any],
    df_ledger: pd.DataFrame | None,
    *,
    sold_rows_only: bool = False,
    float_loan_outbound_only: bool = False,
) -> list[dict[str, Any]]:
    """写真照合 JSON と台帳から、カード候補（全件）を組み立てる。

    ``sold_rows_only=True`` … **販売済** の行のみ（出庫（返品）の写真照合用）。
    ``float_loan_outbound_only=True`` … **在庫中かつ出庫種別が出庫（浮貸）** のみ（出庫（戻入）の写真照合用）。
    """
    if df_ledger is None or df_ledger.empty:
        return []
    if float_loan_outbound_only:
        sub = df_ledger.loc[_mask_ledger_in_stock_outbound_float_loan(df_ledger)]
    else:
        sub = (
            df_ledger.loc[_mask_ledger_sold(df_ledger)]
            if sold_rows_only
            else df_ledger.loc[_mask_ledger_in_stock(df_ledger)]
        )
    if sub.empty:
        return []
    m = result.get("match")
    if not isinstance(m, dict):
        m = {}
    pn0 = str(
        m.get("product_name")
        or result.get("product_name")
        or result.get("商品名")
        or ""
    ).strip()
    su0 = str(
        m.get("supplier")
        or result.get("supplier")
        or result.get("仕入先・取引先")
        or result.get("仕入先")
        or ""
    ).strip()
    conf = float(m.get("confidence") or result.get("confidence") or 0)
    mid_primary = str(m.get("management_id") or "").strip()
    cand = _fuzzy_ledger_match_rows(sub, pn0, su0, limit=None)
    if cand.empty:
        return []
    hits: list[dict[str, Any]] = []
    seen: set[str] = set()
    for _, row in cand.iterrows():
        mid = str(row.get(COL_MANAGEMENT_ID, "") or "").strip()
        if not mid or mid in seen:
            continue
        seen.add(mid)
        c = conf if mid_primary and mid == mid_primary else None
        xc = (
            "写真照合で推定した行（同一管理ID）"
            if mid_primary and mid == mid_primary
            else ""
        )
        hits.append(
            _sale_card_hit_from_series(
                row,
                confidence=c,
                extra_caption=xc or None,
            )
        )

    def _sort_key(h: dict[str, Any]) -> tuple[int, Any]:
        mid_h = str(h.get("management_id") or "")
        return (
            0 if mid_primary and mid_h == mid_primary else 1,
            _management_id_sort_key(mid_h),
        )

    hits.sort(key=_sort_key)
    return hits


def _render_inventory_ledger_data_editor_section(df_sorted: pd.DataFrame) -> None:
    """在庫一覧の表形式: data_editor と保存ボタン。"""
    _ledger_base_for_save = df_sorted
    df_sorted = df_sorted.copy()
    _ledger_evidence_link_mode = _normalize_evidence_urls_for_link_editor(
        df_sorted, COL_VOUCHER_EVIDENCE_URL
    )

    _ledger_col_cfg: dict[str, Any] = {}
    if COL_MANAGEMENT_ID in df_sorted.columns:
        _ledger_col_cfg[COL_MANAGEMENT_ID] = st.column_config.TextColumn(
            COL_MANAGEMENT_ID,
            disabled=True,
            help="1点1行の自動採番（シリアル）。通常は手入力しません。",
        )
    if COL_QTY in df_sorted.columns:
        _ledger_col_cfg[COL_QTY] = st.column_config.NumberColumn(
            COL_QTY,
            min_value=1,
            step=1,
            format="%d",
            help="入出庫集計に使う点数（1以上の整数）。",
        )
    if COL_CATEGORY in df_sorted.columns:
        _ledger_col_cfg[COL_CATEGORY] = st.column_config.TextColumn(
            COL_CATEGORY,
            help="構成比・分析用。空欄のときはキャッシュまたはキーワードで補完されます。下の **AI で一括推定** で埋められます。",
        )
    if COL_STOCK_STATUS in df_sorted.columns:
        _ledger_col_cfg[COL_STOCK_STATUS] = st.column_config.SelectboxColumn(
            COL_STOCK_STATUS,
            options=list(STOCK_STATUS_OPTIONS),
            help="在庫中＝未販売想定、販売済＝実売で粗利、対象外＝在庫・販売の対象外（粗利は計上しません）。",
        )
    if COL_LAST_STOCKTAKE in df_sorted.columns:
        _ledger_col_cfg[COL_LAST_STOCKTAKE] = st.column_config.TextColumn(
            COL_LAST_STOCKTAKE,
            help=f"棚卸・実地確認した日（日本時間の暦日推奨）。例: {_today_jst_date().isoformat()}",
        )
    if COL_VOUCHER_RECORDED_AT in df_sorted.columns:
        _ledger_col_cfg[COL_VOUCHER_RECORDED_AT] = st.column_config.TextColumn(
            COL_VOUCHER_RECORDED_AT,
            help="証憑を台帳に反映した日時（JST・recorded_at に相当）。",
        )
    if COL_VOUCHER_EVIDENCE_URL in df_sorted.columns:
        if _ledger_evidence_link_mode:
            _ledger_col_cfg[COL_VOUCHER_EVIDENCE_URL] = st.column_config.LinkColumn(
                COL_VOUCHER_EVIDENCE_URL,
                help="Google ドライブの証憑 URL。アイコンをクリックするとブラウザで開きます。",
                disabled=True,
                display_text=":material/open_in_new:",
            )
        else:
            _ledger_col_cfg[COL_VOUCHER_EVIDENCE_URL] = st.column_config.TextColumn(
                COL_VOUCHER_EVIDENCE_URL,
                help="http(s) 以外の文字が混ざる行があるためテキスト表示です。URL をコピーしてブラウザで開いてください。",
            )

    _editor_kw: dict[str, Any] = {
        "num_rows": "dynamic",
        "key": LEDGER_DATA_EDITOR_KEY,
        "use_container_width": True,
        "hide_index": True,
        "height": 720,
    }
    if _ledger_col_cfg:
        _editor_kw["column_config"] = _ledger_col_cfg
    edited = st.data_editor(df_sorted, **_editor_kw)

    with st.expander("在庫カテゴリーを AI で一括推定（表に反映・台帳へ保存）", expanded=False):
        st.caption(
            "この表の **現在の内容**（未保存の編集を含む）に対し、Gemini で商品名＋仕入先から **在庫カテゴリー** を推定してセルに入れます。"
            "ローカルの **inventory_category_cache.json** も更新するため、集計・分析の構成比にも反映されます。"
            "セルが1件以上更新されたときは、続けて **Google スプレッドシート**（または **inventory.csv**）へ **自動保存** します。"
            "保存に失敗した場合のみ、反映済みの表を残したうえで「台帳を更新する」から再保存できます。"
        )
        _lc_only_in = st.checkbox("在庫中の行のみ", value=True, key="ledger_cat_bulk_only_in")
        _lc_only_empty = st.checkbox(
            "在庫カテゴリーが空の行のみ（推奨）",
            value=True,
            key="ledger_cat_bulk_only_empty",
        )
        _lc_overwrite = st.checkbox(
            "既にカテゴリーが入っている行も上書き",
            value=False,
            key="ledger_cat_bulk_overwrite",
        )
        _lc_max = int(
            st.number_input(
                "1回の最大件数（ユニークの商品名＋仕入先）",
                min_value=5,
                max_value=120,
                value=50,
                step=5,
                key="ledger_cat_bulk_max_n",
            )
        )
        if st.button(
            "AI で一括推定して表に反映",
            type="secondary",
            key="ledger_cat_bulk_run_btn",
        ):
            ed = edited.copy()
            pool: list[tuple[str, str]] = []
            seen_k: set[str] = set()
            for idx in ed.index:
                if _lc_only_in:
                    if COL_STOCK_STATUS not in ed.columns:
                        continue
                    stt = _normalize_stock_status(
                        str(ed.loc[idx, COL_STOCK_STATUS] or "")
                    )
                    if stt != STATUS_IN_STOCK:
                        continue
                if COL_CATEGORY in ed.columns:
                    cur = str(ed.loc[idx, COL_CATEGORY] or "").strip()
                else:
                    cur = ""
                if _lc_only_empty and cur:
                    continue
                if not _lc_overwrite and cur:
                    continue
                pn = str(ed.loc[idx, COL_NAME] or "").strip() if COL_NAME in ed.columns else ""
                sp = (
                    str(ed.loc[idx, COL_SUPPLIER] or "").strip()
                    if COL_SUPPLIER in ed.columns
                    else ""
                )
                if not pn:
                    continue
                ck = _inventory_category_cache_key(pn, sp)
                if not ck or ck in seen_k:
                    continue
                seen_k.add(ck)
                pool.append((pn, sp))
            targets = pool[:_lc_max]
            if not targets:
                st.info("条件に合う行がありません。")
            else:
                try:
                    with st.spinner(f"Gemini で {len(targets)} 件を推定しています…"):
                        upd = infer_inventory_categories_with_gemini(targets)
                    nmerge = _inventory_category_cache_merge(upd)
                    new_ed, ncell = _apply_inventory_category_map_to_dataframe(
                        ed,
                        upd,
                        only_in_stock=_lc_only_in,
                        only_empty=_lc_only_empty,
                        overwrite=_lc_overwrite,
                    )
                    _dest_lbl = (
                        "inventory.csv"
                        if _uses_local_inventory_csv()
                        else "Google スプレッドシート"
                    )
                    if ncell > 0:
                        try:
                            with st.spinner(f"{_dest_lbl} に保存しています…"):
                                overwrite_inventory_worksheet_from_dataframe(
                                    new_ed.reset_index(drop=True),
                                    previous_df=edited.reset_index(drop=True),
                                )
                        except Exception as save_e:
                            st.session_state[LEDGER_DATA_EDITOR_KEY] = new_ed
                            st.session_state["_ledger_saved_flash"] = (
                                f"在庫カテゴリーを AI で **{ncell}** 行は表に反映しましたが、"
                                f"{_dest_lbl} への保存に失敗しました: {save_e} "
                                "下の「台帳を更新する」で再保存してください。"
                            )
                        else:
                            st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                            st.session_state["_ledger_saved_flash"] = (
                                f"在庫カテゴリーを AI で **{ncell}** 行更新し、**{_dest_lbl}** に保存しました"
                                f"（API {len(upd)} キー・キャッシュの新規/変更 {nmerge}）。"
                            )
                    else:
                        st.session_state[LEDGER_DATA_EDITOR_KEY] = new_ed
                        st.session_state["_ledger_saved_flash"] = (
                            f"API は {len(upd)} 件のキーを返しましたが、選択した条件では表のセルは更新されませんでした"
                            f"（キャッシュの新規/変更 {nmerge}）。"
                        )
                except Exception as e:
                    st.error(str(e))
                st.rerun()

    if st.button(
        "在庫中かつ棚卸日が未入力の行に、今日の日付（JST）を一括入力してすぐ保存",
        key="ledger_bulk_stocktake_today_save",
    ):
        ed_bulk = edited.copy()
        if COL_LAST_STOCKTAKE not in ed_bulk.columns:
            ed_bulk[COL_LAST_STOCKTAKE] = ""
        m_b = _mask_ledger_stocktake_unverified(ed_bulk)
        ed_bulk.loc[m_b, COL_LAST_STOCKTAKE] = _today_jst_date().isoformat()
        with st.spinner("台帳を保存しています…"):
            try:
                overwrite_inventory_worksheet_from_dataframe(
                    ed_bulk.reset_index(drop=True),
                    previous_df=edited.reset_index(drop=True),
                )
            except Exception as e:
                st.error(str(e))
            else:
                filled_mids = set(
                    edited.loc[m_b, COL_MANAGEMENT_ID]
                    .astype(str)
                    .str.strip()
                    .tolist()
                )
                filled_mids.discard("")
                _inv_stocktake_work_remaining_note_done(filled_mids)
                st.session_state["_ledger_saved_flash"] = (
                    "棚卸日（今日・JST）を未確認の在庫中に一括入力し、台帳を保存しました。"
                )
                st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                st.rerun()

    _render_inventory_price_summary(edited)

    if st.button("台帳を更新する", type="primary", key="ledger_save_overwrite"):
        with st.spinner("台帳を保存しています…"):
            try:
                overwrite_inventory_worksheet_from_dataframe(
                    edited.reset_index(drop=True),
                    previous_df=_ledger_base_for_save.reset_index(drop=True),
                )
            except Exception as e:
                st.error(str(e))
                return
        _inv_stocktake_work_remaining_note_done(
            _management_ids_last_stocktake_changed(
                _ledger_base_for_save,
                edited,
            )
        )
        st.session_state["_ledger_saved_flash"] = "台帳を更新しました。"
        st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
        st.rerun()


def _render_inventory_gallery_body(df_sorted_calc: pd.DataFrame) -> None:
    """並び替え済み台帳からギャラリー（タイル）だけを描画する。"""
    st.markdown("### ギャラリー（カタログ）")
    st.caption(
        "Google ドライブの画像 URL はサーバー側で取得して表示します（表示できない場合は **画像を開く** からブラウザで確認できます）。"
    )
    g1, g2, g3 = st.columns([2, 2, 1])
    with g1:
        st.text_input(
            "フリーワード（商品名・管理ID・メモ）",
            key="inv_gallery_search_text",
            placeholder="部分一致で検索",
        )
    with g2:
        sup_opts: list[str] = []
        if COL_SUPPLIER in df_sorted_calc.columns:
            sup_opts = sorted(
                {
                    x
                    for x in df_sorted_calc[COL_SUPPLIER]
                    .astype(str)
                    .str.strip()
                    .tolist()
                    if x
                }
            )
        st.multiselect(
            "仕入先で絞り込み（複数可）",
            options=sup_opts,
            key="inv_gallery_suppliers_filter",
        )
    with g3:
        st.selectbox(
            "ステータス",
            ("すべて", "在庫中", "販売済", "対象外"),
            key="inv_gallery_status_filter",
        )
    gal_f1, gal_f2 = st.columns(2)
    with gal_f1:
        st.selectbox(
            "浮貸で絞り込み（ギャラリー）",
            ("指定なし", "浮貸あり", "浮貸なし"),
            key="inv_gallery_loan_filter",
            help="**浮貸あり** … 浮貸日時が入っている行のみ。**浮貸なし** … 浮貸日時が空の行のみ（販売済も含みます）。",
        )
    with gal_f2:
        st.selectbox(
            "棚卸しで絞り込み（ギャラリー）",
            (
                "指定なし",
                "台帳で棚卸日が未入力の在庫中のみ",
                "今回の作業でまだ未確認（在庫中）",
            ),
            key="inv_gallery_stocktake_filter",
        )

    _fw = str(st.session_state.get("inv_gallery_search_text", "") or "")
    _sup_f = list(st.session_state.get("inv_gallery_suppliers_filter") or [])
    _st_f = str(st.session_state.get("inv_gallery_status_filter", "すべて") or "すべて")
    _stk_f = str(
        st.session_state.get("inv_gallery_stocktake_filter", "指定なし") or "指定なし"
    )
    _loan_f = str(
        st.session_state.get("inv_gallery_loan_filter", "指定なし") or "指定なし"
    )
    _rem_gal = _inv_stocktake_work_remaining_get()
    if _stk_f == "今回の作業でまだ未確認（在庫中）" and _rem_gal is None:
        st.info("「今回の棚卸を開始」を押すと、この絞り込みが使えます。")
    df_view = _filter_inventory_df_for_view(
        df_sorted_calc,
        q=_fw,
        suppliers=_sup_f,
        status_mode=_st_f,
        stocktake_filter=_stk_f,
        stocktake_session_remaining=_rem_gal,
        loan_filter=_loan_f,
    )
    n_total = len(df_view)
    st.caption(
        f"該当 **{n_total:,}** 行（台帳全体 {len(df_sorted_calc):,} 行・粗利は再計算済み）。"
        f"表示は **{INV_GALLERY_PAGE_SIZE}** 件ずつです。"
    )

    if "inv_gallery_page" not in st.session_state:
        st.session_state.inv_gallery_page = 0
    _fp_gal = f"{_fw!r}|{repr(_sup_f)}|{_st_f!r}|{_stk_f!r}|{_loan_f!r}"
    if st.session_state.get("_inv_gallery_filter_fp") != _fp_gal:
        st.session_state._inv_gallery_filter_fp = _fp_gal
        st.session_state.inv_gallery_page = 0

    n_pages = max(1, (n_total + INV_GALLERY_PAGE_SIZE - 1) // INV_GALLERY_PAGE_SIZE)
    page_idx = int(st.session_state.inv_gallery_page)
    if page_idx >= n_pages:
        page_idx = n_pages - 1
        st.session_state.inv_gallery_page = page_idx
    if page_idx < 0:
        page_idx = 0
        st.session_state.inv_gallery_page = 0

    start_idx = page_idx * INV_GALLERY_PAGE_SIZE
    end_idx = min(n_total, start_idx + INV_GALLERY_PAGE_SIZE)
    df_tiles = df_view.iloc[start_idx:end_idx].reset_index(drop=True)

    if n_pages > 1:
        p1, p2, p3 = st.columns([1, 3, 1])
        with p1:
            if st.button(
                "◀ 前のページ",
                disabled=page_idx <= 0,
                key="inv_gallery_prev_page",
            ):
                st.session_state.inv_gallery_page = max(0, page_idx - 1)
                st.rerun()
        with p2:
            st.markdown(
                f"**ページ {page_idx + 1} / {n_pages}**　"
                f"（{n_total:,} 件中 **{start_idx + 1}〜{end_idx}** 件を表示）"
            )
        with p3:
            if st.button(
                "次のページ ▶",
                disabled=page_idx >= n_pages - 1,
                key="inv_gallery_next_page",
            ):
                st.session_state.inv_gallery_page = min(n_pages - 1, page_idx + 1)
                st.rerun()

    ncols = 4
    for i in range(0, len(df_tiles), ncols):
        gc = st.columns(ncols)
        for j in range(ncols):
            ridx = i + j
            if ridx >= len(df_tiles):
                break
            row = df_tiles.iloc[ridx]
            _st_tile = _normalize_stock_status(str(row.get(COL_STOCK_STATUS, "")))
            sold = _st_tile == STATUS_SOLD
            excluded = _st_tile == STATUS_EXCLUDED
            dim = sold or excluded
            mid = str(row.get(COL_MANAGEMENT_ID, "") or "").strip() or f"_{ridx}"
            safe_key = re.sub(r"[^\w\-]", "_", mid)[:48]
            with gc[j]:
                with st.container(border=True):
                    if excluded:
                        st.caption("対象外")
                    elif sold:
                        st.caption("販売済")
                    iu = str(row.get(COL_IMAGE_URL, "") or "").strip()
                    _img_w = 200 if dim else 240
                    _render_inventory_gallery_thumbnail(
                        iu, width=_img_w, sold=sold
                    )
                    st.markdown(
                        f'<p style="opacity:{"0.55" if dim else "1"};margin:0.2rem 0 0 0;font-size:1.05rem;">'
                        f"<b>{mid}</b></p>",
                        unsafe_allow_html=True,
                    )
                    nm = str(row.get(COL_NAME, "") or "").strip() or "—"
                    st.markdown(
                        f'<p style="opacity:{"0.55" if dim else "1"};margin:0;font-size:0.98rem;">'
                        f"{(nm if len(nm) <= 96 else nm[:93] + '…')}</p>",
                        unsafe_allow_html=True,
                    )
                    ps_raw = row.get(COL_PLANNED_SALE, "")
                    try:
                        psv = int(float(ps_raw)) if str(ps_raw).strip() != "" else 0
                    except (TypeError, ValueError):
                        psv = 0
                    _pl_lbl = (
                        f"販売予定（税抜） ¥{psv:,}"
                        if psv > 0
                        else "販売予定（税抜） —"
                    )
                    st.markdown(
                        f'<p style="opacity:{"0.55" if dim else "1"};margin:0;font-size:0.95rem;">'
                        f"{_pl_lbl}</p>",
                        unsafe_allow_html=True,
                    )
                    rd = {str(c): row.get(c) for c in EXPECTED_HEADERS if c in row.index}
                    if st.button(
                        "詳細",
                        key=f"inv_gal_dlg_{ridx}_{safe_key}",
                        use_container_width=True,
                    ):
                        _inventory_gallery_detail_dialog(rd)


def render_inventory_list_page(*, view_mode: str = "table") -> None:
    """在庫データの閲覧・編集。``view_mode`` は ``gallery``（カタログ）または ``table``（台帳表）。"""
    _vm = (view_mode or "table").strip().casefold()
    if _vm not in ("gallery", "table"):
        _vm = "table"
    if _vm == "gallery":
        st.markdown("## ギャラリー（カタログ）")
        st.caption(
            "共有の **inventory.csv** または **スプレッドシート**から読み込んだ在庫を、接客向けに **カード型** で表示します。"
            "**再読込**の直下にある **棚卸し: メトリクス・作業セッション・参照** を開くと、在庫一覧と同じメトリクス・棚卸作業・参照用一覧をまとめて表示できます。"
            f"1ページ **{INV_GALLERY_PAGE_SIZE}** 件ずつ切り替えられます。下の **表示の並び順** はこのページのタイル順にも反映されます。"
        )
    else:
        st.markdown("## 在庫一覧")
        st.caption(
            "共有の **inventory.csv** または **スプレッドシート**の全データを編集できます。行の追加・削除は表から操作し、"
            "「台帳を更新する」で保存します。"
            "「日時」は **保存時にセル内容が変わった行**（および表で追加した新規行）で **JST の現在時刻** に自動更新されます。"
            "「証憑記録日時」は証憑取込の確定時刻、「証憑URL」は Drive 上の証憑です（"
            "台帳内の値がすべて空または http(s) のときはリンク列として表示されクリックで開けます。"
            "http 以外の文字が混ざる行がある場合はテキスト列のままです）。"
            "棚卸し用の「最後に確認した日付（棚卸日）」は **YYYY-MM-DD** 推奨です（例: 今日なら "
            f"{_today_jst_date().isoformat()}）。"
            "カタログ表示はサイドバーの **ギャラリー（カタログ）** から開けます。"
            "棚卸のメトリクス・作業セッション・参照用一覧は下に続きます（参照用一覧のみ展開パネルです）。"
        )

    if msg := st.session_state.pop("_ledger_saved_flash", None):
        st.success(msg)

    if not _uses_local_inventory_csv() and not _secret_str(SECRET_GOOGLE_SPREADSHEET_ID):
        st.info(
            f"{SECRET_GOOGLE_SPREADSHEET_ID} を設定するか、`INVENTORY_SOURCE = \"csv\"` で "
            "共有の inventory.csv を有効にしてください。"
        )
        return

    r1, _ = st.columns([1, 2])
    with r1:
        _reload_label = (
            "inventory.csv から再読込"
            if _uses_local_inventory_csv()
            else "スプレッドシートから再読込"
        )
        if st.button(_reload_label, key="ledger_reload_from_sheet"):
            try:
                _inventory_csv_read_df_cached.clear()
            except Exception:
                pass
            try:
                _inventory_sheet_get_all_values_cached.clear()
            except Exception:
                pass
            _bump_inventory_sheet_cache_bust()
            st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
            st.rerun()

    try:
        df_sheet = load_inventory_dataframe()
    except Exception as e:
        st.error(f"読み込みに失敗しました: {e}")
        return

    if df_sheet is None:
        st.warning("台帳を開けませんでした。サービスアカウントと共有設定を確認してください。")
        return

    _inv_stocktake_work_remaining_prune(df_sheet)
    st_active, st_rem, st_base, st_origin, _st_snap = (
        _inv_stocktake_work_remaining_read_state(df_sheet)
    )
    _stk_outer = (
        st.expander("棚卸し: メトリクス・作業セッション・参照", expanded=False)
        if _vm == "gallery"
        else contextlib.nullcontext()
    )
    with _stk_outer:
        if _vm == "gallery":
            st.caption(
                "クリックで開閉します。**在庫一覧** ページではこのブロックは折りたたまず常に表示されます。"
            )
        n_in_stock = int(_mask_ledger_in_stock(df_sheet).sum())
        n_today_global = int(_mask_ledger_stocktake_today_jst(df_sheet).sum())
        if st_active and COL_MANAGEMENT_ID in df_sheet.columns:
            _m_sess_ct = (
                df_sheet[COL_MANAGEMENT_ID]
                .astype(str)
                .str.strip()
                .isin(st_rem)
                & _mask_ledger_in_stock(df_sheet)
            )
            n_session_in_stock_pending = int(_m_sess_ct.sum())
        else:
            n_session_in_stock_pending = 0
        n_session_confirmed_display: int | None = None
        if st_active:
            n_session_confirmed_display = max(0, n_in_stock - n_session_in_stock_pending)
        sk1, sk2, sk3, sk4 = st.columns(4)
        sk1.metric("在庫中（件数）", f"{n_in_stock:,}")
        sk2.metric("今回の作業でまだ未確認（在庫中）", f"{n_session_in_stock_pending:,}")
        sk3.metric(
            "今回の作業で確認済（在庫中）",
            f"{n_session_confirmed_display:,}" if n_session_confirmed_display is not None else "—",
            help=(
                "「在庫中（件数）」−「今回の作業でまだ未確認（在庫中）」として表示しています。"
                "作業セッション未開始のときは「—」です。セッション開始後に在庫が増えると、この差は今回リスト外の在庫分だけ大きくなることがあります。"
            ),
        )
        with sk4:
            if st_active and st_base > 0:
                n_rem_ids = len(st_rem)
                pct_done = 100.0 * (st_base - n_rem_ids) / st_base
                pct_done = max(0.0, min(100.0, pct_done))
                st.metric("今回リスト（残り／対象）", f"{n_rem_ids:,} / {st_base:,}")
                st.metric("今回リストの進捗", f"{pct_done:.1f}%")
            else:
                st.caption("今回の作業リストは未開始です。")
                st.metric("今回リスト（残り／対象）", "—")
                st.metric("今回リストの進捗", "—")

        st.markdown("##### 棚卸し作業セッション（任意）")
        st.caption(
            "同じ月・年に何度も棚卸しするとき、台帳に前回の棚卸日が入っていても **今回の対象リスト** で追えます。"
            "対象リストは **`"
            + STOCKTAKE_WORK_SESSION_FILENAME
            + "`** に保存されるため、ブラウザやアプリを閉じてもリセットされません。"
            "「今回の棚卸を開始」で在庫中の全管理IDを対象にし、棚卸し登録の確定・一括棚卸日・台帳保存で棚卸日を付けた行は自動でリストから外れます。"
            "残りがゼロになった時点でもセッションは終了します（全数確認済み）。手動で閉じる場合は「今回の対象リストを終了」を押してください（台帳の日付は変わりません）。"
        )
        ss1, ss2 = st.columns(2)
        with ss1:
            if st.button(
                "今回の棚卸を開始（在庫中をすべて今回の対象に）",
                key="inv_stocktake_work_start",
            ):
                _inv_stocktake_work_remaining_start(df_sheet)
                st.session_state.inv_gallery_page = 0
                st.rerun()
        with ss2:
            if st.button(
                "今回の対象リストを終了",
                key="inv_stocktake_work_end",
                disabled=not st_active,
            ):
                _inv_stocktake_work_remaining_clear()
                st.session_state.inv_gallery_page = 0
                st.rerun()

        with st.expander("棚卸し: 参照用一覧（台帳未入力 / 今回の作業）", expanded=False):
            list_kind = st.radio(
                "表示する一覧",
                ("台帳で棚卸日が未入力の在庫中", "今回の作業でまだ未確認の在庫中"),
                horizontal=True,
                key="inv_stocktake_list_kind_radio",
            )
            _ucols = [
                c
                for c in (
                    COL_MANAGEMENT_ID,
                    COL_NAME,
                    COL_SUPPLIER,
                    COL_DATETIME,
                    COL_LAST_STOCKTAKE,
                )
                if c in df_sheet.columns
            ]
            if list_kind.startswith("台帳"):
                st.caption(
                    "「最後に確認した日付（棚卸日）」が空、または日付として解釈できない **在庫中** のみです。"
                    "日付の入力・保存は下の表・スキャン・一括ボタンで行ってください。"
                )
                unv = df_sheet.loc[_mask_ledger_stocktake_unverified(df_sheet)].copy()
                if unv.empty:
                    st.success("在庫中で、かつ棚卸日が未入力の行はありません。")
                else:
                    st.metric("この一覧の件数", f"{len(unv):,}")
                    st.dataframe(unv[_ucols], use_container_width=True, hide_index=True)
            else:
                st.caption(
                    "上で **今回の棚卸を開始** を押したあとのみ有効です。台帳に棚卸日が入っていても、まだ今回のリストに残っている **在庫中** の行です。"
                )
                if not st_active:
                    st.info("作業セッションが未開始です。「今回の棚卸を開始」を押してください。")
                elif not st_rem:
                    st.success(
                        "今回の作業で追っていた在庫中の行は、すべてリストから外れました（または開始時点で在庫中がゼロでした）。"
                    )
                else:
                    m_sess = df_sheet[COL_MANAGEMENT_ID].astype(str).str.strip().isin(st_rem)
                    sess_df = df_sheet.loc[m_sess & _mask_ledger_in_stock(df_sheet)].copy()
                    if COL_MANAGEMENT_ID in sess_df.columns and not sess_df.empty:
                        sess_df = sess_df.copy()
                        sess_df["_sk"] = sess_df[COL_MANAGEMENT_ID].astype(str).str.strip().map(
                            _management_id_sort_key
                        )
                        sess_df = sess_df.sort_values("_sk").drop(columns=["_sk"])
                    st.metric("この一覧の件数（今回の残り・在庫中）", f"{len(sess_df):,}")
                    st.dataframe(sess_df[_ucols], use_container_width=True, hide_index=True)

        if (
            st_active
            and st_origin
            and COL_MANAGEMENT_ID in df_sheet.columns
            and n_session_confirmed_display is not None
            and n_session_confirmed_display > 0
        ):
            _ids_show = _management_ids_origin_cleared_session_in_stock(
                df_sheet, st_origin, st_rem, limit=18
            )
            _cap = (
                f"上記「確認済」＝在庫中（件数）−今回の作業でまだ未確認（在庫中）＝**{n_session_confirmed_display}** 件です。"
            )
            if _ids_show:
                tail = " …" if n_session_confirmed_display > len(_ids_show) else ""
                _cap += (
                    f" 今回の開始対象から外れた在庫中の管理IDの例: "
                    f"{', '.join(_ids_show)}{tail}"
                )
            st.caption(_cap)
        elif n_today_global > 0 and COL_MANAGEMENT_ID in df_sheet.columns:
            _td_rows = df_sheet.loc[_mask_ledger_stocktake_today_jst(df_sheet)]
            _ids_show = (
                _td_rows[COL_MANAGEMENT_ID].astype(str).str.strip().head(18).tolist()
            )
            tail = " …" if len(_td_rows) > len(_ids_show) else ""
            st.caption(
                f"今日（JST {_today_jst_date().isoformat()}）の棚卸日が入っている在庫中（台帳全体）: **{n_today_global}** 件。"
                f"管理IDの例: {', '.join(_ids_show)}{tail}"
            )

    st.markdown("##### 表示の並び順（台帳表に反映・保存時もこの順で書き込みます）")
    s1, s2, s3, s4, s5, s6 = st.columns([2, 1, 2, 1, 2, 1])
    sort_choices = ["日時", "仕入先・取引先", "管理ID", "仕入日時", "販売日時", "なし"]
    with s1:
        prim = st.selectbox("第1ソート", sort_choices, index=5, key="ledger_sort_p")
    with s2:
        prim_ord = st.radio("第1の順序", ["昇順", "降順"], horizontal=True, key="ledger_sort_p_ord")
    with s3:
        sec = st.selectbox("第2ソート", sort_choices, index=5, key="ledger_sort_s")
    with s4:
        sec_ord = st.radio("第2の順序", ["昇順", "降順"], horizontal=True, key="ledger_sort_s_ord")
    with s5:
        ter = st.selectbox("第3ソート", sort_choices, index=5, key="ledger_sort_t")
    with s6:
        ter_ord = st.radio("第3の順序", ["昇順", "降順"], horizontal=True, key="ledger_sort_t_ord")

    df_sorted = _apply_ledger_sort(
        df_sheet,
        prim,
        prim_ord == "昇順",
        sec,
        sec_ord == "昇順",
        ter,
        ter_ord == "昇順",
    )

    df_sorted_calc = _recalc_gross_profit_dataframe(df_sorted.copy())

    if _vm == "gallery":
        _render_inventory_gallery_body(df_sorted_calc)
    else:
        st.markdown("### 在庫一覧（表）")
        st.caption(
            "全列・全行を表示します。行数が多いときは表の **縦スクロール** で移動してください。"
        )
        _render_inventory_ledger_data_editor_section(df_sorted_calc)


def _init_registration_form_session_state() -> None:
    """登録フォーム用の session_state 初期値（キーはウィジェットと連動）。"""
    if "field_product_name" not in st.session_state:
        st.session_state.field_product_name = ""
    if "field_supplier" not in st.session_state:
        st.session_state.field_supplier = ""
    if "field_row_quantity" not in st.session_state:
        st.session_state.field_row_quantity = 1
    if "field_inventory_category" not in st.session_state:
        st.session_state.field_inventory_category = ""
    if "ai_kind" not in st.session_state:
        st.session_state.ai_kind = ""
    if "ai_features" not in st.session_state:
        st.session_state.ai_features = ""
    if "ai_parse_ran" not in st.session_state:
        st.session_state.ai_parse_ran = False
    if "field_memo" not in st.session_state:
        st.session_state.field_memo = ""
    if "field_line_excl_yen" not in st.session_state:
        st.session_state.field_line_excl_yen = 1
    st.session_state.pop("field_unit_price_excl", None)
    if "field_consumption_tax_choice" not in st.session_state:
        st.session_state.field_consumption_tax_choice = "10%"
    if "field_planned_sale_excl" not in st.session_state:
        st.session_state.field_planned_sale_excl = 0
    if "field_actual_sale_excl" not in st.session_state:
        st.session_state.field_actual_sale_excl = 0
    if "field_stock_status" not in st.session_state:
        st.session_state.field_stock_status = STATUS_IN_STOCK
    if "hint_filter_product_name" not in st.session_state:
        st.session_state.hint_filter_product_name = ""
    if "hint_filter_supplier" not in st.session_state:
        st.session_state.hint_filter_supplier = ""
    if "ledger_pick_product_name" not in st.session_state:
        st.session_state.ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
    if "ledger_pick_supplier" not in st.session_state:
        st.session_state.ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
    if "hint_filter_inventory_category" not in st.session_state:
        st.session_state.hint_filter_inventory_category = ""
    if "hint_filter_management_id" not in st.session_state:
        st.session_state.hint_filter_management_id = ""
    if "ledger_pick_inventory_category" not in st.session_state:
        st.session_state.ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER
    if "ledger_pick_management_id" not in st.session_state:
        st.session_state.ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
    if "field_sale_source_mgmt_id" not in st.session_state:
        st.session_state.field_sale_source_mgmt_id = ""
    if "sale_pick_source_id" not in st.session_state:
        st.session_state.sale_pick_source_id = LEDGER_PICK_PLACEHOLDER
    if "sales_hint_filter_product_name" not in st.session_state:
        st.session_state.sales_hint_filter_product_name = ""
    if "sales_hint_filter_supplier" not in st.session_state:
        st.session_state.sales_hint_filter_supplier = ""
    if "sales_hint_filter_inventory_category" not in st.session_state:
        st.session_state.sales_hint_filter_inventory_category = ""
    if "sales_hint_filter_management_id" not in st.session_state:
        st.session_state.sales_hint_filter_management_id = ""
    if "sales_ledger_pick_product_name" not in st.session_state:
        st.session_state.sales_ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
    if "sales_ledger_pick_supplier" not in st.session_state:
        st.session_state.sales_ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
    if "sales_ledger_pick_inventory_category" not in st.session_state:
        st.session_state.sales_ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER
    if "sales_ledger_pick_management_id" not in st.session_state:
        st.session_state.sales_ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
    if "sales_assist_buf_product_name" not in st.session_state:
        st.session_state.sales_assist_buf_product_name = ""
    if "sales_assist_buf_supplier" not in st.session_state:
        st.session_state.sales_assist_buf_supplier = ""
    if "sales_assist_buf_inventory_category" not in st.session_state:
        st.session_state.sales_assist_buf_inventory_category = ""
    if "sales_assist_buf_management_id" not in st.session_state:
        st.session_state.sales_assist_buf_management_id = ""
    if "stocktake_hint_filter_product_name" not in st.session_state:
        st.session_state.stocktake_hint_filter_product_name = ""
    if "stocktake_hint_filter_supplier" not in st.session_state:
        st.session_state.stocktake_hint_filter_supplier = ""
    if "stocktake_hint_filter_inventory_category" not in st.session_state:
        st.session_state.stocktake_hint_filter_inventory_category = ""
    if "stocktake_hint_filter_management_id" not in st.session_state:
        st.session_state.stocktake_hint_filter_management_id = ""
    if "stocktake_ledger_pick_product_name" not in st.session_state:
        st.session_state.stocktake_ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
    if "stocktake_ledger_pick_supplier" not in st.session_state:
        st.session_state.stocktake_ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
    if "stocktake_ledger_pick_inventory_category" not in st.session_state:
        st.session_state.stocktake_ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER
    if "stocktake_ledger_pick_management_id" not in st.session_state:
        st.session_state.stocktake_ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
    if "stocktake_assist_buf_product_name" not in st.session_state:
        st.session_state.stocktake_assist_buf_product_name = ""
    if "stocktake_assist_buf_supplier" not in st.session_state:
        st.session_state.stocktake_assist_buf_supplier = ""
    if "stocktake_assist_buf_inventory_category" not in st.session_state:
        st.session_state.stocktake_assist_buf_inventory_category = ""
    if "stocktake_assist_buf_management_id" not in st.session_state:
        st.session_state.stocktake_assist_buf_management_id = ""
    if "s_reg_qty" not in st.session_state:
        st.session_state.s_reg_qty = 1
    if "s_field_sale_source_mgmt_id" not in st.session_state:
        st.session_state.s_field_sale_source_mgmt_id = ""
    if "s_field_actual_sale_excl" not in st.session_state:
        st.session_state.s_field_actual_sale_excl = 0
    if "s_field_memo" not in st.session_state:
        st.session_state.s_field_memo = ""
    if "sales_tab_memo" not in st.session_state:
        st.session_state.sales_tab_memo = ""
    if "sales_tab_loan_datetime_manual" not in st.session_state:
        st.session_state.sales_tab_loan_datetime_manual = ""
    st.session_state.pop("field_price_excl", None)


def _management_id_sort_key(mid: str) -> tuple[int, str]:
    """管理ID G######## を数値昇順でソート（非標準形式は末尾）。"""
    s = str(mid or "").strip()
    m = re.fullmatch(r"(?i)G(\d+)", s)
    if m:
        return (int(m.group(1)), "")
    return (10**18, s.casefold())


def _stocktake_candidates_from_gemini_response(
    res: dict[str, Any],
    df_ledger: pd.DataFrame,
    *,
    min_conf: float = STOCKTAKE_CAND_MIN_CONFIDENCE,
    max_n: int = STOCKTAKE_CAND_AI_MAX,
    allowed_management_ids: set[str] | None = None,
) -> list[dict[str, Any]]:
    """Gemini の JSON から棚卸し候補を正規化（在庫中・台帳に存在する行のみ、重複除去）。

    ``allowed_management_ids`` … 指定時はその集合に含まれる管理 ID のみ採用（今回の対象リスト）。
    """
    raw_list: list[dict[str, Any]] = []
    if isinstance(res, dict):
        sc = res.get("stocktake_candidates")
        if isinstance(sc, list):
            for item in sc:
                if isinstance(item, dict):
                    raw_list.append(item)
        m0 = res.get("match")
        if isinstance(m0, dict) and str(m0.get("management_id") or "").strip():
            raw_list.append(m0)
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for it in raw_list:
        mid = str(it.get("management_id") or "").strip()
        if not mid or mid in seen:
            continue
        try:
            conf = float(it.get("confidence") or 0)
        except (TypeError, ValueError):
            conf = 0.0
        if conf < min_conf:
            continue
        if allowed_management_ids is not None and mid not in allowed_management_ids:
            continue
        tr = lookup_ledger_row_by_management_id(df_ledger, mid)
        if tr is None:
            continue
        if (
            _normalize_stock_status(str(tr.get(COL_STOCK_STATUS, "")))
            != STATUS_IN_STOCK
        ):
            continue
        seen.add(mid)
        pn = str(it.get("product_name") or "").strip() or str(
            tr.get(COL_NAME, "") or ""
        ).strip()
        su = str(it.get("supplier") or "").strip() or str(
            tr.get(COL_SUPPLIER, "") or ""
        ).strip()
        out.append(
            {
                "management_id": mid,
                "confidence": conf,
                "product_name": pn,
                "supplier": su,
                "last_stocktake": str(tr.get(COL_LAST_STOCKTAKE, "") or "").strip(),
                "image_url": str(tr.get(COL_IMAGE_URL, "") or "").strip(),
                "memo": str(tr.get(COL_MEMO, "") or "").strip(),
                "feature_observation": str(
                    it.get("feature_observation") or ""
                ).strip(),
            }
        )
    out.sort(key=lambda x: _management_id_sort_key(str(x.get("management_id") or "")))
    return out[:max_n]


def _flush_stocktake_multiselect_pending() -> None:
    """ボタンからの multiselect 更新を rerun 先頭で適用し、Streamlit のウィジェット状態競合を避ける。"""
    pk = _PENDING_STOCKTAKE_ASSIST_BATCH_MIDS
    if pk in st.session_state:
        v = st.session_state.pop(pk)
        if v is None:
            st.session_state.pop("stocktake_assist_batch_mids", None)
        else:
            st.session_state["stocktake_assist_batch_mids"] = v
    pk2 = _PENDING_STOCKTAKE_MULTI_DONE_MIDS
    if pk2 in st.session_state:
        v2 = st.session_state.pop(pk2)
        if v2 is None:
            st.session_state.pop("stocktake_multi_done_mids", None)
        else:
            st.session_state["stocktake_multi_done_mids"] = v2


def render_stocktake_scan_tab(
    uploaded,
    df_ledger_hint: pd.DataFrame | None,
) -> None:
    """棚卸し登録: 共通アップロード画像で AI 照合 → 棚卸日の確定更新のみ。"""
    _flush_stocktake_multiselect_pending()
    if df_ledger_hint is not None and not df_ledger_hint.empty:
        # 在庫追加・販売反映のあとにこのタブを開いたときも、
        # 「今回の対象リスト」を最新の在庫中管理IDへ自動同期する。
        _inv_stocktake_work_remaining_prune(df_ledger_hint)
    st.markdown("##### 棚卸し登録（AI 照合）")
    st_rem_scan = _inv_stocktake_work_remaining_get()
    _scan_targets_ok = st_rem_scan is not None and len(st_rem_scan) > 0
    if not _scan_targets_ok:
        for _k in (
            "_stocktake_scan_candidates",
            "stocktake_multi_done_mids",
            "stocktake_cand_page",
        ):
            st.session_state.pop(_k, None)
        if st_rem_scan is not None:
            st.session_state.pop("_stocktake_selected_mid", None)
    st.caption("共通アップロード画像で照合し、候補から1件または複数件を棚卸確定できます。")
    with st.expander("使い方", expanded=False):
        st.markdown(
            f"- 候補は {STOCKTAKE_CAND_PAGE_SIZE} 件ずつ表示します（最大 {STOCKTAKE_CAND_AI_MAX} 件）。\n"
            "- 候補は管理ID昇順で表示されます。\n"
            "- 確定は棚卸日の更新のみで、新規行は追加しません。"
        )
    if not _scan_targets_ok:
        if st_rem_scan is None:
            st.info(
                "先に在庫一覧で **今回の棚卸を開始（在庫中をすべて今回の対象に）** を押してください。"
                "スキャンは **今回のリストに残っている行だけ** を照合対象にします。"
            )
        else:
            st.info(
                "今回の対象リストに **残っている行がありません**（すべて棚卸済みか、台帳更新でリストから外れました）。"
                "続ける場合は **今回の棚卸を開始** でリストを作り直してください。"
            )
    st.caption("写真はページ上部の共通アップローダを使用します。")

    if st.button(
        "AIで台帳と照合",
        type="primary",
        key="stocktake_ai_match_btn",
        disabled=(not _scan_targets_ok) or (uploaded is None),
    ):
        st.session_state.pop("_stocktake_scan_candidates", None)
        st.session_state.pop("_stocktake_selected_mid", None)
        st.session_state.pop("stocktake_multi_done_mids", None)
        st.session_state.pop("_stocktake_scan_warn", None)
        st.session_state.pop("stocktake_cand_page", None)
        st_rem_run = _inv_stocktake_work_remaining_get()
        if st_rem_run is None or not st_rem_run:
            st.session_state["_stocktake_scan_warn"] = (
                "今回の棚卸対象リストがありません。**今回の棚卸を開始** からやり直してください。"
            )
        elif uploaded is None:
            st.session_state["_stocktake_scan_warn"] = (
                "上部の共通アップローダで写真を選択してください。"
            )
        elif df_ledger_hint is None or df_ledger_hint.empty:
            st.session_state["_stocktake_scan_warn"] = "台帳を読み込めないため照合できません。"
        else:
            n_scope_st = len(st_rem_run or [])
            max_lines_st_ctx = min(
                STOCKTAKE_AI_CONTEXT_MAX_LINES,
                max(400, n_scope_st + 20),
            )
            inv_ctx_st = _build_gemini_inventory_context(
                df_ledger_hint,
                max_lines=max_lines_st_ctx,
                only_in_stock=True,
                management_ids_filter=st_rem_run,
            )
            if not (inv_ctx_st or "").strip():
                st.session_state["_stocktake_scan_warn"] = (
                    "今回の対象リストに **在庫中として残っている行** がありません。"
                    "在庫一覧で対象を開始し直すか、残リストを確認してください。"
                )
            else:
                with st.spinner("画像を解析して台帳と照合しています…"):
                    try:
                        raw = analyze_image_with_gemini(
                            uploaded,
                            inventory_context=inv_ctx_st or None,
                            prompt_mode="stocktake_match",
                        )
                        res = _parse_json_from_model(raw or "")
                        if not isinstance(res, dict):
                            res = {}
                        cand_list = _stocktake_candidates_from_gemini_response(
                            res,
                            df_ledger_hint,
                            allowed_management_ids=st_rem_run,
                        )
                        if not cand_list:
                            st.session_state["_stocktake_scan_warn"] = (
                                "今回の対象リストの在庫中の行で、写真に合いそうな候補が得られませんでした。"
                                "明るさ・距離を変えて再撮影するか、在庫一覧で管理IDを確認してください。"
                            )
                        else:
                            st.session_state["_stocktake_scan_candidates"] = cand_list
                            st.session_state["stocktake_cand_page"] = 0
                            st.session_state.pop("_stocktake_selected_mid", None)
                    except Exception as e:
                        st.session_state["_stocktake_scan_warn"] = str(e)

    if "stocktake_assist_visible" not in st.session_state:
        st.session_state.stocktake_assist_visible = False
    if st.button(
        "台帳からの入力補助を表示"
        if not st.session_state.stocktake_assist_visible
        else "台帳からの入力補助を非表示",
        key="stocktake_assist_toggle_btn",
    ):
        st.session_state.stocktake_assist_visible = (
            not st.session_state.stocktake_assist_visible
        )
        st.rerun()

    if (
        st.session_state.stocktake_assist_visible
        and _scan_targets_ok
        and df_ledger_hint is not None
        and not df_ledger_hint.empty
    ):
        _st_hint_df = _stocktake_assist_scope_dataframe(df_ledger_hint, st_rem_scan)
        if _st_hint_df is not None and not _st_hint_df.empty:
            _render_ledger_pick_assist_three_columns(
                _st_hint_df,
                key_prefix="stocktake_",
                body_caption=(
                    "仕入タブと同じ操作で、**今回の棚卸リストに残っている在庫中行** から絞り込めます。"
                    "商品名・仕入先・在庫カテゴリー・**管理ID** は、文字の絞り込みまたはプルダウンから選べます。"
                    "**組み合わせで一致が1件だけ** のときだけ、**選択中の管理ID** を自動セットします。"
                    "複数件のときは下の近い候補カードと AI 照合で確認してください。"
                ),
                on_pick_product_name=_on_stocktake_assist_pick_product_name,
                on_pick_supplier=_on_stocktake_assist_pick_supplier,
                on_pick_inventory_category=_on_stocktake_assist_pick_inventory_category,
                on_pick_management_id=_on_stocktake_assist_pick_management_id,
            )
            stn = int(
                st.session_state.get("stocktake_assist_last_n_matching_mids", 0) or 0
            )
            if stn > 1:
                st.info(
                    f"補助条件に一致する対象行が **{stn}** 件あります。"
                    "一致が1件だけのときのみ **選択中の管理ID** を自動で入れます。"
                )
            elif stn == 1 and str(
                st.session_state.get("_stocktake_selected_mid", "") or ""
            ).strip():
                _tsm = str(st.session_state["_stocktake_selected_mid"]).strip()
                st.caption(f"選択中の管理ID（補助）: **{_tsm}**")
            _refresh_stocktake_assist_quick_candidates(df_ledger_hint, st_rem_scan)
            _stk_c = st.session_state.get("stocktake_assist_quick_candidates")
            if (
                isinstance(_stk_c, pd.DataFrame)
                and not _stk_c.empty
                and df_ledger_hint is not None
            ):
                st.markdown("##### 近い候補（補助から照合・カード表示）")
                st.caption(
                    "今回のリストの在庫中行のうち、補助で確定した内容に近い行を表示します。"
                    "**1件選択** ではカードの **この候補を選ぶ** のあと、下の **台帳入力補助で選んだ管理IDの棚卸確定**（AI 照合と併用可）で確定します。"
                    "**複数選択** では一覧・ボタンでまとめて選び、**選択した行をまとめて棚卸確定** で一括反映します（AI 照合の候補がないときのみ）。"
                )
                _stk_hits = [
                    _sale_card_hit_from_series(row)
                    for _, row in _stk_c.iterrows()
                ]
                _st_assist_mode = st.radio(
                    "台帳入力補助の棚卸確定の仕方",
                    ("1件選択", "複数選択（一括反映）"),
                    horizontal=True,
                    key="stocktake_assist_confirm_mode",
                )
                _st_assist_batch = str(_st_assist_mode or "").startswith("複数")
                _mid_opts_a: list[str] = []
                _seen_m: set[str] = set()
                for _h in _stk_hits:
                    _m = str(_h.get("management_id") or "").strip()
                    if _m and _m not in _seen_m:
                        _seen_m.add(_m)
                        _mid_opts_a.append(_m)
                _mid_opts_a.sort(key=_management_id_sort_key)
                _mid_label_a: dict[str, str] = {}
                for _h in _stk_hits:
                    _m = str(_h.get("management_id") or "").strip()
                    if not _m:
                        continue
                    _pn = str(_h.get("product_name") or "—").strip()
                    if len(_pn) > 36:
                        _pn = _pn[:33] + "…"
                    _mid_label_a[_m] = f"{_m} ／ {_pn}"
                if "stocktake_assist_cand_page" not in st.session_state:
                    st.session_state.stocktake_assist_cand_page = 0
                n_tot_a = len(_mid_opts_a)
                n_pg_a = max(
                    1,
                    (n_tot_a + STOCKTAKE_CAND_PAGE_SIZE - 1) // STOCKTAKE_CAND_PAGE_SIZE,
                )
                pg_a = int(st.session_state.stocktake_assist_cand_page)
                pg_a = max(0, min(n_pg_a - 1, pg_a))
                st.session_state.stocktake_assist_cand_page = pg_a
                si_a = pg_a * STOCKTAKE_CAND_PAGE_SIZE
                ei_a = min(n_tot_a, si_a + STOCKTAKE_CAND_PAGE_SIZE)
                page_mids_a = _mid_opts_a[si_a:ei_a]
                _by_mid_a = {
                    str(h.get("management_id") or "").strip(): h for h in _stk_hits
                }
                page_hits_a = [_by_mid_a[m] for m in page_mids_a if m in _by_mid_a]
                _has_ai_here = isinstance(
                    st.session_state.get("_stocktake_scan_candidates"), list
                ) and len(st.session_state.get("_stocktake_scan_candidates") or []) > 0
                if _st_assist_batch:
                    st.caption(
                        "一覧と **この候補を選ぶ** で選択に追加できます。**選択した行をまとめて棚卸確定** で本日（JST）の棚卸日を一括反映します。"
                    )
                    ba1, ba2, ba3, ba4 = st.columns(4)
                    with ba1:
                        if st.button(
                            "補助候補をすべて選択",
                            key="stocktake_assist_sel_all",
                            disabled=not _mid_opts_a,
                        ):
                            st.session_state[_PENDING_STOCKTAKE_ASSIST_BATCH_MIDS] = list(
                                _mid_opts_a
                            )
                            st.rerun()
                    with ba2:
                        if st.button(
                            "このページの候補をすべて選択",
                            key="stocktake_assist_sel_page",
                            disabled=not page_mids_a,
                        ):
                            _cur = set(
                                st.session_state.get("stocktake_assist_batch_mids")
                                or []
                            )
                            _cur.update(page_mids_a)
                            st.session_state[_PENDING_STOCKTAKE_ASSIST_BATCH_MIDS] = sorted(
                                _cur, key=_management_id_sort_key
                            )
                            st.rerun()
                    with ba3:
                        if st.button("選択をクリア", key="stocktake_assist_clr_sel"):
                            st.session_state[_PENDING_STOCKTAKE_ASSIST_BATCH_MIDS] = None
                            st.rerun()
                    with ba4:
                        st.caption(
                            f"補助候補 **{n_tot_a}** 件中、選択中 **{len(st.session_state.get('stocktake_assist_batch_mids') or [])}** 件"
                        )
                    st.multiselect(
                        "このページ内を任意選択（追加用）",
                        options=page_mids_a,
                        format_func=lambda m: _mid_label_a.get(m, m),
                        key="stocktake_assist_page_pick",
                    )
                    if st.button(
                        "このページの任意選択を追加",
                        key="stocktake_assist_page_pick_add",
                        disabled=not bool(
                            st.session_state.get("stocktake_assist_page_pick")
                        ),
                    ):
                        _cur = set(
                            st.session_state.get("stocktake_assist_batch_mids") or []
                        )
                        _cur.update(
                            [
                                str(x).strip()
                                for x in st.session_state.get(
                                    "stocktake_assist_page_pick", []
                                )
                                if str(x).strip()
                            ]
                        )
                        st.session_state[_PENDING_STOCKTAKE_ASSIST_BATCH_MIDS] = sorted(
                            _cur, key=_management_id_sort_key
                        )
                        st.rerun()
                    st.multiselect(
                        "一括で棚卸確定する管理ID（任意に追加・解除）",
                        options=_mid_opts_a,
                        format_func=lambda m: _mid_label_a.get(m, m),
                        key="stocktake_assist_batch_mids",
                    )
                else:
                    _sel_one = str(
                        st.session_state.get("_stocktake_selected_mid") or ""
                    ).strip()
                    if _sel_one:
                        st.info(f"選択中の管理ID: **{_sel_one}**")
                if n_pg_a > 1:
                    p1, p2, p3 = st.columns([1, 3, 1])
                    with p1:
                        if st.button(
                            "◀ 前へ",
                            disabled=pg_a <= 0,
                            key="stocktake_assist_pg_prev",
                        ):
                            st.session_state.stocktake_assist_cand_page = max(
                                0, pg_a - 1
                            )
                            st.rerun()
                    with p2:
                        st.caption(
                            f"**ページ {pg_a + 1} / {n_pg_a}**（{n_tot_a} 件中 **{si_a + 1}〜{ei_a}** 件）"
                        )
                    with p3:
                        if st.button(
                            "次へ ▶",
                            disabled=pg_a >= n_pg_a - 1,
                            key="stocktake_assist_pg_next",
                        ):
                            st.session_state.stocktake_assist_cand_page = min(
                                n_pg_a - 1, pg_a + 1
                            )
                            st.rerun()
                if _st_assist_batch:
                    _cap_a = (
                        f"全 **{n_tot_a}** 件中 **{si_a + 1}〜{ei_a}** 件（カード）"
                    )
                    _render_mid_pick_candidate_cards(
                        page_hits_a,
                        widget_key_namespace=f"stk_assist_pg_{pg_a}",
                        sold=False,
                        pick_mode="stocktake_merge",
                        pager="hidden",
                        caption_override=_cap_a,
                    )
                else:
                    _render_mid_pick_candidate_cards(
                        _stk_hits,
                        widget_key_namespace="stk_assist",
                        sold=False,
                        pick_mode="stocktake",
                    )
                if (
                    _st_assist_batch
                    and not _has_ai_here
                    and _scan_targets_ok
                ):
                    _picked_a = [
                        str(x).strip()
                        for x in (st.session_state.get("stocktake_assist_batch_mids") or [])
                        if str(x).strip()
                    ]
                    if st.button(
                        "選択した行をまとめて棚卸確定（棚卸日を本日・JST）",
                        type="primary",
                        key="stocktake_assist_confirm_batch",
                        disabled=len(_picked_a) < 1,
                    ):
                        try:
                            with st.spinner("台帳を更新しています…"):
                                n_ok, skips = apply_last_stocktake_jst_for_management_ids(
                                    _picked_a
                                )
                        except Exception as e:
                            st.error(str(e))
                        else:
                            st.session_state.pop("_stocktake_scan_candidates", None)
                            st.session_state.pop("_stocktake_selected_mid", None)
                            st.session_state.pop("stocktake_multi_done_mids", None)
                            st.session_state.pop("stocktake_cand_page", None)
                            st.session_state.pop("stocktake_assist_batch_mids", None)
                            st.session_state.pop("stocktake_assist_page_pick", None)
                            st.success(
                                f"**{n_ok}** 件の棚卸日を本日（JST）に更新しました。"
                            )
                            if skips:
                                st.caption(
                                    "スキップ: "
                                    + "；".join(skips[:8])
                                    + (" …" if len(skips) > 8 else "")
                                )
                            st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                            st.rerun()
        else:
            st.caption("今回のリストに該当する在庫中行が台帳にありません。")
    elif (
        st.session_state.stocktake_assist_visible
        and df_ledger_hint is not None
        and not df_ledger_hint.empty
    ):
        st.caption(
            "**今回の棚卸を開始** して対象リストがあるときだけ、ここに仕入タブと同様の台帳入力補助が表示されます。"
        )

    _cands_pre = st.session_state.get("_stocktake_scan_candidates")
    _has_ai_stocktake_cands = isinstance(_cands_pre, list) and len(_cands_pre) > 0
    _assist_mode_g = str(
        st.session_state.get("stocktake_assist_confirm_mode", "") or ""
    )
    _assist_batch_g = _assist_mode_g.startswith("複数")
    _mid_from_assist = str(
        st.session_state.get("_stocktake_selected_mid") or ""
    ).strip()
    if (
        _scan_targets_ok
        and not _has_ai_stocktake_cands
        and not _assist_batch_g
        and _mid_from_assist
    ):
        st.markdown("##### 台帳入力補助で選んだ管理IDの棚卸確定")
        st.caption(
            "AI 照合の候補がないとき、または補助の近い候補カードで管理IDを選んだあと、ここから棚卸日を本日（JST）に更新できます。"
        )
        st.info(f"選択中の管理ID: **`{_mid_from_assist}`**")
        if st.button(
            "棚卸を確定（棚卸日を本日・JST に更新）",
            type="primary",
            key="stocktake_confirm_from_ledger_assist",
        ):
            try:
                with st.spinner("台帳を更新しています…"):
                    apply_last_stocktake_jst_for_management_ids([_mid_from_assist])
            except Exception as e:
                st.error(str(e))
            else:
                st.session_state.pop("_stocktake_scan_candidates", None)
                st.session_state.pop("_stocktake_selected_mid", None)
                st.session_state.pop("stocktake_multi_done_mids", None)
                st.session_state.pop("stocktake_cand_page", None)
                st.session_state.pop("stocktake_assist_batch_mids", None)
                st.session_state.pop("stocktake_assist_page_pick", None)
                st.success(
                    f"管理ID **{_mid_from_assist}** の棚卸日を更新しました。"
                )
                st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                st.rerun()

    if st.button("入力をクリア", key="stocktake_assist_clear_btn"):
        st.session_state.stocktake_hint_filter_product_name = ""
        st.session_state.stocktake_hint_filter_supplier = ""
        st.session_state.stocktake_hint_filter_inventory_category = ""
        st.session_state.stocktake_hint_filter_management_id = ""
        st.session_state.stocktake_ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
        st.session_state.stocktake_ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
        st.session_state.stocktake_ledger_pick_inventory_category = (
            LEDGER_PICK_PLACEHOLDER
        )
        st.session_state.stocktake_ledger_pick_management_id = (
            LEDGER_PICK_PLACEHOLDER
        )
        st.session_state.stocktake_assist_buf_product_name = ""
        st.session_state.stocktake_assist_buf_supplier = ""
        st.session_state.stocktake_assist_buf_inventory_category = ""
        st.session_state.stocktake_assist_buf_management_id = ""
        st.session_state.pop("stocktake_assist_quick_candidates", None)
        st.session_state.pop("stocktake_assist_last_n_matching_mids", None)
        st.session_state.pop("_stocktake_selected_mid", None)
        st.session_state.pop("_stocktake_scan_candidates", None)
        st.session_state.pop("stocktake_multi_done_mids", None)
        st.session_state.pop("stocktake_cand_page", None)
        st.session_state.pop("_stocktake_scan_warn", None)
        st.session_state.pop("stocktake_assist_batch_mids", None)
        st.session_state.pop("stocktake_assist_page_pick", None)
        st.session_state.pop("stocktake_assist_cand_page", None)
        st.rerun()

    wn = st.session_state.pop("_stocktake_scan_warn", None)
    if wn:
        st.warning(wn)
    cands = st.session_state.get("_stocktake_scan_candidates")
    if isinstance(cands, list) and cands:
        st.markdown("### 照合候補（今回の対象リスト・在庫中）")
        st.caption(
            f"**{len(cands)}** 件（**管理IDの昇順**）。"
            f"**{STOCKTAKE_CAND_PAGE_SIZE}** 件ずつ表示し、下のボタンで全候補をページ送りできます。"
        )
        _st_mode = st.radio(
            "棚卸の確定の仕方",
            (
                "1件選択",
                "複数選択（一括反映）",
            ),
            horizontal=True,
            key="stocktake_scan_confirm_mode",
        )
        _st_batch = _st_mode.startswith("複数")

        if "stocktake_cand_page" not in st.session_state:
            st.session_state.stocktake_cand_page = 0
        n_total = len(cands)
        n_pages = max(1, (n_total + STOCKTAKE_CAND_PAGE_SIZE - 1) // STOCKTAKE_CAND_PAGE_SIZE)
        page_idx = int(st.session_state.stocktake_cand_page)
        if page_idx >= n_pages:
            page_idx = n_pages - 1
            st.session_state.stocktake_cand_page = page_idx
        if page_idx < 0:
            page_idx = 0
            st.session_state.stocktake_cand_page = 0
        start_i = page_idx * STOCKTAKE_CAND_PAGE_SIZE
        end_i = min(n_total, start_i + STOCKTAKE_CAND_PAGE_SIZE)
        page_slice = cands[start_i:end_i]
        _mid_opts = [str(h.get("management_id") or "").strip() for h in cands]
        _mid_opts = [m for m in _mid_opts if m]
        _mid_label: dict[str, str] = {}
        for h in cands:
            _m = str(h.get("management_id") or "").strip()
            if not _m:
                continue
            _pn = str(h.get("product_name") or "—").strip()
            if len(_pn) > 36:
                _pn = _pn[:33] + "…"
            _mid_label[_m] = f"{_m} ／ {_pn}"

        if _st_batch:
            st.caption(
                "下の一覧で **任意の管理IDを複数選択** するか、**すべて選択**／**このページだけ選択**／**クリア** を使ってから一括確定してください。"
            )
            ba1, ba2, ba3, ba4 = st.columns(4)
            with ba1:
                if st.button("すべての候補を選択", key="stocktake_sel_all_cands"):
                    st.session_state[_PENDING_STOCKTAKE_MULTI_DONE_MIDS] = list(_mid_opts)
                    st.rerun()
            with ba2:
                _page_mids = [
                    str(h.get("management_id") or "").strip()
                    for h in page_slice
                    if str(h.get("management_id") or "").strip()
                ]
                if st.button(
                    "このページの候補をすべて選択",
                    key="stocktake_sel_page_cands",
                    disabled=not _page_mids,
                ):
                    _cur = set(st.session_state.get("stocktake_multi_done_mids") or [])
                    _cur.update(_page_mids)
                    st.session_state[_PENDING_STOCKTAKE_MULTI_DONE_MIDS] = sorted(_cur)
                    st.rerun()
            with ba3:
                if st.button("選択をクリア", key="stocktake_clr_multi_sel"):
                    st.session_state[_PENDING_STOCKTAKE_MULTI_DONE_MIDS] = None
                    st.rerun()
            with ba4:
                st.caption(f"候補 **{n_total}** 件中、選択中 **{len(st.session_state.get('stocktake_multi_done_mids') or [])}** 件")
            _page_pick_opts = [
                str(h.get("management_id") or "").strip()
                for h in page_slice
                if str(h.get("management_id") or "").strip()
            ]
            st.multiselect(
                "このページ内の候補を任意選択（追加用）",
                options=_page_pick_opts,
                format_func=lambda m: _mid_label.get(m, m),
                key="stocktake_multi_page_pick",
            )
            if st.button(
                "このページの任意選択を追加",
                key="stocktake_sel_page_partial_add",
                disabled=not bool(st.session_state.get("stocktake_multi_page_pick")),
            ):
                _cur = set(st.session_state.get("stocktake_multi_done_mids") or [])
                _cur.update(
                    [
                        str(x).strip()
                        for x in st.session_state.get("stocktake_multi_page_pick", [])
                        if str(x).strip()
                    ]
                )
                st.session_state[_PENDING_STOCKTAKE_MULTI_DONE_MIDS] = sorted(_cur)
                st.rerun()
            st.multiselect(
                "一括で棚卸確定する管理ID（任意に追加・解除）",
                options=_mid_opts,
                format_func=lambda m: _mid_label.get(m, m),
                key="stocktake_multi_done_mids",
            )
        else:
            sel_cur = str(st.session_state.get("_stocktake_selected_mid") or "").strip()
            if sel_cur:
                st.info(f"選択中の管理ID: **{sel_cur}**")

        if n_pages > 1:
            p1, p2, p3 = st.columns([1, 3, 1])
            with p1:
                if st.button(
                    "◀ 前の候補",
                    disabled=page_idx <= 0,
                    key="stocktake_cand_prev",
                ):
                    st.session_state.stocktake_cand_page = max(0, page_idx - 1)
                    st.rerun()
            with p2:
                st.caption(
                    f"**ページ {page_idx + 1} / {n_pages}**（{n_total} 件中 **{start_i + 1}〜{end_i}** 件）"
                )
            with p3:
                if st.button(
                    "次の候補 ▶",
                    disabled=page_idx >= n_pages - 1,
                    key="stocktake_cand_next",
                ):
                    st.session_state.stocktake_cand_page = min(
                        n_pages - 1, page_idx + 1
                    )
                    st.rerun()

        for j, hit in enumerate(page_slice):
            mid = str(hit.get("management_id") or "").strip()
            if not mid:
                continue
            gidx = start_i + j
            with st.container(border=True):
                h1, h2 = st.columns([1, 2])
                with h1:
                    _render_inventory_gallery_thumbnail(
                        str(hit.get("image_url") or ""),
                        width=200,
                        sold=False,
                    )
                with h2:
                    st.markdown(f"**管理ID:** `{mid}`")
                    st.write(f"**商品名:** {hit.get('product_name') or '—'}")
                    st.write(f"**仕入先:** {hit.get('supplier') or '—'}")
                    st.write(
                        f"**前回の棚卸日:** {hit.get('last_stocktake') or '—（未入力）'}"
                    )
                    st.caption(
                        f"AI 確信度: {float(hit.get('confidence') or 0):.2f}（参考）"
                    )
                    if not _st_batch and st.button(
                        "この候補を選ぶ",
                        key=f"stocktake_pick_{gidx}_{mid}",
                        type="secondary",
                    ):
                        st.session_state["_stocktake_selected_mid"] = mid
                        st.rerun()

        if _st_batch:
            _picked = list(st.session_state.get("stocktake_multi_done_mids") or [])
            _picked = [str(x).strip() for x in _picked if str(x).strip()]
            if st.button(
                "選択した候補をまとめて棚卸確定（棚卸日を本日・JST）",
                type="primary",
                key="stocktake_confirm_multi",
                disabled=len(_picked) < 1,
            ):
                try:
                    with st.spinner("台帳を更新しています…"):
                        n_ok, skips = apply_last_stocktake_jst_for_management_ids(_picked)
                except Exception as e:
                    st.error(str(e))
                else:
                    st.session_state.pop("_stocktake_scan_candidates", None)
                    st.session_state.pop("_stocktake_selected_mid", None)
                    st.session_state.pop("stocktake_multi_done_mids", None)
                    st.session_state.pop("stocktake_cand_page", None)
                    st.success(
                        f"**{n_ok}** 件の棚卸日を本日（JST）に更新しました。"
                    )
                    if skips:
                        st.caption(
                            "スキップ: " + "；".join(skips[:8])
                            + (" …" if len(skips) > 8 else "")
                        )
                    st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                    st.rerun()
        elif st.button(
            "棚卸を確定（棚卸日を本日・JST に更新）",
            type="primary",
            key="stocktake_confirm_selected",
            disabled=not str(st.session_state.get("_stocktake_selected_mid") or "").strip(),
        ):
            mid = str(st.session_state.get("_stocktake_selected_mid") or "").strip()
            try:
                with st.spinner("台帳を更新しています…"):
                    apply_last_stocktake_jst_for_management_ids([mid])
            except Exception as e:
                st.error(str(e))
            else:
                st.session_state.pop("_stocktake_scan_candidates", None)
                st.session_state.pop("_stocktake_selected_mid", None)
                st.session_state.pop("stocktake_multi_done_mids", None)
                st.session_state.pop("stocktake_cand_page", None)
                st.success(f"管理ID **{mid}** の棚卸日を更新しました。")
                st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                st.rerun()


def _mask_ledger_loan_datetime_nonblank(df: pd.DataFrame) -> pd.Series:
    """浮貸日時が実質入力されている行（在庫一覧・ギャラリー用）。"""
    if df.empty or COL_LOAN_DATETIME not in df.columns:
        return pd.Series(False, index=df.index, dtype=bool)
    s = df[COL_LOAN_DATETIME].astype(str).str.strip()
    low = s.str.lower()
    blank = s.eq("") | low.isin(("nan", "none", "<na>", "nat"))
    return ~(blank.fillna(True))


def _filter_inventory_df_for_view(
    df: pd.DataFrame,
    *,
    q: str,
    suppliers: list[str],
    status_mode: str,
    stocktake_filter: str = "指定なし",
    stocktake_session_remaining: set[str] | None = None,
    loan_filter: str = "指定なし",
) -> pd.DataFrame:
    """在庫一覧の検索・フィルタ（ギャラリー／表の共通ビュー用）。"""
    out = df.copy()
    if status_mode == "在庫中":
        out = out.loc[_mask_ledger_in_stock(out)]
    elif status_mode == "販売済":
        if COL_STOCK_STATUS in out.columns:
            out = out.loc[
                out[COL_STOCK_STATUS].astype(str).str.strip().map(_normalize_stock_status)
                == STATUS_SOLD
            ]
    elif status_mode == "対象外":
        if COL_STOCK_STATUS in out.columns:
            out = out.loc[
                out[COL_STOCK_STATUS].astype(str).str.strip().map(_normalize_stock_status)
                == STATUS_EXCLUDED
            ]
    if stocktake_filter == "台帳で棚卸日が未入力の在庫中のみ":
        out = out.loc[_mask_ledger_stocktake_unverified(out)]
    elif (
        stocktake_filter == "今回の作業でまだ未確認（在庫中）"
        and stocktake_session_remaining is not None
        and COL_MANAGEMENT_ID in out.columns
    ):
        rem = stocktake_session_remaining
        m_rem = out[COL_MANAGEMENT_ID].astype(str).str.strip().isin(rem)
        out = out.loc[m_rem & _mask_ledger_in_stock(out)]
    if loan_filter == "浮貸あり":
        if COL_LOAN_DATETIME in out.columns:
            out = out.loc[_mask_ledger_loan_datetime_nonblank(out)]
        else:
            out = out.iloc[0:0].copy()
    elif loan_filter == "浮貸なし":
        if COL_LOAN_DATETIME in out.columns:
            out = out.loc[~_mask_ledger_loan_datetime_nonblank(out)]
    if suppliers and COL_SUPPLIER in out.columns:
        sup_m = out[COL_SUPPLIER].astype(str).str.strip().isin(set(suppliers))
        out = out.loc[sup_m]
    qt = (q or "").strip()
    if qt and not out.empty:
        qf = qt.casefold()
        m = pd.Series(False, index=out.index)
        for col in (COL_NAME, COL_MANAGEMENT_ID, COL_MEMO):
            if col in out.columns:
                m = m | out[col].astype(str).str.casefold().str.contains(
                    qf, na=False, regex=False
                )
        out = out.loc[m]
    return out.reset_index(drop=True)


@st.dialog("在庫の詳細")
def _inventory_gallery_detail_dialog(row_dict: dict[str, Any]) -> None:
    """ギャラリーから開く行の全列表示。"""
    for k in EXPECTED_HEADERS:
        if k not in row_dict:
            continue
        v = row_dict.get(k)
        if k in (COL_IMAGE_URL, COL_SALE_IMAGE_URL):
            st.markdown(f"**{k}**")
            _iu = str(v or "").strip()
            if _iu.startswith("http://") or _iu.startswith("https://"):
                _render_http_product_image_from_url(
                    _iu, pixel_width=None, use_container_width=True
                )
            else:
                st.write("—")
            continue
        st.markdown(f"**{k}**")
        st.write(str(v) if v is not None and str(v).strip() != "" else "—")


def _render_sales_management_tab(
    uploaded,
    df_ledger_hint: pd.DataFrame | None,
) -> None:
    """販売管理タブ: 出庫区分に応じて管理ID・実売または浮貸日時・返品・戻入を記録する。"""
    st.markdown("##### 販売管理")
    st.caption(
        "区分に応じて **在庫中** の管理IDを指定して反映します。"
        "**出庫（除外）** では結果を **販売済** または **対象外** にできます。"
    )
    with st.expander("使い方", expanded=False):
        st.markdown(
            "- **出庫（販売）** / **出庫（除外）**: **在庫中** の行を **販売済** または（除外のみ）**対象外** に更新します。\n"
            "- **出庫（浮貸）**: 在庫中のまま浮貸日時を記録するか、販売済へ更新できます。\n"
            "- **出庫（返品）**: **販売済** の行を **在庫中** に戻します（写真照合・台帳補助は販売済行が対象）。\n"
            "- **出庫（戻入）**: **在庫中** の行に浮貸日時と出庫（戻入）を記録するか、販売済へ更新できます（浮貸と同様にステータスを選べます）。\n"
            "- 写真照合は上部の共通アップローダ画像を使用します。"
        )
    outbound_kind = st.radio(
        "出庫区分",
        (
            "出庫（販売）",
            OUTBOUND_KIND_EXCLUDE,
            "出庫（浮貸）",
            "出庫（返品）",
            "出庫（戻入）",
        ),
        horizontal=True,
        key="sales_tab_outbound_kind",
    )
    loan_target_status: str | None = None
    receipt_target_status: str | None = None
    if outbound_kind == "出庫（浮貸）":
        loan_target_status = st.radio(
            "出庫（浮貸）の結果ステータス",
            (STATUS_IN_STOCK, STATUS_SOLD),
            horizontal=True,
            key="sales_tab_loan_stock_status",
        )
    elif outbound_kind == "出庫（戻入）":
        receipt_target_status = st.radio(
            "出庫（戻入）の結果ステータス",
            (STATUS_IN_STOCK, STATUS_SOLD),
            horizontal=True,
            key="sales_tab_receipt_stock_status",
        )
    disposal_target_status: str | None = None
    if _is_outbound_exclude_kind(outbound_kind):
        disposal_target_status = st.radio(
            "出庫（除外）の結果ステータス",
            (STATUS_SOLD, STATUS_EXCLUDED),
            format_func=lambda x: (
                "販売済（実売で計上）" if x == STATUS_SOLD else "対象外"
            ),
            horizontal=True,
            key="sales_tab_disposal_stock_status",
        )
    _loan_keep_stock = (
        outbound_kind == "出庫（浮貸）" and loan_target_status == STATUS_IN_STOCK
    )
    _loan_as_sale = (
        outbound_kind == "出庫（浮貸）" and loan_target_status == STATUS_SOLD
    )
    _receipt_keep_stock = (
        outbound_kind == "出庫（戻入）" and receipt_target_status == STATUS_IN_STOCK
    )
    _receipt_as_sale = (
        outbound_kind == "出庫（戻入）" and receipt_target_status == STATUS_SOLD
    )
    _plain_sale = outbound_kind == "出庫（販売）" or (
        _is_outbound_exclude_kind(outbound_kind)
        and disposal_target_status == STATUS_SOLD
    )
    _disposal_excluded = (
        _is_outbound_exclude_kind(outbound_kind)
        and disposal_target_status == STATUS_EXCLUDED
    )
    _return_flow = outbound_kind == "出庫（返品）"
    _receipt_flow = outbound_kind == "出庫（戻入）"
    _loan_keep_eff = _loan_keep_stock or _receipt_keep_stock
    do_match = st.button(
        "AIで写真と照合",
        type="primary",
        disabled=uploaded is None,
        key="sales_tab_photo_match_btn",
    )
    if "sales_assist_visible" not in st.session_state:
        st.session_state.sales_assist_visible = False
    if st.button(
        "台帳からの入力補助を表示"
        if not st.session_state.sales_assist_visible
        else "台帳からの入力補助を非表示",
        key="sales_assist_toggle_btn",
    ):
        st.session_state.sales_assist_visible = not st.session_state.sales_assist_visible
        st.rerun()
    if st.button("入力をクリア", key="sales_tab_clear_fields_btn"):
        st.session_state.field_sale_source_mgmt_id = ""
        st.session_state.field_actual_sale_excl = 0
        st.session_state.sales_tab_memo = ""
        st.session_state.sales_tab_loan_datetime_manual = ""
        st.session_state.sale_pick_source_id = LEDGER_PICK_PLACEHOLDER
        st.session_state.sales_hint_filter_product_name = ""
        st.session_state.sales_hint_filter_supplier = ""
        st.session_state.sales_hint_filter_inventory_category = ""
        st.session_state.sales_hint_filter_management_id = ""
        st.session_state.sales_ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
        st.session_state.sales_ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
        st.session_state.sales_ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER
        st.session_state.sales_ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
        st.session_state.sales_assist_buf_product_name = ""
        st.session_state.sales_assist_buf_supplier = ""
        st.session_state.sales_assist_buf_inventory_category = ""
        st.session_state.sales_assist_buf_management_id = ""
        st.session_state.sales_pick_mode = "1件選択"
        st.session_state.pop("sales_assist_quick_candidates", None)
        st.session_state.pop("sales_assist_last_n_matching_mids", None)
        st.session_state.pop("_sale_link_management_id", None)
        st.session_state.pop("_sale_link_warn", None)
        st.session_state.pop("_sales_photo_match_card_hits", None)
        st.session_state.pop("sales_assist_page_partial_pick", None)
        st.session_state.pop("sales_assist_cand_page", None)
        st.rerun()

    if do_match and uploaded is not None:
        inv_ctx_sale = ""
        if df_ledger_hint is not None and not df_ledger_hint.empty:
            if _return_flow:
                n_ctx = int(_mask_ledger_sold(df_ledger_hint).sum())
                max_lines_sale_ctx = min(
                    SALES_AI_CONTEXT_MAX_LINES,
                    max(400, n_ctx + 20),
                )
                inv_ctx_sale = _build_gemini_inventory_context(
                    df_ledger_hint,
                    only_sold=True,
                    max_lines=max_lines_sale_ctx,
                )
            elif _receipt_flow:
                n_ctx = int(
                    _mask_ledger_in_stock_outbound_float_loan(df_ledger_hint).sum()
                )
                max_lines_sale_ctx = min(
                    SALES_AI_CONTEXT_MAX_LINES,
                    max(400, n_ctx + 20),
                )
                inv_ctx_sale = _build_gemini_inventory_context(
                    df_ledger_hint,
                    only_in_stock=True,
                    max_lines=max_lines_sale_ctx,
                    sale_outbound_type_eq="出庫（浮貸）",
                )
            else:
                n_in_stock_ctx = int(_mask_ledger_in_stock(df_ledger_hint).sum())
                max_lines_sale_ctx = min(
                    SALES_AI_CONTEXT_MAX_LINES,
                    max(400, n_in_stock_ctx + 20),
                )
                inv_ctx_sale = _build_gemini_inventory_context(
                    df_ledger_hint,
                    only_in_stock=True,
                    max_lines=max_lines_sale_ctx,
                )
        if do_match and not (inv_ctx_sale or "").strip():
            if _return_flow:
                st.warning(
                    "返品の写真照合には、台帳に **販売済** の行が少なくとも1行必要です。"
                    "該当がない場合や **管理ID** が空の行しかない場合はリストを作れません。"
                    "ページ先頭の台帳読み込みエラーが出ていないかも確認してください。"
                )
            elif _receipt_flow:
                st.warning(
                    "戻入の写真照合には、台帳に **在庫中かつ出庫種別が出庫（浮貸）** の行が少なくとも1行必要です。"
                    "該当がない場合や **管理ID** が空の行しかない場合はリストを作れません。"
                    "ページ先頭の台帳読み込みエラーが出ていないかも確認してください。"
                )
            else:
                st.warning(
                    "販売元の写真照合には、台帳に **在庫中** の行が少なくとも1行必要です。"
                    "在庫がすべて販売済の場合や、**管理ID** が空の行しかない場合はリストを作れません。"
                    "ページ先頭の台帳読み込みエラーが出ていないかも確認してください。"
                )
        elif do_match:
            _spin_msg = (
                "画像を解析して返品対象を照合しています…"
                if _return_flow
                else (
                    "画像を解析して戻入対象を照合しています…"
                    if _receipt_flow
                    else "画像を解析して販売元を照合しています…"
                )
            )
            with st.spinner(_spin_msg):
                try:
                    raw_text = analyze_image_with_gemini(
                        uploaded,
                        inventory_context=inv_ctx_sale or None,
                        prompt_mode=(
                            "sale_link_sold"
                            if _return_flow
                            else (
                                "sale_link_float_loan"
                                if _receipt_flow
                                else "sale_link"
                            )
                        ),
                    )
                    result = _parse_json_from_model(raw_text or "")
                    _apply_gemini_sale_link_to_session(
                        result,
                        df_ledger_hint,
                        fill_product_preview_fields=False,
                        restrict_to_sold=_return_flow,
                        restrict_to_float_loan_outbound=_receipt_flow,
                    )
                    _pm_cards = _sales_photo_match_card_hits_from_result(
                        result,
                        df_ledger_hint,
                        sold_rows_only=_return_flow,
                        float_loan_outbound_only=_receipt_flow,
                    )
                    if _pm_cards:
                        st.session_state["_sales_photo_match_card_hits"] = _pm_cards
                    else:
                        st.session_state.pop("_sales_photo_match_card_hits", None)
                    st.success("照合が完了しました。管理IDを確認してください。")
                except Exception as e:
                    st.session_state.pop("_sales_photo_match_card_hits", None)
                    st.warning(
                        "現在混み合っているか、無料枠の上限に達している可能性があります。"
                        "1分ほど待ってから再試行してください。"
                    )
                    st.caption(f"詳細: {e}")

    _swarn = st.session_state.pop("_sale_link_warn", None)
    if _swarn:
        st.warning(_swarn)
    _sale_link_flash = st.session_state.pop("_sale_link_management_id", None)
    if _sale_link_flash:
        st.info(
            f"販売元として **{_sale_link_flash}** をセットしました。"
            "内容を確認してから確定してください。"
        )

    _sale_id_opts: list[str] = []
    if df_ledger_hint is not None and not df_ledger_hint.empty:
        if _return_flow:
            _sale_id_opts = _ledger_sold_management_ids(df_ledger_hint)
        elif _receipt_flow:
            _sale_id_opts = _ledger_in_stock_outbound_float_loan_management_ids(
                df_ledger_hint
            )
        else:
            _sale_id_opts = _ledger_in_stock_management_ids(df_ledger_hint)
    _sale_pick_mode = st.radio(
        "販売対象の選択",
        ("1件選択", "複数選択（一括反映）"),
        horizontal=True,
        key="sales_pick_mode",
    )
    if not _sale_pick_mode.startswith("複数"):
        if _sale_id_opts:
            st.selectbox(
                (
                    "販売済の管理ID（すぐ選ぶ）"
                    if _return_flow
                    else (
                        "出庫（浮貸）の在庫の管理ID（すぐ選ぶ）"
                        if _receipt_flow
                        else "在庫中の管理ID（すぐ選ぶ）"
                    )
                ),
                options=[LEDGER_PICK_PLACEHOLDER] + _sale_id_opts,
                key="sale_pick_source_id",
                on_change=_on_sale_pick_source_id,
            )
    _spm_hits = st.session_state.get("_sales_photo_match_card_hits")
    if isinstance(_spm_hits, list) and _spm_hits:
        st.markdown("##### 写真照合の近い候補（カード）")
        st.caption(
            (
                "AI の商品名・仕入先・管理IDと表記が近い **販売済** の行です。"
                "**この候補を販売元にする** でその管理IDへ切り替えられます。"
                if _return_flow
                else (
                    "AI の商品名・仕入先・管理IDと表記が近い、**在庫中かつ出庫種別が出庫（浮貸）** の行です。"
                    "**この候補を販売元にする** でその管理IDへ切り替えられます。"
                    if _receipt_flow
                    else (
                        "AI の商品名・仕入先・管理IDと表記が近い **在庫中** の行です。"
                        "**この候補を販売元にする** でその管理IDへ切り替えられます。"
                    )
                )
            )
        )
        _spm_mids = [str(h.get("management_id") or "").strip() for h in _spm_hits]
        _spm_mids = [m for m in _spm_mids if m]
        _spm_page_key = "_sales_photo_match_cards_page"
        _spm_page_size = 5
        _spm_total = len(_spm_hits)
        _spm_pages = max(1, (_spm_total + _spm_page_size - 1) // _spm_page_size)
        _spm_cur = int(st.session_state.get(_spm_page_key, 0) or 0)
        _spm_cur = max(0, min(_spm_pages - 1, _spm_cur))
        _spm_start = _spm_cur * _spm_page_size
        _spm_end = min(_spm_total, _spm_start + _spm_page_size)
        _spm_page_mids = [
            str(h.get("management_id") or "").strip()
            for h in _spm_hits[_spm_start:_spm_end]
            if str(h.get("management_id") or "").strip()
        ]
        if _sale_pick_mode.startswith("複数"):
            c1, c2, c3 = st.columns(3)
            with c1:
                if st.button(
                    "候補全部を選択",
                    key="sales_cand_pick_all",
                    disabled=not _spm_mids,
                ):
                    st.session_state.field_sale_source_mgmt_id = ", ".join(_spm_mids)
                    st.rerun()
            with c2:
                if st.button(
                    "表示ページの候補をすべて選択",
                    key="sales_cand_pick_page_all",
                    disabled=not _spm_page_mids,
                ):
                    _cur = set(
                        _split_management_ids_from_field(
                            str(st.session_state.get("field_sale_source_mgmt_id", "") or "")
                        )
                    )
                    _cur.update(_spm_page_mids)
                    _new = sorted(x for x in _cur if str(x).strip())
                    st.session_state.field_sale_source_mgmt_id = ", ".join(_new)
                    st.rerun()
            with c3:
                if st.button("候補選択をクリア", key="sales_cand_pick_clear"):
                    st.session_state.field_sale_source_mgmt_id = ""
                    st.rerun()
            st.multiselect(
                "表示ページ内の候補を任意選択（追加用）",
                options=_spm_page_mids,
                key="sales_page_partial_pick",
            )
            if st.button(
                "表示ページの任意選択を追加",
                key="sales_cand_pick_page_partial_add",
                disabled=not bool(st.session_state.get("sales_page_partial_pick")),
            ):
                _cur = set(
                    _split_management_ids_from_field(
                        str(st.session_state.get("field_sale_source_mgmt_id", "") or "")
                    )
                )
                _cur.update(
                    [
                        str(x).strip()
                        for x in st.session_state.get("sales_page_partial_pick", [])
                        if str(x).strip()
                    ]
                )
                _new = sorted(x for x in _cur if str(x).strip())
                st.session_state.field_sale_source_mgmt_id = ", ".join(_new)
                st.rerun()
        _render_mid_pick_candidate_cards(
            _spm_hits,
            widget_key_namespace="sales_photo_match_cards",
            sold=_return_flow,
            pick_mode="sale",
        )

    if (
        st.session_state.sales_assist_visible
        and df_ledger_hint is not None
        and not df_ledger_hint.empty
    ):
        _render_ledger_pick_assist_three_columns(
            df_ledger_hint,
            key_prefix="sales_",
            sales_restrict_to_sold=_return_flow,
            sales_restrict_to_float_loan_outbound=_receipt_flow,
            body_caption=(
                (
                    "仕入タブと同様に、商品名・仕入先・在庫カテゴリー・**管理ID** を選べます。"
                    "**販売済** に限定した上でフィルタを **AND** した結果がちょうど1件のときのみ、自動で "
                    "**販売する管理ID** に反映します。"
                )
                if _return_flow
                else (
                    (
                        "仕入タブと同様に、商品名・仕入先・在庫カテゴリー・**管理ID** を、文字での絞り込みまたはプルダウンから選べます。"
                        "**在庫中かつ出庫種別が出庫（浮貸）** に限定した上でフィルタを **AND** した結果がちょうど1件のときのみ、自動で "
                        "**販売する管理ID** に反映します。その他は下のカードから選ぶか手入力してください。"
                    )
                    if _receipt_flow
                    else (
                        "仕入タブと同様に、商品名・仕入先・在庫カテゴリー・**管理ID** を、文字での絞り込みまたはプルダウンから選べます。"
                        "**在庫中** に限定した上でフィルタを **AND** した結果がちょうど1件のときのみ、自動で "
                        "**販売する管理ID** に反映します。その他は下のカードから選ぶか手入力してください。"
                    )
                )
            ),
            on_pick_product_name=_on_sales_assist_pick_product_name,
            on_pick_supplier=_on_sales_assist_pick_supplier,
            on_pick_inventory_category=_on_sales_assist_pick_inventory_category,
            on_pick_management_id=_on_sales_assist_pick_management_id,
        )
        nm = int(st.session_state.get("sales_assist_last_n_matching_mids", 0) or 0)
        if nm > 1:
            _scope_lbl = (
                "販売済"
                if _return_flow
                else (
                    "出庫種別が出庫（浮貸）の在庫中"
                    if _receipt_flow
                    else "在庫中"
                )
            )
            st.info(
                f"補助条件に一致する **{_scope_lbl}** が **{nm}** 件あります。"
                "一致が1件だけのときだけ **販売する管理ID** が自動入力されます。それ以外は一覧か手入力で特定してください。"
            )

        _refresh_sales_assist_quick_candidates(df_ledger_hint)
        _sac = st.session_state.get("sales_assist_quick_candidates")
        if (
            isinstance(_sac, pd.DataFrame)
            and not _sac.empty
            and df_ledger_hint is not None
        ):
            st.markdown("##### 近い候補（補助で選んだ項目・入力から照合・カード）")
            st.caption(
                (
                    "入力補助で確定した項目と表記が近い **販売済** を表示します。"
                    "上の **販売対象の選択** が **複数選択** のときは、一覧・ボタンで **販売する管理ID** に追加できます。"
                    "カードの **この候補を販売元にする** は選択に追加（既存のIDは残します）。"
                )
                if _return_flow
                else (
                    (
                        "入力補助で確定した項目と表記が近い、**在庫中かつ出庫種別が出庫（浮貸）** の行を表示します。"
                        "上の **販売対象の選択** が **複数選択** のときは、AI 写真照合と同様に一覧・ボタンで **販売する管理ID** に追加できます。"
                        "カードの **この候補を販売元にする** は選択に追加（既存のIDは残します）。"
                    )
                    if _receipt_flow
                    else (
                        "入力補助で確定した項目と表記が近い **在庫中** を表示します。"
                        "上の **販売対象の選択** が **複数選択** のときは、AI 写真照合と同様に一覧・ボタンで **販売する管理ID** に追加できます。"
                        "カードの **この候補を販売元にする** は選択に追加（既存のIDは残します）。"
                    )
                )
            )
            _s_hits = [_sale_card_hit_from_series(row) for _, row in _sac.iterrows()]
            _sale_mbatch = str(_sale_pick_mode or "").startswith("複数")
            _mid_opts_s: list[str] = []
            _seen_s: set[str] = set()
            for _h in _s_hits:
                _m = str(_h.get("management_id") or "").strip()
                if _m and _m not in _seen_s:
                    _seen_s.add(_m)
                    _mid_opts_s.append(_m)
            _mid_opts_s.sort(key=_management_id_sort_key)
            _mid_label_s: dict[str, str] = {}
            for _h in _s_hits:
                _m = str(_h.get("management_id") or "").strip()
                if not _m:
                    continue
                _pn = str(_h.get("product_name") or "—").strip()
                if len(_pn) > 36:
                    _pn = _pn[:33] + "…"
                _mid_label_s[_m] = f"{_m} ／ {_pn}"
            if "sales_assist_cand_page" not in st.session_state:
                st.session_state.sales_assist_cand_page = 0
            n_tot_s = len(_mid_opts_s)
            n_pg_s = max(
                1,
                (n_tot_s + STOCKTAKE_CAND_PAGE_SIZE - 1) // STOCKTAKE_CAND_PAGE_SIZE,
            )
            pg_s = int(st.session_state.sales_assist_cand_page)
            pg_s = max(0, min(n_pg_s - 1, pg_s))
            st.session_state.sales_assist_cand_page = pg_s
            si_s = pg_s * STOCKTAKE_CAND_PAGE_SIZE
            ei_s = min(n_tot_s, si_s + STOCKTAKE_CAND_PAGE_SIZE)
            page_mids_s = _mid_opts_s[si_s:ei_s]
            _by_mid_s = {str(h.get("management_id") or "").strip(): h for h in _s_hits}
            page_hits_s = [_by_mid_s[m] for m in page_mids_s if m in _by_mid_s]
            if _sale_mbatch:
                c1, c2, c3 = st.columns(3)
                with c1:
                    if st.button(
                        "補助候補をすべて選択",
                        key="sales_assist_pick_all",
                        disabled=not _mid_opts_s,
                    ):
                        st.session_state.field_sale_source_mgmt_id = ", ".join(
                            _mid_opts_s
                        )
                        st.rerun()
                with c2:
                    if st.button(
                        "このページの候補をすべて選択",
                        key="sales_assist_pick_page_all",
                        disabled=not page_mids_s,
                    ):
                        _cur = set(
                            _split_management_ids_from_field(
                                str(
                                    st.session_state.get(
                                        "field_sale_source_mgmt_id", ""
                                    )
                                    or ""
                                )
                            )
                        )
                        _cur.update(page_mids_s)
                        st.session_state.field_sale_source_mgmt_id = ", ".join(
                            sorted(_cur, key=_management_id_sort_key)
                        )
                        st.rerun()
                with c3:
                    if st.button("補助の選択をクリア", key="sales_assist_pick_clear"):
                        st.session_state.field_sale_source_mgmt_id = ""
                        st.rerun()
                st.multiselect(
                    "このページ内を任意選択（追加用）",
                    options=page_mids_s,
                    format_func=lambda m: _mid_label_s.get(m, m),
                    key="sales_assist_page_partial_pick",
                )
                if st.button(
                    "このページの任意選択を追加",
                    key="sales_assist_page_partial_add",
                    disabled=not bool(
                        st.session_state.get("sales_assist_page_partial_pick")
                    ),
                ):
                    _cur = set(
                        _split_management_ids_from_field(
                            str(
                                st.session_state.get("field_sale_source_mgmt_id", "")
                                or ""
                            )
                        )
                    )
                    _cur.update(
                        [
                            str(x).strip()
                            for x in st.session_state.get(
                                "sales_assist_page_partial_pick", []
                            )
                            if str(x).strip()
                        ]
                    )
                    st.session_state.field_sale_source_mgmt_id = ", ".join(
                        sorted(_cur, key=_management_id_sort_key)
                    )
                    st.rerun()
            if n_pg_s > 1:
                s1, s2, s3 = st.columns([1, 3, 1])
                with s1:
                    if st.button(
                        "◀ 前へ",
                        disabled=pg_s <= 0,
                        key="sales_assist_cand_prev",
                    ):
                        st.session_state.sales_assist_cand_page = max(0, pg_s - 1)
                        st.rerun()
                with s2:
                    st.caption(
                        f"**ページ {pg_s + 1} / {n_pg_s}**（{n_tot_s} 件中 **{si_s + 1}〜{ei_s}** 件）"
                    )
                with s3:
                    if st.button(
                        "次へ ▶",
                        disabled=pg_s >= n_pg_s - 1,
                        key="sales_assist_cand_next",
                    ):
                        st.session_state.sales_assist_cand_page = min(
                            n_pg_s - 1, pg_s + 1
                        )
                        st.rerun()
            if _sale_mbatch:
                _cap_s = f"全 **{n_tot_s}** 件中 **{si_s + 1}〜{ei_s}** 件（カード）"
                _render_mid_pick_candidate_cards(
                    page_hits_s,
                    widget_key_namespace=f"sales_assist_pg_{pg_s}",
                    sold=_return_flow,
                    pick_mode="sale",
                    pager="hidden",
                    caption_override=_cap_s,
                    sale_merge_selection=True,
                )
            else:
                _render_mid_pick_candidate_cards(
                    _s_hits,
                    widget_key_namespace="sales_assist_cards",
                    sold=_return_flow,
                    pick_mode="sale",
                )

    st.text_input(
        "販売する管理ID（手入力可）",
        key="field_sale_source_mgmt_id",
        placeholder="例: G00000001 または G00000001, G00000002",
        help="複数選択モードではこの欄に自動反映されます。",
    )
    if _loan_keep_eff:
        st.text_input(
            "浮貸日時（空欄＝確定ボタン押下の JST）",
            key="sales_tab_loan_datetime_manual",
            help=(
                "出庫（浮貸）または **出庫（戻入）で在庫中のまま** のとき、**浮貸日時** 列に入る値です。"
                "未入力なら確定実行の JST を記録します。"
            ),
        )
    st.number_input(
        "実売金額（税抜・1点あたり）",
        min_value=0,
        step=1,
        key="field_actual_sale_excl",
        disabled=not (_plain_sale or _loan_as_sale or _receipt_as_sale),
        help=(
            "出庫（販売）・出庫（除外）・出庫（浮貸）で **販売済** にするとき、または **出庫（戻入）で販売済** にするときに使用（0円以上）。"
            "在庫中のまま（浮貸・戻入）記録のみのときは不要です。**出庫（返品）** のときも不要です。"
        ),
    )
    memo_sales = st.text_area(
        "販売メモ（任意・台帳のメモに追記）",
        key="sales_tab_memo",
        height=80,
    )

    _ids_pv = _split_management_ids_from_field(
        str(st.session_state.get("field_sale_source_mgmt_id", "") or "")
    )
    _q = len(_ids_pv)
    _act_u = int(st.session_state.get("field_actual_sale_excl", 0))
    _pv_ok_rows: list[pd.Series] = []
    if _ids_pv and len(set(_ids_pv)) != len(_ids_pv):
        st.warning("管理IDに **重複** があります。")
    if _ids_pv and df_ledger_hint is not None:
        _pv_msgs: list[str] = []
        for _mid_one in _ids_pv:
            _tr_pv = lookup_ledger_row_by_management_id(df_ledger_hint, _mid_one)
            if _tr_pv is None:
                st.warning(f"管理ID **{_mid_one}** が台帳に見つかりません。")
                continue
            _row_st = _normalize_stock_status(str(_tr_pv.get(COL_STOCK_STATUS, "")))
            if _return_flow:
                if _row_st != STATUS_SOLD:
                    st.warning(
                        f"管理ID **{_mid_one}** は販売済ではありません（現在: {_row_st}）。"
                        "出庫（返品）は **販売済** の行のみ対象です。"
                    )
                    continue
            elif _row_st != STATUS_IN_STOCK:
                st.warning(
                    f"管理ID **{_mid_one}** は在庫中ではありません（現在: {_row_st}）。"
                )
                continue
            _pv_ok_rows.append(_tr_pv)
            _cg1 = _finite_int(_tr_pv.get(COL_PRICE_EXCL), 0)
            _pnv = str(_tr_pv.get(COL_NAME, "") or "").strip()
            _suv = str(_tr_pv.get(COL_SUPPLIER, "") or "").strip()
            _pv_msgs.append(
                f"**{_mid_one}** … {_pnv or '—'} ／ {_suv or '—'} ／ 原価税抜 ¥{_cg1:,}"
            )
        if _pv_msgs:
            _hdr = "紐付け元（販売済）" if _return_flow else "紐付け元（在庫中）"
            st.info(_hdr + ":\n" + "\n".join(_pv_msgs))
    elif _ids_pv:
        st.warning("台帳を読み込めないため、紐付け元の原価を表示できません。")

    _sale_pv_agg = (
        bool(_ids_pv)
        and len(set(_ids_pv)) == len(_ids_pv)
        and len(_pv_ok_rows) == len(_ids_pv)
        and len(_pv_ok_rows) > 0
    )
    _cogs_preview = 0
    _pl_u_gp = 0
    _tax_preview = float(CONSUMPTION_TAX_RATE)
    _plex = _pin = _aex = _ain = 0
    _gp_preview: int | None = None
    if _loan_keep_eff and _sale_pv_agg:
        _cogs_preview = sum(_finite_int(x.get(COL_PRICE_EXCL), 0) for x in _pv_ok_rows)
        _tr0 = _pv_ok_rows[0]
        _tax_preview = _infer_tax_rate_from_main_line(
            _finite_int(_tr0.get(COL_PRICE_EXCL), 0),
            _finite_int(_tr0.get(COL_PRICE_INCL), 0),
        )
        _plex = sum(_finite_int(x.get(COL_PLANNED_SALE), 0) for x in _pv_ok_rows)
        _pin = price_incl_tax(_plex, _tax_preview) if _plex > 0 else 0
        _aex = _ain = 0
        _gp_acc = 0
        _gp_any = False
        for _xr in _pv_ok_rows:
            _cgx = _finite_int(_xr.get(COL_PRICE_EXCL), 0)
            _plx_u = _finite_int(_xr.get(COL_PLANNED_SALE), 0)
            _plex1, _, _, _ = _planned_actual_line_amounts(
                1, _plx_u, 0, STATUS_IN_STOCK, _tax_preview
            )
            _gpx = _compute_gross_profit_row(_cgx, _plex1, 0, STATUS_IN_STOCK)
            if _gpx is not None:
                _gp_acc += int(_gpx)
                _gp_any = True
        _gp_preview = _gp_acc if _gp_any else None
    elif _sale_pv_agg and not _loan_keep_eff and not _return_flow and not _disposal_excluded:
        _cogs_preview = sum(_finite_int(x.get(COL_PRICE_EXCL), 0) for x in _pv_ok_rows)
        _tr0 = _pv_ok_rows[0]
        _tax_preview = _infer_tax_rate_from_main_line(
            _finite_int(_tr0.get(COL_PRICE_EXCL), 0),
            _finite_int(_tr0.get(COL_PRICE_INCL), 0),
        )
        _pl_u_gp = _finite_int(_tr0.get(COL_PLANNED_SALE), 0)
        _plex = sum(_finite_int(x.get(COL_PLANNED_SALE), 0) for x in _pv_ok_rows)
        _pin = price_incl_tax(_plex, _tax_preview) if _plex > 0 else 0
        _aex = _act_u * max(1, _q) if _act_u > 0 else 0
        _ain = price_incl_tax(_aex, _tax_preview) if _aex > 0 else 0
        _gp_acc = 0
        _gp_any = False
        for _xr in _pv_ok_rows:
            _cgx = _finite_int(_xr.get(COL_PRICE_EXCL), 0)
            _plx = _finite_int(_xr.get(COL_PLANNED_SALE), 0)
            _gpx = _compute_gross_profit_row(_cgx, _plx, _act_u, STATUS_SOLD)
            if _gpx is not None:
                _gp_acc += int(_gpx)
                _gp_any = True
        _gp_preview = _gp_acc if _gp_any else None
    elif _disposal_excluded and _sale_pv_agg:
        _cogs_preview = sum(_finite_int(x.get(COL_PRICE_EXCL), 0) for x in _pv_ok_rows)
        _plex = _pin = _aex = _ain = 0
        _gp_preview = None
    elif _return_flow and _sale_pv_agg:
        _cogs_preview = sum(_finite_int(x.get(COL_PRICE_EXCL), 0) for x in _pv_ok_rows)
        _plex = _pin = _aex = _ain = 0
        _gp_preview = None
    elif _loan_keep_eff and len(_pv_ok_rows) == 1:
        _tr_pv = _pv_ok_rows[0]
        _cogs_preview = _finite_int(_tr_pv.get(COL_PRICE_EXCL), 0)
        _pl_u_gp = _finite_int(_tr_pv.get(COL_PLANNED_SALE), 0)
        _tax_preview = _infer_tax_rate_from_main_line(
            _finite_int(_tr_pv.get(COL_PRICE_EXCL), 0),
            _finite_int(_tr_pv.get(COL_PRICE_INCL), 0),
        )
        _plex, _pin, _aex, _ain = _planned_actual_line_amounts(
            1, _pl_u_gp, 0, STATUS_IN_STOCK, _tax_preview
        )
        _gp_preview = _compute_gross_profit_row(
            _cogs_preview,
            _plex,
            0,
            STATUS_IN_STOCK,
        )
    elif len(_pv_ok_rows) == 1 and not _loan_keep_eff and not _return_flow and not _disposal_excluded:
        _tr_pv = _pv_ok_rows[0]
        _cogs_preview = _finite_int(_tr_pv.get(COL_PRICE_EXCL), 0)
        _pl_u_gp = _finite_int(_tr_pv.get(COL_PLANNED_SALE), 0)
        _tax_preview = _infer_tax_rate_from_main_line(
            _finite_int(_tr_pv.get(COL_PRICE_EXCL), 0),
            _finite_int(_tr_pv.get(COL_PRICE_INCL), 0),
        )
        _plex, _pin, _aex, _ain = _planned_actual_line_amounts(
            1, _pl_u_gp, _act_u, STATUS_SOLD, _tax_preview
        )
        _gp_preview = _compute_gross_profit_row(
            _cogs_preview,
            _plex,
            _aex,
            STATUS_SOLD,
        )

    pm1, pm2, pm3, pm4, pm5 = st.columns(5)
    with pm1:
        _cogs_lbl = (
            "原価（税抜・合計）" if _sale_pv_agg and len(_ids_pv) > 1 else "原価（税抜・参考）"
        )
        if _pv_ok_rows:
            st.metric(_cogs_lbl, f"¥{_cogs_preview:,}")
        else:
            st.metric(_cogs_lbl, "—")
    with pm2:
        st.metric(
            "販売予定（税抜・行計）",
            "—" if _plex <= 0 else f"¥{_plex:,}",
        )
    with pm3:
        st.metric(
            "販売予定（税込・総額）",
            "—" if _pin <= 0 else f"¥{_pin:,}",
        )
    with pm4:
        st.metric(
            "実売（税抜・行計）",
            "—" if _aex <= 0 else f"¥{_aex:,}",
        )
    with pm5:
        st.metric(
            "実売（税込・総額）",
            "—" if _ain <= 0 else f"¥{_ain:,}",
        )
    pm6, _, _ = st.columns([1, 1, 3])
    with pm6:
        st.metric(
            "粗利（税抜・プレビュー）",
            "—" if _gp_preview is None else f"¥{int(_gp_preview):,}",
        )
    if _loan_keep_stock:
        st.caption(
            f"在庫中のまま出庫（浮貸）を確定すると、各行の **{COL_LOAN_DATETIME}** に日時が入ります（手入力が空なら確定の JST）。"
        )
    elif _receipt_keep_stock:
        st.caption(
            f"在庫中のまま出庫（戻入）を確定すると、各行の **{COL_LOAN_DATETIME}** に日時が入ります（手入力が空なら確定の JST）。"
        )

    _confirm_lbl = (
        "返品を確定（販売済→在庫中・新規行なし）"
        if _return_flow
        else (
            "戻入を確定（在庫中のまま・浮貸日時のみ）"
            if _receipt_keep_stock
            else (
                "戻入を確定（販売済へ・実売で記録）"
                if _receipt_as_sale
                else (
                    "浮貸を確定（在庫中のまま・浮貸日時のみ）"
                    if _loan_keep_stock
                    else (
                        "除外を確定（対象外・実売クリア・新規行なし）"
                        if _disposal_excluded
                        else (
                            "除外を確定（在庫行のみ更新・新規行なし）"
                            if _is_outbound_exclude_kind(outbound_kind)
                            else "販売を確定（在庫行のみ更新・新規行なし）"
                        )
                    )
                )
            )
        )
    )
    confirm_sale = st.button(_confirm_lbl, type="primary", key="sales_tab_confirm_btn")

    if confirm_sale:
        _sale_src_save = str(
            st.session_state.get("field_sale_source_mgmt_id", "") or ""
        ).strip()
        _act_ex2 = int(st.session_state.get("field_actual_sale_excl", 0))
        _ids_sale_val = _split_management_ids_from_field(_sale_src_save)
        memo_s = (memo_sales or "").strip()
        validation_ok = True
        _need_actual = _plain_sale or _loan_as_sale or _receipt_as_sale
        if not _sale_src_save:
            st.error("**販売する管理ID** の入力が必須です。")
            validation_ok = False
        elif _need_actual and _act_ex2 < 0:
            st.error("**実売金額（税抜）** を0円以上で入力してください。")
            validation_ok = False
        elif df_ledger_hint is None:
            st.error("台帳を読み込めないため、反映できません。")
            validation_ok = False
        elif len(set(_ids_sale_val)) != len(_ids_sale_val):
            st.error("管理IDに **重複** があります。")
            validation_ok = False
        else:
            for _sid_v in _ids_sale_val:
                trv = lookup_ledger_row_by_management_id(df_ledger_hint, _sid_v)
                if trv is None:
                    st.error(f"管理ID {_sid_v} が台帳に見つかりません。")
                    validation_ok = False
                    break
                _st_v = _normalize_stock_status(str(trv.get(COL_STOCK_STATUS, "")))
                if _return_flow:
                    if _st_v != STATUS_SOLD:
                        st.error(
                            f"管理ID **{_sid_v}** は販売済ではありません（現在: {_st_v}）。"
                            "出庫（返品）は **販売済** の行のみ確定できます。"
                        )
                        validation_ok = False
                        break
                elif _st_v != STATUS_IN_STOCK:
                    st.error(
                        f"管理ID **{_sid_v}** は在庫中ではありません（現在: {_st_v}）。"
                        "販売済・対象外などへ変わっている可能性があります。"
                    )
                    validation_ok = False
                    break

        if validation_ok:
            urls: list[str] = [""]
            if uploaded is not None:
                with st.spinner("画像をリサイズしてドライブに保存しています…"):
                    raw0 = uploaded.getvalue()
                    try:
                        data0, mime0 = prepare_upload_image_jpeg(raw0)
                    except Exception as e:
                        st.error(f"画像の処理に失敗しました: {e}")
                        st.warning("画像なしで台帳の反映のみ続行します。")
                    else:
                        safe_base = re.sub(
                            r"[^\w\-_.]", "_", uploaded.name.rsplit(".", 1)[0]
                        )[:80]
                        fname0 = f"{jst_now().strftime('%Y%m%d_%H%M%S')}_{safe_base}_{uuid.uuid4().hex[:8]}.jpg"
                        try:
                            urls[0] = upload_image_to_drive(fname0, mime0, data0)
                        except Exception as e:
                            st.error(f"ドライブ保存に失敗しました: {e}")
                            st.warning("画像URLは付けずに反映のみ続行します。")
            if not _uses_local_inventory_csv() and not _secret_str(
                SECRET_GOOGLE_SPREADSHEET_ID
            ):
                st.warning("台帳の保存先が未設定のため、反映をスキップしました。")
            else:
                try:
                    _n_ids = len(_ids_sale_val)
                    _loan_manual = str(
                        st.session_state.get("sales_tab_loan_datetime_manual", "")
                        or ""
                    ).strip()
                    if _return_flow:
                        _spin_r = (
                            "返品を台帳に反映しています…"
                            if _n_ids <= 1
                            else f"**{_n_ids} 件** の返品を台帳に反映しています…"
                        )
                        with st.spinner(_spin_r):
                            for _sid_save in _ids_sale_val:
                                apply_outbound_sale_return_to_in_stock_by_management_id(
                                    _sid_save,
                                    new_image_url=(urls[0] if urls else "") or "",
                                    memo_suffix=memo_s,
                                )
                    elif _loan_keep_stock:
                        _spin = (
                            "浮貸日時を記録しています…"
                            if _n_ids <= 1
                            else f"在庫行 **{_n_ids} 件** に浮貸日時を記録しています…"
                        )
                        with st.spinner(_spin):
                            for _sid_save in _ids_sale_val:
                                apply_outbound_loan_in_stock_datetime_by_management_id(
                                    _sid_save,
                                    loan_datetime_jst=_loan_manual or None,
                                    new_image_url=(urls[0] if urls else "") or "",
                                    memo_suffix=memo_s,
                                )
                    elif _receipt_keep_stock:
                        _spin_rc = (
                            "戻入（在庫中）を記録しています…"
                            if _n_ids <= 1
                            else f"在庫行 **{_n_ids} 件** に戻入を記録しています…"
                        )
                        with st.spinner(_spin_rc):
                            for _sid_save in _ids_sale_val:
                                apply_outbound_receipt_in_stock_by_management_id(
                                    _sid_save,
                                    receipt_datetime_jst=_loan_manual or None,
                                    new_image_url=(urls[0] if urls else "") or "",
                                    memo_suffix=memo_s,
                                )
                    elif _receipt_as_sale:
                        _rdt = _loan_manual or jst_now_str()
                        _spin_rs = (
                            "戻入（販売済）を記録しています…"
                            if _n_ids <= 1
                            else f"在庫行 **{_n_ids} 件** を戻入（販売済）で更新しています…"
                        )
                        with st.spinner(_spin_rs):
                            for _sid_save in _ids_sale_val:
                                apply_outbound_sale_to_ledger_by_management_id(
                                    _sid_save,
                                    actual_sale_unit_excl_yen=_act_ex2,
                                    new_image_url=(urls[0] if urls else "") or "",
                                    memo_suffix=memo_s,
                                    sale_outbound_type="出庫（戻入）",
                                    loan_datetime_jst=_rdt,
                                )
                    elif _disposal_excluded:
                        _spin_dx = (
                            "除外（対象外）を台帳に反映しています…"
                            if _n_ids <= 1
                            else f"**{_n_ids} 件** の除外（対象外）を反映しています…"
                        )
                        with st.spinner(_spin_dx):
                            for _sid_save in _ids_sale_val:
                                apply_outbound_disposal_excluded_by_management_id(
                                    _sid_save,
                                    new_image_url=(urls[0] if urls else "") or "",
                                    memo_suffix=memo_s,
                                )
                    else:
                        if _loan_as_sale:
                            _ot = "出庫（浮貸）"
                        else:
                            _ot = (outbound_kind or "").strip() or "出庫（販売）"
                        _spin_sale = (
                            "該当の在庫行を販売済に更新しています…"
                            if _n_ids <= 1
                            else f"在庫行 **{_n_ids} 件** を順に販売済に更新しています…"
                        )
                        with st.spinner(_spin_sale):
                            for _sid_save in _ids_sale_val:
                                apply_outbound_sale_to_ledger_by_management_id(
                                    _sid_save,
                                    actual_sale_unit_excl_yen=_act_ex2,
                                    new_image_url=(urls[0] if urls else "") or "",
                                    memo_suffix=memo_s,
                                    sale_outbound_type=_ot,
                                )
                except Exception as e:
                    st.error(f"台帳の更新に失敗しました: {e}")
                else:
                    st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                    if _return_flow:
                        if len(_ids_sale_val) <= 1:
                            st.success(
                                f"管理ID **{_ids_sale_val[0]}** を **在庫中** に戻しました（出庫（返品）・販売日時・日時は確定の JST）。"
                            )
                        else:
                            st.success(
                                f"**{len(_ids_sale_val)} 件** を在庫中に戻しました（管理ID: {'、'.join(_ids_sale_val)}）。"
                            )
                    elif _receipt_keep_stock:
                        if len(_ids_sale_val) <= 1:
                            st.success(
                                f"管理ID **{_ids_sale_val[0]}** に **出庫（戻入）** と **{COL_LOAN_DATETIME}** を記録しました（在庫中のまま）。"
                            )
                        else:
                            st.success(
                                f"**{len(_ids_sale_val)} 件** に戻入を記録しました（管理ID: {'、'.join(_ids_sale_val)}）。"
                            )
                    elif _receipt_as_sale:
                        if len(_ids_sale_val) <= 1:
                            st.success(
                                f"管理ID **{_ids_sale_val[0]}** を **出庫（戻入）** で販売済に更新しました（実売 ¥{_act_ex2:,}）。"
                            )
                        else:
                            _ids_join = "、".join(_ids_sale_val)
                            st.success(
                                f"**{len(_ids_sale_val)} 件** を出庫（戻入）で販売済に更新しました（管理ID: {_ids_join}）。"
                                f"実売（税抜・1点あたり）は各行 ¥{_act_ex2:,}。"
                            )
                    elif _disposal_excluded:
                        if len(_ids_sale_val) <= 1:
                            st.success(
                                f"管理ID **{_ids_sale_val[0]}** を **対象外** に更新しました（{OUTBOUND_KIND_EXCLUDE}・日時は確定の JST）。"
                            )
                        else:
                            st.success(
                                f"**{len(_ids_sale_val)} 件** を対象外に更新しました（管理ID: {'、'.join(_ids_sale_val)}）。"
                            )
                    elif _loan_keep_stock:
                        if len(_ids_sale_val) <= 1:
                            st.success(
                                f"管理ID **{_ids_sale_val[0]}** に **{COL_LOAN_DATETIME}** を記録しました（在庫中のまま）。"
                            )
                        else:
                            st.success(
                                f"**{len(_ids_sale_val)} 件** に {COL_LOAN_DATETIME} を記録しました（管理ID: {'、'.join(_ids_sale_val)}）。"
                            )
                    elif len(_ids_sale_val) <= 1:
                        _sid_one = _ids_sale_val[0]
                        st.success(
                            f"管理ID **{_sid_one}** の行を販売済に更新しました（実売 ¥{_act_ex2:,}・販売日時は確定実行の JST）。"
                        )
                    else:
                        _ids_join = "、".join(_ids_sale_val)
                        st.success(
                            f"**{len(_ids_sale_val)} 件** の在庫行を販売済に更新しました（管理ID: {_ids_join}）。"
                            f"実売（税抜・1点あたり）は各行 ¥{_act_ex2:,}、販売日時は確定実行の JST を記録しています。"
                        )
                    if urls[0]:
                        st.markdown(f"[保存した画像を開く]({urls[0]})")
                    st.balloons()


def main():
    st.set_page_config(page_title="商品在庫・販売", layout="wide")
    _nav_opts = (
        "登録（インプット）",
        "ギャラリー（カタログ）",
        "在庫一覧",
        "集計・分析（ダッシュボード）",
    )
    _nav_legacy = "ギャラリー（カタログ）・在庫一覧"
    if "nav_page" not in st.session_state:
        st.session_state.nav_page = _nav_opts[0]
    elif st.session_state.nav_page == _nav_legacy:
        st.session_state.nav_page = "ギャラリー（カタログ）"
    elif st.session_state.nav_page not in _nav_opts:
        st.session_state.nav_page = _nav_opts[0]
    with st.sidebar:
        st.markdown("### メニュー")
        page = st.radio("ページ", _nav_opts, key="nav_page")
    st.title("商品在庫・販売管理")
    st.caption(
        "写真は任意。台帳の必須項目のみの記録、または写真＋AI解析・ドライブ保存・"
        "**inventory.csv** またはスプレッドシートへの記録ができます。"
    )
    if page == "ギャラリー（カタログ）":
        render_inventory_list_page(view_mode="gallery")
        return
    if page == "在庫一覧":
        render_inventory_list_page(view_mode="table")
        return
    if page == "集計・分析（ダッシュボード）":
        render_analytics_dashboard_page()
        return

    _init_registration_form_session_state()
    _init_voucher_sidebar_state()
    df_ledger_hint = _ledger_hint_dataframe()
    _ledger_err = st.session_state.pop("_ledger_hint_load_error", None)
    _sheet_err = st.session_state.pop("_inventory_sheet_load_error", None)
    if _ledger_err or _sheet_err:
        st.warning(
            "台帳の参照用データの読み込みに失敗しました（**仕入のAI解析・販売の写真照合・候補一覧**に影響します）。"
            f"\n\n{_ledger_err or _sheet_err}"
        )

    st.markdown("## 台帳登録")
    st.caption(
        "仕入れ・販売・棚卸しは下のタブで切り替えます。写真は1枚を全タブで共通利用します。"
    )
    uploaded = st.file_uploader(
        "商品写真（任意・1枚まで・カメラやギャラリーから）",
        type=["jpg", "jpeg", "png", "webp"],
        key="shared_reg_photo_uploader",
    )
    st.caption("写真は1枚まで。複数行登録時は同じ画像URLを各行に保存します。")

    _inject_prominent_main_tabs_style()
    st.markdown("## 入力モード")
    st.caption("タブで入力モードを切り替えます。")
    tab_purchase, tab_sales, tab_stock = st.tabs(
        ("仕入れ登録", "販売管理", "棚卸し登録")
    )

    with tab_purchase:
        _render_voucher_inventory_panel()
        st.divider()
        st.markdown("##### クイック検索（写真から検索）")
        st.caption(
            "**AIで画像を解析** で商品名・色・シルエットなどを推定しつつ在庫と照合します（洋服・帽子・雑貨など **和装以外も** 想定）。"
            "パッケージに大きな文字がある商品より、**無包装の衣料**は判別が難しいことがあります。**タグ・下札**が写っていると有利です。"
            "解析後は下の「近い候補」も併せて確認してください。"
        )

        movement = st.radio(
            "区分（仕入れ・在庫の増減）",
            ("入庫（購入）", "入庫（浮貸）"),
            horizontal=True,
            key="tab_purchase_movement",
        )
    
        col_a, col_c = st.columns([1, 1])
        with col_a:
            analyze = st.button(
                "AIで画像を解析",
                type="primary",
                disabled=uploaded is None,
            )
        with col_c:
            if st.button("候補の自動入力をクリア"):
                st.session_state.field_product_name = ""
                st.session_state.field_supplier = ""
                st.session_state.field_row_quantity = 1
                st.session_state.field_inventory_category = ""
                st.session_state.ai_kind = ""
                st.session_state.ai_features = ""
                st.session_state.ai_parse_ran = False
                st.session_state.field_memo = ""
                st.session_state.field_line_excl_yen = 1
                st.session_state.field_planned_sale_excl = 0
                st.session_state.field_actual_sale_excl = 0
                st.session_state.field_stock_status = STATUS_IN_STOCK
                st.session_state.hint_filter_product_name = ""
                st.session_state.hint_filter_supplier = ""
                st.session_state.hint_filter_inventory_category = ""
                st.session_state.hint_filter_management_id = ""
                st.session_state.ledger_pick_product_name = LEDGER_PICK_PLACEHOLDER
                st.session_state.ledger_pick_supplier = LEDGER_PICK_PLACEHOLDER
                st.session_state.ledger_pick_inventory_category = LEDGER_PICK_PLACEHOLDER
                st.session_state.ledger_pick_management_id = LEDGER_PICK_PLACEHOLDER
                st.session_state.field_sale_source_mgmt_id = ""
                st.session_state.sale_pick_source_id = LEDGER_PICK_PLACEHOLDER
                st.session_state.pop("ledger_quick_candidates", None)
                st.session_state.pop("_gemini_match_management_id", None)
                st.session_state.pop("_sale_link_management_id", None)
                st.session_state.pop("_sale_link_warn", None)
                st.rerun()
    
        if analyze and uploaded is not None:
            with st.spinner("画像を解析しています…"):
                try:
                    inv_ctx = ""
                    if df_ledger_hint is not None and not df_ledger_hint.empty:
                        inv_ctx = _build_gemini_inventory_context(
                            df_ledger_hint, only_in_stock=False
                        )
                    raw_text = analyze_image_with_gemini(
                        uploaded,
                        inventory_context=inv_ctx or None,
                    )
                    result = _parse_json_from_model(raw_text or "")
                    result = _apply_purchase_ledger_match_supplement(
                        result, df_ledger_hint
                    )
                    _apply_gemini_json_to_session(result, df_ledger_hint)
                    _refresh_ledger_quick_search_candidates(df_ledger_hint)
                    st.success(
                        "解析が完了しました。必要に応じて商品名・仕入先・在庫カテゴリー・仕入金額（税抜）を修正してください。"
                    )
                except Exception as e:
                    st.warning(
                        "現在混み合っているか、無料枠の上限に達している可能性があります。"
                        "1分ほど待ってから再試行してください。"
                    )
                    st.caption(f"詳細: {e}")
    
        if st.session_state.get("ai_parse_ran"):
            st.subheader("AI解析結果（参考）")
            st.write(f"**推定種類:** {st.session_state.ai_kind or '—'}")
            st.write(
                f"**推定数量（同時追記行数）:** {int(st.session_state.field_row_quantity)}"
            )
            st.write(
                f"**推定仕入金額（税抜・1点）:** ¥{int(st.session_state.field_line_excl_yen):,}"
            )
            st.write(
                f"**推定在庫カテゴリー:** {str(st.session_state.get('field_inventory_category', '') or '').strip() or '—'}"
            )
            mid_hit = st.session_state.get("_gemini_match_management_id")
            if mid_hit:
                st.info(
                    f"台帳照合: 管理ID **{mid_hit}** の在庫行に合わせて、商品名・仕入先・仕入金額（税抜）"
                    f"・{COL_CATEGORY}を反映しました。"
                )
    
        if "purchase_assist_visible" not in st.session_state:
            st.session_state.purchase_assist_visible = False
        if st.button(
            "台帳からの入力補助を表示"
            if not st.session_state.purchase_assist_visible
            else "台帳からの入力補助を非表示",
            key="purchase_assist_toggle_btn",
        ):
            st.session_state.purchase_assist_visible = (
                not st.session_state.purchase_assist_visible
            )
            st.rerun()

        if (
            st.session_state.purchase_assist_visible
            and df_ledger_hint is not None
            and not df_ledger_hint.empty
        ):
            _render_ledger_pick_assist_three_columns(
                df_ledger_hint,
                key_prefix="",
                body_caption=(
                    "絞り込み欄に文字を入れると候補が絞られます。商品名・仕入先・在庫カテゴリーは個別にプルダウンで選べます。"
                    "**管理ID** をプルダウンで選ぶと、その台帳1行から **商品名・仕入先・カテゴリー・仕入金額（税抜）・販売予定（税抜）** を一度に反映します。"
                    "在庫中の行に一致したときは **販売予定金額（税抜・任意）** にも、台帳の1点あたりの値を入れます（選んだ行の値をそのまま反映）。"
                    f"**{COL_CATEGORY}** も同様に、台帳の既存値から選べます。"
                ),
                on_pick_product_name=_on_ledger_pick_product_name,
                on_pick_supplier=_on_ledger_pick_supplier,
                on_pick_inventory_category=_on_ledger_pick_inventory_category,
                on_pick_management_id=_on_ledger_pick_management_id,
            )
        elif (
            st.session_state.purchase_assist_visible
            and (_uses_local_inventory_csv() or _secret_str(SECRET_GOOGLE_SPREADSHEET_ID))
        ):
            st.caption("台帳が空か読み込めないため、入力補助の候補は表示できません。")
    
        st.markdown("##### 必須入力項目")
        st.caption(
            "このタブの確定は **在庫中** の行の追加（入庫（購入）／入庫（浮貸））です。"
            "顧客返品で在庫へ戻す操作は **販売管理** の **出庫（返品）** を使用してください。"
            "**出庫（浮貸）・出庫（販売）・出庫（除外）・出庫（返品）・出庫（戻入）** は **販売管理** タブで行ってください。"
        )
        product_name = st.text_input("商品名（必須）", key="field_product_name")
        supplier = st.text_input("仕入先・取引先（必須）", key="field_supplier")
        st.text_input(
            "在庫カテゴリー（必須）",
            key="field_inventory_category",
            placeholder="例: 帯 / 雑貨 / 飲料（上の台帳補助・AI解析で入力可）",
        )
        _refresh_ledger_quick_search_candidates(df_ledger_hint)
        _cand = st.session_state.get("ledger_quick_candidates")
        if (
            isinstance(_cand, pd.DataFrame)
            and not _cand.empty
            and df_ledger_hint is not None
        ):
            with st.expander("近い候補（写真解析・入力文字から照合・カード）", expanded=False):
                st.caption(
                    "Gemini で管理IDが一致しないときは、**名前＋仕入先** → **名前** → **仕入先** の部分一致で候補を出します。"
                    "カードから選ぶと、仕入入力の必須項目へ反映されます（在庫中・販売済どちらも候補対象）。"
                )
                _p_hits: list[dict[str, Any]] = []
                for _, row in _cand.iterrows():
                    _p_hits.append(
                        _sale_card_hit_from_series(row)
                    )
                _render_mid_pick_candidate_cards(
                    _p_hits,
                    widget_key_namespace="purchase_quick_cards",
                    sold=False,
                    pick_mode="purchase",
                    page_size=5,
                )
    
        st.number_input(
            f"数量（{COL_QTY}・同時追記行数）",
            min_value=1,
            max_value=2000,
            step=1,
            key="field_row_quantity",
            help=(
                f"在庫として計上する点数です。**N にすると同一内容のレコードを N 行** 追記します（1点＝1行）。"
                f"各行の「{COL_QTY}」列には **1** を入れます。仕入金額（税抜）は各行とも **1点あたり** の金額です。"
                "写真をアップロードした場合、複数行でも **同じ画像URL** を各行に記録します。"
                "「AIで画像を解析」の推定数量がここに入ります（いつでも上書きできます）。"
            ),
        )
        st.caption(
            "2 以上にすると、同じ商品名・仕入先・単価で **管理IDだけ異なる複数行** を一度に作成します。"
        )

        line_excl_yen = st.number_input(
            "仕入金額（税抜・必須）",
            min_value=1,
            step=1,
            key="field_line_excl_yen",
            help="1点あたりの税抜の仕入金額（円）。台帳の各行は1点あたりの金額として保存されます。",
        )
    
        st.radio(
            "消費税（仕入金額（税込）の計算）",
            options=list(CONSUMPTION_TAX_CHOICE_TO_RATE.keys()),
            horizontal=True,
            key="field_consumption_tax_choice",
            help="仕入金額（税抜）の税込行計に使用します。非課税のときは税込＝税抜です。",
        )
        _tax_r = _consumption_tax_rate_from_choice_label(
            str(st.session_state.get("field_consumption_tax_choice", "10%"))
        )
    
        _n_save = max(1, min(2000, int(st.session_state.get("field_row_quantity", 1))))
        _lex_inp = int(line_excl_yen)
        _line_ex_one = _lex_inp
        _line_in_one = price_incl_tax(_line_ex_one, _tax_r)
    
        price_row = st.columns([1, 1, 1])
        with price_row[0]:
            st.metric("仕入金額（税抜・1点）", f"¥{_line_ex_one:,}")
            _cap_rows = (
                f"確定時は **{_n_save} 行**（各行「{COL_QTY}」**1**）。税抜合計（参考） ¥{_line_ex_one * _n_save:,}。"
            )
            if _n_save > 1:
                _cap_rows += (
                    "写真があるとき、複数行なら **同じ画像URLを全行** に記録します。"
                )
            st.caption(_cap_rows)
        with price_row[1]:
            st.metric("仕入金額（税込・1点・自動）", f"¥{_line_in_one:,}")
            _tl = st.session_state.get("field_consumption_tax_choice", "10%")
            if _tl == "非課税":
                st.caption("非課税のため税込＝税抜行合計")
            else:
                st.caption(f"消費税{_tl}を行合計に四捨五入")
        with price_row[2]:
            st.caption(
                "原価は各行の仕入金額（税抜）です。販売予定・実売の詳細は下の **価格管理／販売管理** で入力します。"
            )
    
        st.markdown("##### 価格管理（任意）")
        st.caption(
            "「販売予定金額（税抜）」は **1点あたりの税抜金額（円）** です。"
            "「在庫中」のときは販売予定行計−原価で粗利の参考になります（下のプレビュー）。"
        )
        planned_sale_excl = st.number_input(
            "販売予定金額（税抜・任意）",
            min_value=0,
            step=1,
            key="field_planned_sale_excl",
            help="1点あたり。0 のとき台帳では空欄。税抜行計・税込総額は各行1点として自動計算します。",
        )
    
        st.markdown("##### 販売・実売について")
        st.caption(
            "通常は **新規行の追加**（入庫（購入）など）です。"
            "**実売・販売済の更新** は **販売管理** タブを使用してください。"
        )
        _pl_u = int(planned_sale_excl)
        _act_u = 0
        _cogs_preview = _lex_inp * _n_save
        _pl_u_gp = _pl_u
        _tax_preview = _tax_r
        _st_gp = STATUS_IN_STOCK
        _plex, _pin, _aex, _ain = _planned_actual_line_amounts(
            _n_save, _pl_u_gp, _act_u, _st_gp, _tax_preview
        )
        _gp_preview = _compute_gross_profit_row(
            _cogs_preview,
            _plex,
            0,
            _st_gp,
        )
        pm1, pm2, pm3, pm4, pm5 = st.columns(5)
        with pm1:
            _cogs_lbl = "原価（税抜・行合計）"
            st.metric(_cogs_lbl, f"¥{_cogs_preview:,}")
        with pm2:
            st.metric(
                "販売予定（税抜・行計）",
                "—" if _plex <= 0 else f"¥{_plex:,}",
            )
        with pm3:
            st.metric(
                "販売予定（税込・総額）",
                "—" if _pin <= 0 else f"¥{_pin:,}",
            )
        with pm4:
            st.metric(
                "実売（税抜・行計）",
                "—" if _aex <= 0 else f"¥{_aex:,}",
            )
        with pm5:
            st.metric(
                "実売（税込・総額）",
                "—" if _ain <= 0 else f"¥{_ain:,}",
            )
        pm6, _, _ = st.columns([1, 1, 3])
        with pm6:
            st.metric(
                "粗利（税抜・プレビュー）",
                "—" if _gp_preview is None else f"¥{int(_gp_preview):,}",
            )
    
        st.markdown("##### 補足情報（任意）")
        memo = st.text_area(
            "メモ（任意）",
            key="field_memo",
            height=100,
            placeholder="備考・社内メモなどがあれば入力してください",
        )
    
        confirm = st.button(
            "確定（台帳に記録・写真は任意でドライブ保存）",
            type="primary",
        )
    
        if confirm:
            validation_ok = True
            _act_ex2 = 0
            if not (product_name or "").strip():
                st.error("商品名を入力してください。")
                validation_ok = False
            elif not (supplier or "").strip():
                st.error("仕入先・取引先を入力してください。")
                validation_ok = False
            elif int(line_excl_yen) < 1:
                st.error("仕入金額（税抜）を1円以上で入力してください。")
                validation_ok = False
            elif not str(
                st.session_state.get("field_inventory_category", "") or ""
            ).strip():
                st.error("在庫カテゴリーを入力してください。")
                validation_ok = False

            if validation_ok:
                _lex_one = int(line_excl_yen)
                _tax_r2 = _consumption_tax_rate_from_choice_label(
                    str(st.session_state.get("field_consumption_tax_choice", "10%"))
                )
                _lin_one = price_incl_tax(_lex_one, _tax_r2)
                _plan2 = int(st.session_state.get("field_planned_sale_excl", 0))
                _stat2 = STATUS_IN_STOCK
                memo_s = (memo or "").strip()
                _rq2 = max(1, min(2000, int(st.session_state.get("field_row_quantity", 1))))
                _icat2 = str(
                    st.session_state.get("field_inventory_category", "") or ""
                ).strip()
                n_save = _rq2

                if validation_ok:
                    urls: list[str] = [""] * n_save
                    _record_dt = jst_now_str()
                    ready_for_sheet = True

                    if uploaded is not None:
                        with st.spinner("画像をリサイズ・圧縮してドライブに保存しています…"):
                            raw0 = uploaded.getvalue()
                            _record_dt = (
                                capture_datetime_jst_from_bytes(raw0) or _record_dt
                            )
                            try:
                                data0, mime0 = prepare_upload_image_jpeg(
                                    raw0,
                                    max_long_edge=PURCHASE_DRIVE_JPEG_MAX_LONG_EDGE,
                                    quality=PURCHASE_DRIVE_JPEG_QUALITY,
                                )
                            except Exception as e:
                                st.error(f"画像の処理に失敗しました: {e}")
                                ready_for_sheet = False
                            else:
                                safe_base = re.sub(
                                    r"[^\w\-_.]", "_", uploaded.name.rsplit(".", 1)[0]
                                )[:80]
                                fname0 = f"{jst_now().strftime('%Y%m%d_%H%M%S')}_{safe_base}_{uuid.uuid4().hex[:8]}.jpg"
                                try:
                                    shared_url = upload_image_to_drive(fname0, mime0, data0)
                                except Exception as e:
                                    st.error(f"ドライブ保存に失敗しました: {e}")
                                    ready_for_sheet = False
                                else:
                                    urls = [shared_url] * n_save

                    if ready_for_sheet:
                        ws0 = (
                            None
                            if _uses_local_inventory_csv()
                            else ensure_worksheet_header()
                        )
                        if not _uses_local_inventory_csv() and ws0 is None:
                            st.warning(
                                "スプレッドシート未設定のため、行の追記をスキップしました。"
                            )
                        else:
                            try:
                                ids = allocate_management_ids(ws0, n_save)
                                _spin_msg = (
                                    "inventory.csv に記録しています…"
                                    if _uses_local_inventory_csv()
                                    else "スプレッドシートに記録しています…"
                                )
                                with st.spinner(_spin_msg):
                                    _reg_row_dt = jst_now_str()
                                    _reg_batch: list[list[Any]] = []
                                    for i in range(n_save):
                                        _reg_batch.append(
                                            _inventory_row_values_for_append(
                                                _reg_row_dt,
                                                _record_dt,
                                                movement,
                                                product_name.strip(),
                                                supplier.strip(),
                                                _lex_one,
                                                _lin_one,
                                                urls[i],
                                                ids[i],
                                                memo_s,
                                                quantity=1,
                                                inventory_category=_icat2,
                                                planned_sale_unit_excl_yen=_plan2,
                                                actual_sale_unit_excl_yen=_act_ex2,
                                                stock_status=_stat2,
                                                consumption_tax_rate=_tax_r2,
                                            )
                                        )
                                    _append_inventory_data_rows(_reg_batch)
                            except Exception as e:
                                st.error(f"台帳の更新に失敗しました: {e}")
                                if any(urls):
                                    st.warning(
                                        "一部の画像はドライブに保存済みの可能性があります。台帳の内容を確認してください。"
                                    )
                            else:
                                st.session_state.pop(LEDGER_DATA_EDITOR_KEY, None)
                                st.success(
                                    f"記録しました（{n_save} 行・1点1行）。管理IDを自動付与しています。"
                                )
                                _link_urls = list(dict.fromkeys(u for u in urls if u))
                                for _uurl in _link_urls[:8]:
                                    st.markdown(f"[保存した画像を開く]({_uurl})")
                                if len(_link_urls) > 8:
                                    st.caption(
                                        f"ほか {len(_link_urls) - 8} 件の画像URLは台帳の「{COL_IMAGE_URL}」列を参照してください。"
                                    )
                                st.balloons()

    with tab_sales:
        _render_sales_management_tab(uploaded, df_ledger_hint)

    with tab_stock:
        render_stocktake_scan_tab(uploaded, df_ledger_hint)


if __name__ == "__main__":
    if check_password():
        main()
